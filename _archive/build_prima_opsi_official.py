from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DOCX = Path("Proposal_Penelitian_PRIMA_OPSI_2026_Format_Resmi.docx")
OUT_TABLE = Path("Tabel_Revisi_Sistematika_OPSI_PRIMA.docx")
FLOWCHART = Path("prima_prosedur_ADDIE_HIGH.png")


REFERENCES = [
    "Al-Dosakee, K., & Ozdamli, F. (2021). Gamification in teaching and learning languages: A systematic literature review. Revista Romaneasca pentru Educatie Multidimensionala, 13(2), 559-577. https://doi.org/10.18662/rrem/13.2/436",
    "Amelia, S. R., Siroj, M. B., & Afgani, M. W. (2024). Analisis perkembangan bahasa Indonesia di era digital: Tantangan dan peluang. Jurnal Multidisiplin Dehasen, 3(3), 125-132. https://doi.org/10.37676/mude.v3i3.5792",
    "Balai Bahasa Provinsi Jambi. (2025). UKBI jadi instrumen strategis wujudkan kemahiran berbahasa Indonesia yang terstandar. Kementerian Pendidikan Dasar dan Menengah.",
    "Branch, R. M. (2009). Instructional Design: The ADDIE Approach. Springer. https://doi.org/10.1007/978-0-387-09506-6",
    "Carter, R. (2003). Language awareness. ELT Journal, 57(1), 64-65. https://doi.org/10.1093/elt/57.1.64",
    "Chaer, A., & Agustina, L. (2014). Sosiolinguistik: Perkenalan awal. Rineka Cipta.",
    "Crystal, D. (2011). Internet Linguistics: A Student Guide. Routledge.",
    "Fairclough, N. (1992). Critical Language Awareness. Longman.",
    "Garvin, P. L., & Mathiot, M. (1968). The urbanization of the Guarani language: A problem in language and culture. In J. A. Fishman (Ed.), Readings in the Sociology of Language (pp. 365-374). De Gruyter Mouton. https://doi.org/10.1515/9783110805376.365",
    "Luo, Z. (2023). The effectiveness of gamified tools for foreign language learning (FLL): A systematic review. Behavioral Sciences, 13(4), 331. https://doi.org/10.3390/bs13040331",
    "Prensky, M. (2001). Digital natives, digital immigrants. On the Horizon, 9(5), 1-6.",
    "Purba, E. N., Togatorop, D. P., Simbolon, A., & Sari, Y. (2024). Analisis pengaruh media sosial terhadap keberagaman bahasa: Campur kode sebagai tren komunikasi anak muda. Atmosfer: Jurnal Pendidikan, Bahasa, Sastra, Seni, Budaya, dan Sosial Humaniora, 2(4), 184-194. https://doi.org/10.59024/atmosfer.v2i4.1060",
    "Shortt, M., Tilak, S., Kuznetcova, I., Martens, B., & Akinkuolie, B. (2021). Gamification in mobile-assisted language learning: A systematic review of Duolingo literature from public release of 2012 to early 2020. Computer Assisted Language Learning. https://doi.org/10.1080/09588221.2021.1933540",
    "Sugiyono. (2019). Metode Penelitian dan Pengembangan: Research and Development. Alfabeta.",
]


def set_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in [("Heading 1", 16), ("Heading 2", 12), ("Heading 3", 11)]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True


def add_center(doc, text, size=12, bold=False, after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    set_font(r)
    return p


def add_plain(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r)
    return p


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        if p.runs:
            p.runs[0].font.name = "Times New Roman"
        r = p.add_run(item)
        set_font(r)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        set_font(r)


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def write_cell(cell, text, bold=False, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if align:
        p.alignment = align
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc, headers, rows, widths=None, size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        write_cell(t.rows[0].cells[i], header, bold=True, size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_values in rows:
        row = t.add_row()
        for i, value in enumerate(row_values):
            write_cell(row.cells[i], value, size=size)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return t


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size=16 if level == 1 else 12, bold=True)
    return p


def cover(doc):
    add_center(doc, "PROPOSAL PENELITIAN OPSI", size=14, bold=True, after=16)
    add_center(
        doc,
        "PENGEMBANGAN PLATFORM PRIMA+ BERBASIS KESADARAN BERBAHASA\nUNTUK MENGUATKAN LOYALITAS BAHASA INDONESIA REMAJA",
        size=14,
        bold=True,
        after=20,
    )
    add_center(doc, "Nama Tim Peneliti", size=12, bold=True)
    add_center(doc, "Ketua : Larisa Dianti", size=12)
    add_center(doc, "Anggota : Gusti Nazwa Azizah", size=12, after=14)
    add_center(doc, "Bidang Kompetisi Penelitian: Ilmu Sosial dan Humaniora (ISH)", size=12, bold=True)
    add_center(doc, "Subbidang: Bahasa dan Sastra", size=12)
    add_center(doc, "MAN KOTAWARINGIN TIMUR", size=12, bold=True)
    add_center(doc, "Kotawaringin Timur, Kalimantan Tengah", size=12)
    add_center(doc, "Tahun 2026", size=12, bold=True)
    doc.add_page_break()


def front_matter(doc):
    add_heading(doc, "DAFTAR ISI", level=1)
    for line in [
        "BAB 1. PENDAHULUAN",
        "1.1 Latar Belakang",
        "1.2 Rumusan Masalah",
        "1.3 Tujuan Penelitian",
        "1.4 Hipotesis",
        "1.5 Manfaat Penelitian",
        "BAB 2. TINJAUAN PUSTAKA",
        "2.1 Landasan Teori",
        "2.2 Studi Pustaka",
        "BAB 3. METODE PENELITIAN",
        "3.1 Waktu dan Tempat Penelitian",
        "3.2 Alat dan Bahan",
        "3.3 Rancangan dan Prosedur Penelitian",
        "3.4 Rancangan Pengolahan Data",
        "BAB 4. RANCANGAN ANGGARAN BIAYA DAN JADWAL KEGIATAN",
        "DAFTAR PUSTAKA",
        "LAMPIRAN",
    ]:
        add_plain(doc, line)
    add_heading(doc, "DAFTAR TABEL", level=1)
    for line in [
        "Tabel 1. Studi pustaka terdahulu",
        "Tabel 2. Alat dan bahan penelitian",
        "Tabel 3. Konstruk, indikator, dan instrumen penelitian",
        "Tabel 4. Rancangan pengolahan data",
        "Tabel 5. Rancangan anggaran biaya penelitian",
        "Tabel 6. Jadwal penelitian",
    ]:
        add_plain(doc, line)
    add_heading(doc, "DAFTAR GAMBAR", level=1)
    add_plain(doc, "Gambar 1. Alur prosedur penelitian PRIMA+ berbasis ADDIE")
    doc.add_page_break()


def build_proposal():
    doc = Document()
    style_doc(doc)
    cover(doc)
    front_matter(doc)

    add_heading(doc, "BAB 1. PENDAHULUAN", level=1)
    add_heading(doc, "1.1 LATAR BELAKANG", level=2)
    for text in [
        "Di lingkungan sekolah, bahasa Indonesia hadir hampir di semua kegiatan siswa: belajar di kelas, berdiskusi, menulis tugas, mengikuti organisasi, hingga membuat unggahan di media sosial. Namun, cara siswa menggunakan bahasa tidak selalu sama pada setiap situasi. Ada saatnya mereka memakai bahasa santai, dan ada saatnya mereka perlu memakai bahasa Indonesia yang lebih tertib dan jelas.",
        "Di media sosial, remaja terbiasa menulis dengan cepat dan ringkas. Mereka memakai slang, emotikon, campur kode Indonesia-Inggris, serta istilah seperti happy, healing, deadline, vibes, atau relate. Gaya seperti ini tidak selalu salah karena bahasa berubah mengikuti lingkungan pemakainya. Masalah muncul ketika siswa tidak lagi peka terhadap perbedaan situasi. Bahasa untuk story Instagram tentu tidak selalu cocok dipakai untuk laporan, pengumuman kelas, atau komunikasi resmi di sekolah.",
        "Kebutuhan penguatan bahasa Indonesia pada pelajar juga tampak dari data kebahasaan. Balai Bahasa Provinsi Jambi (2025) menjelaskan bahwa UKBI digunakan sebagai alat untuk melihat kemahiran berbahasa Indonesia secara lebih terstandar. Data pelajar tahun 2024 yang disoroti dalam laporan tersebut menunjukkan bahwa pembinaan kemahiran berbahasa masih perlu diperkuat.",
        "Penelitian terdahulu telah banyak membahas campur kode, bahasa gaul, dan perubahan bahasa di media sosial. Di sisi lain, penelitian tentang gamifikasi bahasa lebih sering membahas pembelajaran bahasa asing. Masih terbatas model yang secara khusus menghubungkan kesadaran berbahasa, loyalitas bahasa Indonesia, dan media digital interaktif untuk remaja. Oleh karena itu, penelitian ini mengembangkan PRIMA+ sebagai platform digital yang memuat kasus bahasa digital, kuis kontekstual, refleksi pilihan bahasa, skor, dan umpan balik.",
        "Kebaruan penelitian ini terletak pada pengembangan media digital berbasis kesadaran berbahasa untuk menguatkan loyalitas bahasa Indonesia remaja. Temuan yang ditargetkan adalah rancangan platform PRIMA+ yang layak digunakan dan data awal mengenai efektivitasnya dalam meningkatkan loyalitas berbahasa Indonesia.",
    ]:
        add_para(doc, text)

    add_heading(doc, "1.2 RUMUSAN MASALAH", level=2)
    add_numbered(
        doc,
        [
            "Bagaimana rancangan platform PRIMA+ berbasis kesadaran berbahasa untuk menguatkan loyalitas bahasa Indonesia remaja?",
            "Bagaimana kelayakan materi, bahasa, media, dan instrumen PRIMA+ berdasarkan validasi guru atau ahli?",
            "Bagaimana efektivitas penggunaan PRIMA+ terhadap peningkatan loyalitas berbahasa Indonesia remaja?",
        ],
    )

    add_heading(doc, "1.3 TUJUAN PENELITIAN", level=2)
    add_numbered(
        doc,
        [
            "Mengembangkan platform PRIMA+ berbasis kesadaran berbahasa sesuai kebutuhan remaja pengguna media digital.",
            "Menilai kelayakan PRIMA+ dari aspek materi, bahasa, media, dan instrumen penelitian.",
            "Menguji efektivitas PRIMA+ dalam menguatkan loyalitas berbahasa Indonesia remaja melalui desain pretest-posttest.",
        ],
    )

    add_heading(doc, "1.4 HIPOTESIS", level=2)
    add_plain(doc, "H0: Tidak terdapat peningkatan loyalitas berbahasa Indonesia remaja setelah menggunakan platform PRIMA+.")
    add_plain(doc, "H1: Terdapat peningkatan loyalitas berbahasa Indonesia remaja setelah menggunakan platform PRIMA+.")

    add_heading(doc, "1.5 MANFAAT PENELITIAN", level=2)
    add_bullets(
        doc,
        [
            "Bagi siswa, PRIMA+ membantu mereka lebih peka saat memilih bahasa sesuai konteks sekolah dan media digital.",
            "Bagi guru Bahasa Indonesia, PRIMA+ dapat menjadi contoh media pembelajaran yang mengaitkan materi bahasa dengan situasi nyata siswa.",
            "Bagi sekolah, penelitian ini mendukung program literasi dan pembiasaan berbahasa Indonesia yang dekat dengan keseharian siswa.",
            "Bagi kajian sosiolinguistik terapan, penelitian ini memberi contoh pengembangan media digital untuk membangun sikap dan kesadaran berbahasa.",
        ],
    )

    add_heading(doc, "BAB 2. TINJAUAN PUSTAKA", level=1)
    add_heading(doc, "2.1 LANDASAN TEORI", level=2)
    add_para(doc, "Komunikasi remaja di ruang digital berlangsung cepat dan sering memadukan teks, gambar, suara, emotikon, serta video. Crystal (2011) menjelaskan bahwa teknologi digital dapat mengubah praktik bahasa, pola interaksi, dan norma komunikasi. Dalam konteks Indonesia, bahasa gaul dan campur kode di media sosial perlu dilihat sebagai gejala sosial sekaligus tantangan pembinaan bahasa Indonesia.")
    add_para(doc, "Kesadaran berbahasa menekankan kemampuan seseorang memperhatikan bentuk, fungsi, konteks, dan akibat sosial dari bahasa yang digunakan. Carter (2003) memandang language awareness sebagai kesadaran terhadap bentuk dan fungsi bahasa, sedangkan Fairclough (1992) menekankan pentingnya sikap kritis dalam melihat hubungan bahasa dan kehidupan sosial.")
    add_para(doc, "Loyalitas bahasa dapat dipahami sebagai kemauan untuk tetap menggunakan dan menjaga bahasa sebagai bagian dari identitas sosial. Garvin dan Mathiot (1968) menempatkan loyalitas bahasa sebagai salah satu unsur sikap bahasa, bersama kebanggaan bahasa dan kesadaran terhadap norma. Dalam penelitian ini, loyalitas berbahasa Indonesia dimaknai sebagai kemauan menempatkan bahasa Indonesia secara tepat, terutama pada situasi yang membutuhkan bahasa yang jelas dan tertib.")
    add_para(doc, "Gamifikasi memanfaatkan unsur permainan seperti poin, tantangan, misi, umpan balik, dan capaian agar proses belajar terasa lebih menarik. Beberapa tinjauan sistematis menunjukkan bahwa gamifikasi dapat mendukung motivasi belajar bahasa, tetapi hasilnya bergantung pada desain pembelajaran, kualitas umpan balik, dan cara mengukur capaian belajar (Al-Dosakee & Ozdamli, 2021; Luo, 2023; Shortt et al., 2021).")

    add_heading(doc, "2.2 STUDI PUSTAKA", level=2)
    add_table(
        doc,
        ["No.", "Sumber", "Temuan Utama", "Posisi terhadap PRIMA+"],
        [
            ["1", "Crystal (2011)", "Teknologi digital mengubah praktik dan norma komunikasi.", "Menjadi dasar fenomena bahasa digital remaja."],
            ["2", "Carter (2003); Fairclough (1992)", "Kesadaran berbahasa berkaitan dengan bentuk, fungsi, konteks, dan sikap kritis.", "Menjadi dasar pendekatan language awareness."],
            ["3", "Garvin & Mathiot (1968)", "Loyalitas bahasa merupakan bagian dari sikap bahasa.", "Menjadi dasar konstruk loyalitas berbahasa Indonesia."],
            ["4", "Shortt dkk. (2021); Luo (2023)", "Gamifikasi dapat mendukung pembelajaran bahasa jika desain dan evaluasinya jelas.", "Menjadi dasar rancangan PRIMA+ sebagai media interaktif."],
            ["5", "Amelia dkk. (2024); Purba dkk. (2024)", "Bahasa digital dan campur kode menjadi gejala komunikasi remaja.", "Memperkuat konteks masalah penelitian."],
        ],
        widths=[1.0, 3.8, 5.6, 5.4],
        size=8,
    )
    add_para(doc, "Berdasarkan studi pustaka tersebut, penelitian PRIMA+ memosisikan diri pada pengembangan media digital yang menggabungkan kesadaran berbahasa, gamifikasi, dan penguatan loyalitas bahasa Indonesia. Dengan demikian, penelitian ini tidak hanya mendeskripsikan fenomena campur kode, tetapi menawarkan media interaktif untuk melatih pengambilan keputusan berbahasa.")

    add_heading(doc, "BAB 3. METODE PENELITIAN", level=1)
    add_heading(doc, "3.1 WAKTU DAN TEMPAT PENELITIAN", level=2)
    add_para(doc, "Penelitian direncanakan berlangsung pada April-Agustus 2026 di MAN Kotawaringin Timur, Kalimantan Tengah. Kegiatan penelitian meliputi analisis kebutuhan, perancangan, pengembangan prototipe, validasi, uji coba terbatas, analisis data, revisi produk, dan penyusunan laporan akhir.")

    add_heading(doc, "3.2 ALAT DAN BAHAN", level=2)
    add_table(
        doc,
        ["No.", "Alat/Bahan", "Jumlah", "Fungsi"],
        [
            ["1", "Laptop pengembang", "1 unit", "Mendesain platform, mengolah data, dan menyusun laporan."],
            ["2", "Perangkat siswa", "30-40 perangkat", "Menguji akses dan penggunaan PRIMA+."],
            ["3", "Godot Engine atau platform web sederhana", "1 sistem", "Membangun prototipe interaktif."],
            ["4", "Skenario kasus bahasa digital", "1 paket", "Materi kuis dan refleksi bahasa."],
            ["5", "Kuesioner loyalitas berbahasa", "20 butir", "Mengukur konstruk loyalitas berbahasa Indonesia."],
            ["6", "Lembar validasi guru/ahli", "2-3 validator", "Menilai kelayakan materi, bahasa, media, dan instrumen."],
        ],
        widths=[1.0, 4.8, 3.0, 7.0],
        size=8,
    )

    add_heading(doc, "3.3 RANCANGAN DAN PROSEDUR PENELITIAN", level=2)
    add_para(doc, "Penelitian menggunakan metode Research and Development (R&D) dengan model ADDIE, yaitu Analyze, Design, Develop, Implement, dan Evaluate (Branch, 2009). Untuk menguji efektivitas awal, penelitian menggunakan desain one-group pretest-posttest. Subjek penelitian adalah siswa kelas X atau XI MAN Kotawaringin Timur yang aktif menggunakan media digital. Uji coba terbatas direncanakan melibatkan 30-40 siswa dengan purposive sampling.")
    add_table(
        doc,
        ["Konstruk/Aspek", "Definisi", "Indikator", "Instrumen"],
        [
            ["Kesadaran berbahasa", "Kepekaan terhadap bentuk, fungsi, konteks, dan dampak sosial bahasa.", "Mengenali ragam; menilai konteks; memberi alasan pilihan bahasa.", "Kuis PRIMA+ dan refleksi."],
            ["Loyalitas berbahasa Indonesia", "Kemauan menempatkan bahasa Indonesia secara tepat sebagai identitas dan alat komunikasi.", "Sikap positif; kesetiaan; norma; kebanggaan; pemilihan ragam.", "Kuesioner Likert 20 butir."],
            ["Kelayakan PRIMA+", "Kelayakan produk sebagai media pembelajaran.", "Materi; bahasa; tampilan; navigasi; umpan balik.", "Lembar validasi guru/ahli."],
            ["Respons pengguna", "Tanggapan siswa setelah menggunakan PRIMA+.", "Kemudahan; kemenarikan; manfaat; kendala.", "Angket respons/refleksi."],
        ],
        widths=[3.5, 4.7, 4.2, 3.4],
        size=8,
    )
    add_numbered(
        doc,
        [
            "Analyze: menganalisis kebutuhan siswa, contoh bahasa digital, dan indikator loyalitas berbahasa.",
            "Design: merancang fitur PRIMA+, kisi-kisi instrumen, skenario kuis, dan rubrik validasi.",
            "Develop: membangun prototipe, konten kuis, aset visual/audio, dan umpan balik.",
            "Implement: melakukan pretest, uji coba PRIMA+, posttest, dan pengumpulan respons siswa.",
            "Evaluate: menganalisis data, menafsirkan hasil, dan merevisi produk.",
        ],
    )
    if FLOWCHART.exists():
        doc.add_picture(str(FLOWCHART), width=Cm(12.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Gambar 1. Alur prosedur penelitian PRIMA+ berbasis ADDIE")
        set_font(r, size=10)

    add_heading(doc, "3.4 RANCANGAN PENGOLAHAN DATA", level=2)
    add_para(doc, "Data validasi guru/ahli dihitung dalam bentuk persentase kelayakan. Skor pretest dan posttest loyalitas berbahasa dibandingkan melalui rata-rata, persentase peningkatan, dan N-gain. Jika data memenuhi syarat, peningkatan skor diuji dengan paired sample t-test; jika tidak memenuhi syarat, digunakan uji Wilcoxon. Refleksi siswa dianalisis secara tematik untuk melihat pengalaman, kendala, dan saran setelah menggunakan PRIMA+.")
    add_table(
        doc,
        ["Data", "Teknik Pengolahan", "Interpretasi"],
        [
            ["Validasi produk", "Persentase kelayakan", "Kategori sangat layak, layak, cukup, atau perlu revisi."],
            ["Pretest-posttest", "Rata-rata, persentase peningkatan, N-gain", "Posttest lebih tinggi menunjukkan peningkatan awal."],
            ["Uji beda", "Paired t-test atau Wilcoxon", "Hipotesis diterima bila peningkatan bermakna."],
            ["Refleksi siswa", "Analisis tema ringkas", "Menjelaskan pengalaman, kendala, dan saran pengguna."],
        ],
        widths=[3.5, 5.2, 7.1],
        size=8,
    )

    add_heading(doc, "BAB 4. RANCANGAN ANGGARAN BIAYA DAN JADWAL KEGIATAN", level=1)
    add_heading(doc, "4.1 RANCANGAN ANGGARAN BIAYA PENELITIAN", level=2)
    add_table(
        doc,
        ["No.", "Komponen Pengeluaran", "Jumlah Anggaran (Rp)"],
        [
            ["1", "Biaya alat dan bahan: aset visual, aset audio, perangkat pendukung, cetak instrumen", "3.700.000"],
            ["2", "Biaya jasa: pengembangan prototipe, validasi guru/ahli, revisi materi", "3.300.000"],
            ["3", "Biaya sewa/akses: internet, hosting/domain uji coba, platform pendukung", "1.000.000"],
            ["4", "Akomodasi dan transportasi: koordinasi uji coba, operasional, dokumentasi", "2.000.000"],
            ["", "Total", "10.000.000"],
        ],
        widths=[1.2, 10.0, 4.5],
        size=9,
    )
    add_para(doc, "Total anggaran dirancang sebesar Rp10.000.000, masih berada di bawah batas maksimal Rp15.000.000. Biaya diarahkan pada pengembangan prototipe, penyediaan aset, penyusunan instrumen, pelaksanaan uji coba, validasi, dan dokumentasi.")

    add_heading(doc, "4.2 JADWAL PENELITIAN", level=2)
    add_table(
        doc,
        ["No.", "Kegiatan", "Tanggal Pelaksanaan"],
        [
            ["1", "Analisis kebutuhan dan studi literatur", "6-16 April 2026"],
            ["2", "Perancangan platform dan instrumen", "17-30 April 2026"],
            ["3", "Pengembangan prototipe PRIMA+", "1-25 Mei 2026"],
            ["4", "Validasi dan revisi awal", "26 Mei-5 Juni 2026"],
            ["5", "Pretest, uji coba, dan posttest", "6-30 Juni 2026"],
            ["6", "Pengolahan data dan revisi akhir", "1-25 Juli 2026"],
            ["7", "Penyusunan laporan akhir", "26 Juli-20 Agustus 2026"],
        ],
        widths=[1.2, 8.5, 5.8],
        size=9,
    )

    add_heading(doc, "PERNYATAAN PENGGUNAAN KECERDASAN ARTIFISIAL (AI)", level=1)
    add_para(doc, "Penelitian ini menggunakan bantuan Artificial Intelligence (AI), yaitu Gemini dan ChatGPT, secara terbatas untuk mencari referensi awal, memahami teori, merapikan kerangka metodologi, dan mencari alternatif ide alur permainan atau kuis bahasa. Keputusan penelitian, penyusunan instrumen, pengembangan produk, pengambilan data, analisis data, dan penarikan kesimpulan tetap dilakukan oleh peneliti. AI tidak digunakan untuk membuat data palsu, menggantikan proses uji coba, atau menulis hasil penelitian tanpa verifikasi.")

    add_heading(doc, "DAFTAR PUSTAKA", level=1)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(ref)
        set_font(r, size=11)

    add_heading(doc, "LAMPIRAN", level=1)
    add_plain(doc, "Lampiran 1. Kisi-kisi instrumen konstruk loyalitas berbahasa Indonesia")
    add_table(
        doc,
        ["Konstruk", "Indikator", "Jumlah Butir", "Contoh Pernyataan"],
        [
            ["Loyalitas berbahasa Indonesia", "Sikap positif", "4", "Bahasa Indonesia tetap penting digunakan di media digital."],
            ["Loyalitas berbahasa Indonesia", "Kesetiaan penggunaan", "4", "Saya memilih bahasa Indonesia yang jelas saat menulis informasi sekolah."],
            ["Loyalitas berbahasa Indonesia", "Kesadaran norma", "4", "Saya dapat membedakan bahasa santai dan bahasa formal."],
            ["Loyalitas berbahasa Indonesia", "Kebanggaan bahasa", "4", "Saya bangga menggunakan bahasa Indonesia dengan baik."],
            ["Loyalitas berbahasa Indonesia", "Pemilihan ragam bahasa", "4", "Saya menyesuaikan bahasa dengan lawan bicara dan tujuan komunikasi."],
        ],
        widths=[4.0, 3.5, 2.3, 6.0],
        size=8,
    )

    props = doc.core_properties
    props.title = "Proposal Penelitian PRIMA+ Format OPSI 2026"
    props.author = ""
    props.last_modified_by = ""
    props.comments = ""
    doc.save(OUT_DOCX)


def build_revision_table():
    doc = Document()
    style_doc(doc)
    add_center(doc, "TABEL REVISI SISTEMATIKA PROPOSAL PRIMA+", size=14, bold=True)
    add_para(doc, "Tabel ini mencatat perubahan setelah proposal disesuaikan kembali dengan sistematika resmi OPSI SMA/MA/Sederajat 2026.")
    add_table(
        doc,
        ["No.", "Bagian", "Masalah Sebelumnya", "Perbaikan", "Acuan/Letak"],
        [
            ["1", "Sistematika bab", "Subbab tambahan dibuat di luar sistematika OPSI.", "Struktur dikembalikan ke BAB 1-4 sesuai Lampiran 4 OPSI 2026.", "Seluruh naskah"],
            ["2", "Gap dan batasan", "Gap dan batasan berdiri sebagai subbab sendiri.", "Gap, kebaruan, dan target temuan digabung ke Latar Belakang.", "BAB 1.1"],
            ["3", "Hipotesis", "Belum ditampilkan padahal ada uji pretest-posttest.", "Ditambahkan H0 dan H1.", "BAB 1.4"],
            ["4", "Landasan teori", "Teori dipecah menjadi banyak subbab tidak resmi.", "Dikonsolidasikan ke Landasan Teori.", "BAB 2.1"],
            ["5", "Studi pustaka", "Posisi penelitian belum mengikuti nama subbab resmi.", "Dijadikan Studi Pustaka dengan tabel state of the art.", "BAB 2.2"],
            ["6", "Metode", "Desain, subjek, konstruk, dan instrumen berdiri sebagai subbab terpisah.", "Digabung dalam Rancangan dan Prosedur Penelitian.", "BAB 3.3"],
            ["7", "Pengolahan data", "Teknik analisis dan kriteria keberhasilan berdiri sendiri.", "Dimasukkan ke Rancangan Pengolahan Data.", "BAB 3.4"],
            ["8", "RAB", "Format RAB belum mengikuti kategori OPSI.", "RAB disusun ke biaya alat/bahan, jasa, sewa/akses, akomodasi/transportasi.", "BAB 4.1"],
            ["9", "AI", "Pernyataan AI dipertahankan tetapi disesuaikan istilah resmi.", "Judul menjadi Pernyataan Penggunaan Kecerdasan Artifisial (AI).", "Setelah BAB 4"],
        ],
        widths=[1.0, 3.2, 4.7, 4.7, 2.2],
        size=8,
    )
    doc.core_properties.title = "Tabel Revisi Sistematika Proposal PRIMA+"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.save(OUT_TABLE)


if __name__ == "__main__":
    build_proposal()
    build_revision_table()
    print(OUT_DOCX)
    print(OUT_TABLE)
