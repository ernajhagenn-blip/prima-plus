from docx import Document


BUDGET_ROWS = [
    ["1", "Pengembangan prototipe platform PRIMA+", "1 paket", "Rp2.500.000", "Rp2.500.000", "Desain alur, fitur kuis, skor, umpan balik, dan integrasi aset"],
    ["2", "Aset visual game dan antarmuka", "1 paket", "Rp1.200.000", "Rp1.200.000", "Karakter, ikon, map sederhana, item, dan elemen UI"],
    ["3", "Aset audio permainan", "1 paket", "Rp600.000", "Rp600.000", "Efek suara, musik latar, dan audio pendukung"],
    ["4", "Perangkat pendukung uji coba", "1 paket", "Rp1.200.000", "Rp1.200.000", "Headset, flashdisk/penyimpanan, kabel, dan kebutuhan teknis"],
    ["5", "Akses internet dan hosting/domain uji coba", "1 paket", "Rp1.000.000", "Rp1.000.000", "Internet pengembangan, unggah prototipe, dan akses uji coba"],
    ["6", "Cetak instrumen penelitian", "1 paket", "Rp700.000", "Rp700.000", "Kuesioner pretest-posttest, lembar validasi, dan formulir observasi"],
    ["7", "Dokumentasi dan publikasi hasil", "1 paket", "Rp800.000", "Rp800.000", "Dokumentasi kegiatan, poster, dan bahan presentasi"],
    ["8", "Konsumsi dan operasional uji coba", "1 paket", "Rp1.200.000", "Rp1.200.000", "Koordinasi responden, konsumsi terbatas, dan kebutuhan lapangan"],
    ["9", "Validasi ahli/guru dan revisi materi", "1 paket", "Rp800.000", "Rp800.000", "Masukan guru/ahli bahasa/media dan perbaikan konten"],
    ["", "Total", "", "", "Rp10.000.000", ""],
]


def set_cell(cell, text):
    cell.text = str(text)


def update_budget(path):
    doc = Document(path)
    table = None
    for candidate in doc.tables:
        header = [cell.text.strip() for cell in candidate.rows[0].cells]
        if header[:6] == ["No.", "Komponen", "Jumlah", "Harga Satuan", "Subtotal", "Keterangan"]:
            table = candidate
            break
    if table is None:
        raise RuntimeError(f"Budget table not found in {path}")

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for row_data in BUDGET_ROWS:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            set_cell(row.cells[idx], value)

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.startswith("Catatan: Godot Engine"):
            paragraph.text = (
                "Catatan: total anggaran disesuaikan menjadi Rp10.000.000 agar lebih realistis "
                "untuk pengembangan prototipe, penyediaan aset, uji coba responden, validasi, "
                "dokumentasi, dan operasional penelitian. Godot Engine tidak dimasukkan sebagai "
                "biaya lisensi karena bersifat gratis/open-source; biaya pengembangan diarahkan "
                "pada pembuatan prototipe, aset, dan kebutuhan uji coba."
            )
            break

    doc.save(path)


def update_review_table():
    path = "tabel_revisi_reviewer.docx"
    doc = Document(path)
    table = doc.tables[0]
    for row in table.rows:
        vals = [cell.text.strip() for cell in row.cells]
        if vals and vals[0] == "13":
            row.cells[2].text = (
                "RAB disesuaikan dari Rp1.192.000 menjadi Rp10.000.000 dengan rincian komponen "
                "yang lebih realistis: pengembangan prototipe, aset visual/audio, perangkat uji coba, "
                "internet/hosting, cetak instrumen, dokumentasi, operasional, dan validasi ahli/guru."
            )
            row.cells[3].text = "BAB 4, bagian 4.1 Rancangan Anggaran Biaya."
            row.cells[4].text = (
                "Perbaikan membuat anggaran lebih sesuai dengan kebutuhan pengembangan platform digital "
                "dan uji efektivitas. Total dipilih Rp10.000.000 agar tetap wajar di bawah batas maksimal "
                "Rp15.000.000, tetapi cukup untuk mendukung prototipe, instrumen, uji coba, dan dokumentasi."
            )
            break
    doc.save(path)


if __name__ == "__main__":
    for filename in ["prima_revisi.docx", "prima_revisi_DAFTAR_PUSTAKA_UPDATE.docx"]:
        update_budget(filename)
    update_review_table()
    print("updated budget to Rp10.000.000")
