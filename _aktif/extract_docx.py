from docx import Document
from pathlib import Path
import json


doc = Document("prima.docx")

paragraphs = []
for idx, paragraph in enumerate(doc.paragraphs, 1):
    text = paragraph.text.strip()
    if text:
        paragraphs.append(
            {
                "idx": idx,
                "style": paragraph.style.name if paragraph.style else "",
                "text": text,
            }
        )

tables = []
for table_idx, table in enumerate(doc.tables, 1):
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    tables.append({"idx": table_idx, "rows": rows})

lines = []
for paragraph in paragraphs:
    lines.append(
        f"[{paragraph['idx']}] ({paragraph['style']}) {paragraph['text']}"
    )

for table in tables:
    lines.append(f"\n[TABLE {table['idx']}]")
    for row in table["rows"]:
        lines.append(" | ".join(cell.replace("\n", " / ") for cell in row))

Path("prima_extracted.txt").write_text("\n".join(lines), encoding="utf-8")
Path("prima_structure.json").write_text(
    json.dumps({"paragraphs": paragraphs, "tables": tables}, ensure_ascii=False, indent=2),
    encoding="utf-8",
)

print(f"paragraphs={len(paragraphs)} tables={len(tables)}")
print("wrote prima_extracted.txt and prima_structure.json")
