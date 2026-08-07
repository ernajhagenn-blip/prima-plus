from copy import deepcopy

from docx import Document


PROPOSAL = "Proposal_Penelitian_PRIMA_OPSI_2026_Format_Resmi.docx"
REVISION_TABLE = "Tabel_Revisi_Sistematika_OPSI_PRIMA.docx"
PROPOSAL_OUT = "Proposal_Penelitian_PRIMA_OPSI_2026_Format_Resmi_RAB_Rinci.docx"
REVISION_TABLE_OUT = "Tabel_Revisi_Sistematika_OPSI_PRIMA_RAB_Rinci.docx"


RAB_ROWS = [
    ("1", "Alat/bahan: cetak instrumen pretest-posttest, lembar validasi, dan angket respons", "600.000"),
    ("2", "Alat/bahan: kertas, map, label, tinta, dan perlengkapan administrasi uji coba", "650.000"),
    ("3", "Alat/bahan: aset visual antarmuka, ikon, dan ilustrasi ringan untuk PRIMA+", "950.000"),
    ("4", "Alat/bahan: aset audio pendek dan elemen pendukung kuis bahasa", "700.000"),
    ("5", "Alat/bahan: dokumentasi kegiatan, pencetakan laporan, dan penjilidan", "750.000"),
    ("6", "Jasa: perancangan prototipe antarmuka PRIMA+ tahap awal", "950.000"),
    ("7", "Jasa: pengembangan modul kuis, refleksi bahasa, dan umpan balik skor", "950.000"),
    ("8", "Jasa: penataan materi, penyuntingan bahasa, dan revisi konten setelah validasi", "750.000"),
    ("9", "Jasa: validasi guru/ahli untuk materi, media, bahasa, dan instrumen", "700.000"),
    ("10", "Sewa/akses: paket internet selama pengembangan, validasi, dan uji coba", "600.000"),
    ("11", "Sewa/akses: domain, hosting, atau ruang uji coba platform selama penelitian", "600.000"),
    ("12", "Sewa/akses: penyimpanan awan dan platform pendukung kolaborasi data", "400.000"),
    ("13", "Akomodasi/transportasi: transportasi lokal koordinasi, validasi, dan uji coba terbatas", "750.000"),
    ("14", "Akomodasi/transportasi: konsumsi ringan saat validasi dan uji coba siswa", "650.000"),
    ("", "Total", "10.000.000"),
]


def clear_table_rows(table):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold


def append_row_like(table, values, bold=False):
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value, bold=bold)
    return row


def revise_proposal():
    doc = Document(PROPOSAL)
    rab_table = None
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers == ["No.", "Komponen Pengeluaran", "Jumlah Anggaran (Rp)"]:
            rab_table = table
            break

    if rab_table is None:
        raise RuntimeError("RAB table not found.")

    clear_table_rows(rab_table)
    for no, component, amount in RAB_ROWS:
        append_row_like(rab_table, [no, component, amount], bold=(component == "Total"))

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Total anggaran dirancang sebesar Rp10.000.000"):
            paragraph.text = (
                "Total anggaran dirancang sebesar Rp10.000.000, masih berada di bawah "
                "batas maksimal Rp15.000.000. RAB disusun secara rinci agar tidak ada "
                "komponen pengeluaran tunggal di atas Rp1.000.000. Setiap item "
                "dirumuskan sebagai kebutuhan operasional yang dapat ditelusuri, meliputi "
                "alat/bahan, jasa, sewa/akses, serta akomodasi/transportasi."
            )
            break

    try:
        doc.save(PROPOSAL)
    except PermissionError:
        doc.save(PROPOSAL_OUT)


def revise_revision_table():
    doc = Document(REVISION_TABLE)
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if cells and cells[0].text.strip() == "8" and cells[1].text.strip() == "RAB":
                cells[2].text = (
                    "Format RAB sebelumnya masih memuat komponen besar/glondongan "
                    "dengan nominal di atas Rp1.000.000."
                )
                cells[3].text = (
                    "RAB dipecah menjadi item operasional rinci. Setiap baris pengeluaran "
                    "dibuat maksimal Rp1.000.000, tetap mengikuti kategori OPSI: "
                    "alat/bahan, jasa, sewa/akses, dan akomodasi/transportasi. Total "
                    "anggaran tetap Rp10.000.000."
                )
                cells[4].text = "BAB 4.1"
                break
    try:
        doc.save(REVISION_TABLE)
    except PermissionError:
        doc.save(REVISION_TABLE_OUT)


if __name__ == "__main__":
    revise_proposal()
    revise_revision_table()
    print("Revised RAB into detailed non-glondongan line items.")
