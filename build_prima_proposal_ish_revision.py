from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


PROPOSAL_OUT = "Proposal_Penelitian_PRIMA_OPSI_2026_ISH_Revisi_RAB_Konkret_v2.docx"
REVISION_OUT = "Tabel_Revisi_Reviewer_PRIMA_OPSI_2026_Revisi_RAB_Konkret.docx"
FIGURE = "prima_prosedur_ADDIE_HIGH.png"

TITLE = "KESADARAN BERBAHASA REMAJA MELALUI MEDIA PRIMA+ UNTUK MENGUATKAN LOYALITAS BAHASA INDONESIA DI LINGKUNGAN SEKOLAH"


def set_run(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_p(paragraph, align=None, before=0, after=6, line=1.15):
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_p(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, bold=False, italic=False, after=6, before=0):
    p = doc.add_paragraph()
    set_p(p, align=align, after=after, before=before)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, italic=italic)
    return p


def add_center(doc, text, size=12, bold=True, after=8):
    return add_p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=size, bold=bold, after=after)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    set_p(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=10 if level == 1 else 6, after=5)
    r = p.add_run(text)
    set_run(r, size=14 if level == 1 else 12, bold=True)
    return p


def add_list(doc, text):
    return add_p(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3)


def shade_cell(cell, fill="EDEDED"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    set_p(p, align=align, after=0, line=1.0)
    r = p.add_run(str(text))
    set_run(r, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, caption=None, widths=None):
    if caption:
        p = add_p(doc, caption, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, after=3)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    return doc


def build_proposal():
    doc = setup_doc()
    for _ in range(2):
        doc.add_paragraph()
    add_center(doc, "PROPOSAL PENELITIAN OPSI", size=14, bold=True, after=18)
    add_center(doc, TITLE, size=14, bold=True, after=22)
    add_center(doc, "Nama Tim Peneliti", size=12, bold=True, after=4)
    add_center(doc, "Ketua : Larisa Dianti", size=12, bold=True, after=2)
    add_center(doc, "Anggota : Gusti Nazwa Azizah", size=12, bold=True, after=14)
    add_center(doc, "Bidang Kompetisi Penelitian: Ilmu Sosial dan Humaniora (ISH)", size=12, bold=True, after=2)
    add_center(doc, "Subbidang: Bahasa dan Sastra", size=12, bold=True, after=14)
    add_center(doc, "MAN KOTAWARINGIN TIMUR", size=12, bold=True, after=2)
    add_center(doc, "Kotawaringin Timur, Kalimantan Tengah", size=12, bold=True, after=14)
    add_center(doc, "Tahun 2026", size=12, bold=True)
    doc.add_page_break()

    add_heading(doc, "DAFTAR ISI", 1)
    for line in [
        "BAB 1. PENDAHULUAN",
        "1.1 Latar Belakang",
        "1.2 Rumusan Masalah",
        "1.3 Tujuan Penelitian",
        "1.4 Manfaat Penelitian",
        "BAB 2. TINJAUAN PUSTAKA",
        "2.1 Landasan Teori",
        "2.2 Studi Pustaka",
        "BAB 3. METODE PENELITIAN",
        "3.1 Waktu dan Tempat Penelitian",
        "3.2 Alat dan Bahan",
        "3.3 Rancangan dan Prosedur Penelitian",
        "3.4 Rancangan Pengolahan Data",
        "BAB 4. RANCANGAN ANGGARAN BIAYA DAN JADWAL KEGIATAN",
        "PERNYATAAN PENGGUNAAN KECERDASAN ARTIFISIAL (AI)",
        "DAFTAR PUSTAKA",
        "LAMPIRAN",
        "Lampiran 1. Kisi-kisi instrumen loyalitas berbahasa Indonesia",
        "Lampiran 2. Butir instrumen kuesioner loyalitas berbahasa Indonesia",
        "Lampiran 3. Kisi-kisi skenario kasus bahasa PRIMA+",
        "Lampiran 4. Lembar validasi guru/ahli",
        "Lampiran 5. Angket respons/refleksi siswa",
    ]:
        add_p(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)

    add_heading(doc, "DAFTAR TABEL", 1)
    for line in [
        "Tabel 1. Studi pustaka terdahulu",
        "Tabel 2. Alat dan bahan penelitian",
        "Tabel 3. Konstruk, dimensi, instrumen, dan skala pengukuran",
        "Tabel 4. Rancangan pengolahan data",
        "Tabel 5. Rancangan anggaran biaya penelitian",
        "Tabel 6. Jadwal penelitian",
    ]:
        add_p(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)
    add_heading(doc, "DAFTAR GAMBAR", 1)
    add_p(doc, "Gambar 1. Alur prosedur penelitian PRIMA+ berbasis ADDIE", align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    add_heading(doc, "BAB 1. PENDAHULUAN", 1)
    add_heading(doc, "1.1 LATAR BELAKANG", 2)
    for text in [
        "Bahasa Indonesia digunakan siswa dalam kegiatan belajar, diskusi kelas, organisasi, penulisan tugas, dan komunikasi digital. Di lingkungan sekolah, kemampuan memilih ragam bahasa menjadi bagian penting dari kedisiplinan akademik dan sikap berbahasa. Tantangannya, remaja kini bergerak di dua ruang sekaligus: ruang sekolah yang menuntut bahasa Indonesia yang jelas dan tertib, serta ruang digital yang cenderung cepat, ringkas, santai, dan campur kode.",
        "Konteks penelitian ini ditempatkan di MAN Kotawaringin Timur, Sampit. Data Referensi Kemendikdasmen mencatat MAN Kotawaringin Timur sebagai madrasah negeri di Jalan H.M. Arsyad No. 68 Sampit, Kecamatan Mentawa Baru Ketapang, Kabupaten Kotawaringin Timur, dengan total peserta didik 402 orang pada pembaruan data 26 Juni 2026. Jumlah peserta didik tersebut menunjukkan bahwa sekolah memiliki lingkungan sosial yang cukup besar untuk mengamati praktik bahasa remaja dalam komunikasi akademik dan digital.",
        "Masalah yang tampak bukan sekadar penggunaan slang atau istilah asing, melainkan melemahnya kepekaan siswa dalam membedakan kapan bahasa santai dapat digunakan dan kapan bahasa Indonesia yang lebih tertib diperlukan. Dalam konteks pembelajaran Bahasa Indonesia, kondisi ini berkaitan dengan kesadaran berbahasa, sikap terhadap bahasa Indonesia, dan loyalitas berbahasa Indonesia sebagai bagian dari identitas dan komunikasi sekolah.",
        "Kajian tentang remaja Sampit juga menunjukkan bahwa media dapat berperan dalam membentuk cara remaja memaknai identitas dan kehidupan sosialnya. Widyaningrum, Yumarnamto, dan Prijambodo (2020) meneliti remaja Kota Sampit sebagai khalayak media dan menunjukkan pentingnya memahami remaja melalui konteks lokal, identitas, dan pengalaman sosialnya. Oleh karena itu, penelitian bahasa remaja di Sampit perlu ditempatkan sebagai kajian sosial-humaniora, bukan semata-mata sebagai pengembangan teknologi.",
        "PRIMA+ dalam penelitian ini diposisikan sebagai media pembelajaran/kebahasaan, bukan sebagai produk rekayasa perangkat lunak. Unsur digital dipakai untuk menghadirkan kasus bahasa, pilihan ragam, kuis kontekstual, refleksi, dan umpan balik. Fokus penelitian tetap berada pada perubahan kesadaran berbahasa dan loyalitas bahasa Indonesia remaja di lingkungan sekolah.",
        "Kebaruan penelitian ini terletak pada penggunaan media PRIMA+ sebagai sarana latihan pengambilan keputusan berbahasa. Penelitian tidak hanya mendeskripsikan fenomena bahasa digital remaja, tetapi mengembangkan media yang dapat membantu siswa menilai konteks, memilih ragam bahasa, dan merefleksikan alasan penggunaan bahasa Indonesia secara tepat.",
    ]:
        add_p(doc, text)

    add_heading(doc, "1.2 RUMUSAN MASALAH", 2)
    for text in [
        "1. Bagaimana rancangan media PRIMA+ berbasis kesadaran berbahasa untuk menguatkan loyalitas bahasa Indonesia remaja di lingkungan sekolah?",
        "2. Bagaimana kelayakan materi, bahasa, media, dan instrumen PRIMA+ berdasarkan validasi guru atau ahli?",
        "3. Bagaimana efektivitas penggunaan PRIMA+ terhadap peningkatan loyalitas berbahasa Indonesia remaja?",
    ]:
        add_list(doc, text)
    add_heading(doc, "1.3 TUJUAN PENELITIAN", 2)
    for text in [
        "1. Mengembangkan media PRIMA+ berbasis kesadaran berbahasa sesuai kebutuhan remaja di lingkungan sekolah.",
        "2. Menilai kelayakan PRIMA+ dari aspek materi, bahasa, media, dan instrumen penelitian.",
        "3. Menguji efektivitas PRIMA+ dalam menguatkan loyalitas berbahasa Indonesia remaja melalui desain pretest-posttest.",
    ]:
        add_list(doc, text)
    add_heading(doc, "1.4 MANFAAT PENELITIAN", 2)
    for text in [
        "Bagi siswa, PRIMA+ membantu mereka lebih sadar dalam memilih ragam bahasa sesuai konteks sekolah dan media digital.",
        "Bagi guru Bahasa Indonesia, PRIMA+ dapat menjadi contoh media pembelajaran yang menghubungkan materi bahasa dengan situasi komunikasi nyata siswa.",
        "Bagi sekolah, penelitian ini mendukung pembiasaan berbahasa Indonesia yang tertib, santun, dan sesuai konteks akademik.",
        "Bagi kajian bahasa dan sastra, penelitian ini memberi contoh pengembangan media berbasis kesadaran berbahasa untuk menguatkan sikap dan loyalitas bahasa Indonesia remaja.",
    ]:
        add_list(doc, text)

    add_heading(doc, "BAB 2. TINJAUAN PUSTAKA", 1)
    add_heading(doc, "2.1 LANDASAN TEORI", 2)
    for text in [
        "Kesadaran berbahasa menekankan kemampuan memperhatikan bentuk, fungsi, konteks, dan dampak sosial dari bahasa yang digunakan. Carter (2003) memandang language awareness sebagai kesadaran terhadap bentuk dan fungsi bahasa, sedangkan Fairclough (1992) menekankan pentingnya sikap kritis dalam melihat hubungan bahasa dan kehidupan sosial.",
        "Loyalitas bahasa merupakan bagian dari sikap bahasa. Garvin dan Mathiot (1968) menempatkan loyalitas bahasa bersama kebanggaan bahasa dan kesadaran norma. Dalam penelitian ini, loyalitas berbahasa Indonesia dimaknai sebagai kemauan siswa untuk menempatkan bahasa Indonesia secara tepat, terutama pada situasi sekolah, komunikasi resmi, dan ruang digital yang berkaitan dengan kegiatan akademik.",
        "Bahasa remaja di media sosial memperlihatkan perubahan pilihan kata, campur kode, dan ragam santai. Penelitian tentang pengaruh media sosial terhadap bahasa remaja menunjukkan bahwa media digital memengaruhi gaya penulisan bahasa Indonesia (Prasetyaningrum, 2024) dan penggunaan ragam informal di kalangan remaja (Parlindungan Siahaan dkk., 2024). Temuan tersebut menjadi dasar bahwa pembinaan bahasa perlu dekat dengan kebiasaan komunikasi remaja.",
        "Media digital dalam penelitian ini dipahami sebagai sarana pedagogis untuk menghadirkan kasus bahasa dan refleksi, bukan sebagai objek utama rekayasa komputer. Dengan demikian, pengembangan PRIMA+ tetap ditempatkan dalam kajian sosial-humaniora, terutama pembelajaran bahasa, sikap bahasa, dan perilaku komunikasi remaja.",
    ]:
        add_p(doc, text)
    add_heading(doc, "2.2 STUDI PUSTAKA", 2)
    add_table(
        doc,
        ["No.", "Sumber", "Temuan Utama", "Posisi terhadap PRIMA+"],
        [
            ["1", "Carter (2003); Fairclough (1992)", "Kesadaran berbahasa berkaitan dengan bentuk, fungsi, konteks, dan sikap kritis.", "Menjadi dasar latihan refleksi pilihan bahasa."],
            ["2", "Garvin dan Mathiot (1968); Chaer dan Agustina (2014)", "Loyalitas bahasa merupakan bagian dari sikap bahasa dan berkaitan dengan norma serta kebanggaan bahasa.", "Menjadi dasar konstruk loyalitas berbahasa Indonesia."],
            ["3", "Widyaningrum dkk. (2020)", "Remaja Kota Sampit dapat dikaji melalui konteks media, identitas, dan pengalaman sosial lokal.", "Menegaskan konteks lokal Sampit dan kategori sosial-humaniora."],
            ["4", "Prasetyaningrum (2024); Parlindungan Siahaan dkk. (2024)", "Media sosial memengaruhi gaya bahasa dan ragam informal remaja.", "Memperkuat masalah bahasa digital remaja."],
            ["5", "Shortt dkk. (2021); Luo (2023)", "Gamifikasi dapat mendukung pembelajaran bahasa jika desain dan evaluasinya jelas.", "Menjadi pendukung bentuk media interaktif PRIMA+."],
        ],
        caption="Tabel 1. Studi pustaka terdahulu",
        widths=[0.45, 1.55, 2.35, 2.15],
    )

    add_heading(doc, "BAB 3. METODE PENELITIAN", 1)
    add_heading(doc, "3.1 WAKTU DAN TEMPAT PENELITIAN", 2)
    add_p(doc, "Penelitian direncanakan berlangsung pada April-Agustus 2026 di MAN Kotawaringin Timur, Sampit, Kabupaten Kotawaringin Timur, Kalimantan Tengah. Kegiatan penelitian meliputi analisis kebutuhan, perancangan media PRIMA+, penyusunan instrumen, validasi, uji coba terbatas, pengolahan data, revisi media, dan penyusunan laporan.")
    add_heading(doc, "3.2 ALAT DAN BAHAN", 2)
    add_table(
        doc,
        ["No.", "Alat/Bahan", "Jumlah", "Fungsi"],
        [
            ["1", "Laptop peneliti", "1 unit", "Menyusun materi, mengolah data, dan mengembangkan PRIMA+."],
            ["2", "Perangkat siswa", "30 perangkat/sesuai ketersediaan", "Mengakses PRIMA+ (web/PWA) saat uji coba terbatas."],
            ["3", "ReactJS (frontend)", "1 framework", "Membangun antarmuka web interaktif dan PWA untuk PRIMA+."],
            ["4", "Express.js + Node.js (backend)", "1 sistem", "Mengelola API, autentikasi, dan logika aplikasi PRIMA+."],
            ["5", "Firebase (database & autentikasi)", "1 layanan", "Menyimpan data pengguna, skor pretest-posttest, dan progress."],
            ["6", "Gemini API (LLM)", "1 layanan", "Menghasilkan soal kuis dinamis, feedback otomatis, dan refleksi bahasa."],
            ["7", "Google AI Studio", "1 platform", "Mengembangkan dan menguji prompt Gemini API untuk konten PRIMA+."],
            ["8", "Skenario kasus bahasa remaja", "1 paket", "Materi latihan pilihan ragam bahasa dalam PRIMA+."],
            ["9", "Kuesioner loyalitas berbahasa", "20 butir", "Mengukur sikap positif, kesetiaan penggunaan, kesadaran norma, kebanggaan, dan pemilihan ragam."],
            ["10", "Lembar validasi guru/ahli", "2-3 validator", "Menilai kelayakan materi, bahasa, media, dan instrumen."],
        ],
        caption="Tabel 2. Alat dan bahan penelitian",
        widths=[0.45, 1.7, 1.25, 3.1],
    )
    add_heading(doc, "3.3 RANCANGAN DAN PROSEDUR PENELITIAN", 2)
    add_p(doc, "Penelitian menggunakan pendekatan Research and Development (R&D) sederhana dengan model ADDIE. Model ini dipilih karena penelitian bertujuan mengembangkan media pembelajaran/kebahasaan dan menguji penggunaannya secara terbatas. Desain uji efektivitas menggunakan one-group pretest-posttest, yaitu siswa mengisi kuesioner loyalitas berbahasa sebelum dan sesudah menggunakan PRIMA+.")
    add_p(doc, "Subjek uji coba terbatas berjumlah 30 siswa kelas X atau XI MAN Kotawaringin Timur. Jumlah 30 siswa dipilih karena penelitian ini bukan survei populasi besar, melainkan uji coba terbatas media dalam penelitian pengembangan. Dengan 30 siswa, peneliti memperoleh 30 pasangan data pretest-posttest sehingga analisis peningkatan dapat dilakukan melalui rata-rata, persentase peningkatan, N-gain, serta uji beda berpasangan. Jumlah tersebut juga memenuhi ukuran minimum yang lazim digunakan untuk analisis awal berbasis distribusi sampel pada desain berpasangan, sekaligus masih realistis dilaksanakan di lingkungan sekolah.")
    add_p(doc, "Teknik pengambilan sampel menggunakan purposive sampling. Kriteria responden adalah siswa kelas X atau XI, aktif mengikuti pembelajaran Bahasa Indonesia, menggunakan media digital dalam komunikasi sehari-hari, bersedia mengikuti pretest, uji coba PRIMA+, dan posttest, serta memperoleh izin dari sekolah sesuai ketentuan penelitian yang melibatkan manusia.")
    for text in [
        "Analyze: menganalisis kebutuhan siswa, contoh bahasa remaja di lingkungan sekolah dan media digital, serta indikator loyalitas berbahasa.",
        "Design: merancang skenario kasus bahasa, kisi-kisi instrumen, alur penggunaan PRIMA+, dan rubrik validasi.",
        "Develop: menyusun media PRIMA+ sebagai aplikasi web berbasis ReactJS dengan backend Express.js, database Firebase, serta integrasi Gemini API untuk menghasilkan soal kuis dinamis, feedback otomatis, dan refleksi pilihan bahasa.",
        "Implement: melaksanakan pretest, uji coba PRIMA+ kepada 30 siswa, posttest, dan pengumpulan respons siswa.",
        "Evaluate: menganalisis data, menafsirkan hasil, dan merevisi media berdasarkan validasi serta respons siswa.",
    ]:
        add_list(doc, text)
    if __import__("os").path.exists(FIGURE):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(FIGURE, width=Inches(5.6))
            add_center(doc, "Gambar 1. Alur prosedur penelitian PRIMA+ berbasis ADDIE", size=10, bold=False)
        except Exception:
            add_center(doc, "Gambar 1. Alur prosedur penelitian PRIMA+ berbasis ADDIE", size=10, bold=False)
    add_table(
        doc,
        ["Konstruk", "Dimensi/Indikator", "Jumlah Butir", "Instrumen", "Skala"],
        [
            ["Kesadaran berbahasa", "Mengenali ragam bahasa; menilai kesesuaian konteks; memberi alasan pilihan bahasa.", "5 skenario kuis", "Kuis PRIMA+ (skenario kasus bahasa)", "Skor 0-100"],
            ["Loyalitas berbahasa Indonesia", "Sikap positif terhadap bahasa Indonesia", "4", "Kuesioner Likert", "1-4"],
            ["", "Kesetiaan menggunakan bahasa Indonesia", "4", "Kuesioner Likert", "1-4"],
            ["", "Kesadaran norma bahasa", "4", "Kuesioner Likert", "1-4"],
            ["", "Kebanggaan terhadap bahasa Indonesia", "4", "Kuesioner Likert", "1-4"],
            ["", "Pemilihan ragam sesuai konteks", "4", "Kuesioner Likert", "1-4"],
            ["Kelayakan media", "Materi; bahasa; tampilan; navigasi; umpan balik.", "15 aspek", "Lembar validasi guru/ahli", "1-4"],
            ["Respons pengguna", "Kemudahan; kemenarikan; manfaat; kendala.", "8 butir", "Angket respons/refleksi", "1-4 + terbuka"],
        ],
        caption="Tabel 3. Konstruk, dimensi, instrumen, dan skala pengukuran",
        widths=[1.2, 2.0, 0.7, 1.6, 0.6],
    )
    add_heading(doc, "3.4 RANCANGAN PENGOLAHAN DATA", 2)
    add_p(doc, "Rancangan pengolahan data disesuaikan dengan tiga rumusan masalah dan tujuan penelitian. Rumusan masalah pertama dijawab melalui deskripsi proses pengembangan media PRIMA+ berdasarkan tahapan ADDIE. Rumusan masalah kedua dijawab melalui persentase kelayakan dari validasi guru atau ahli. Rumusan masalah ketiga dijawab melalui perbandingan skor pretest dan posttest loyalitas berbahasa Indonesia.")
    add_table(
        doc,
        ["Data", "Teknik Pengolahan", "Interpretasi"],
        [
            ["Validasi media", "Persentase kelayakan", "Kategori sangat layak, layak, cukup, atau perlu revisi."],
            ["Pretest-posttest", "Rata-rata, persentase peningkatan, dan N-gain", "Posttest lebih tinggi menunjukkan peningkatan awal loyalitas berbahasa."],
            ["Uji beda berpasangan", "Paired sample t-test jika selisih data berdistribusi normal; Wilcoxon signed-rank test jika tidak normal", "Peningkatan dinilai bermakna bila hasil uji menunjukkan perbedaan yang signifikan."],
            ["Refleksi siswa", "Analisis tema ringkas", "Menjelaskan pengalaman, kendala, dan saran pengguna."],
        ],
        caption="Tabel 4. Rancangan pengolahan data",
        widths=[1.45, 2.6, 2.45],
    )

    add_heading(doc, "BAB 4. RANCANGAN ANGGARAN BIAYA DAN JADWAL KEGIATAN", 1)
    add_heading(doc, "4.1 RANCANGAN ANGGARAN BIAYA PENELITIAN", 2)
    add_p(doc, "Total anggaran dirancang sebesar Rp1.987.100. Walaupun pedoman OPSI memperbolehkan anggaran sampai Rp15.000.000, penelitian ini dibatasi maksimal Rp2.000.000 agar lebih proporsional dengan ruang lingkup uji coba terbatas kepada 30 siswa. RAB ditulis sebagai item konkret yang akan dibeli, dicetak, disewa, atau digunakan dalam kegiatan penelitian. Layanan Firebase (Spark Plan) dan Gemini API (Free Tier) digunakan dalam batas gratis sehingga tidak memerlukan biaya tambahan.")
    add_table(
        doc,
        ["No.", "Komponen Pengeluaran", "Jumlah Anggaran (Rp)"],
        [
            ["1", "Barang: kertas HVS A4 80 gsm (2 rim)", "112.000"],
            ["2", "Barang: tinta printer hitam refill (1 botol)", "62.500"],
            ["3", "Barang: tinta printer warna refill (1 botol)", "78.500"],
            ["4", "Barang: map plastik/snelhecter untuk berkas responden (35 lembar)", "52.500"],
            ["5", "Barang: label kode responden (2 pak)", "24.000"],
            ["6", "Barang: pulpen untuk pengisian instrumen (2 lusin)", "48.000"],
            ["7", "Barang: binder clip, staples, dan isolasi arsip (1 paket)", "35.750"],
            ["8", "Barang: flashdisk 32 GB untuk cadangan data penelitian (1 unit)", "44.850"],
            ["9", "Barang: materai Rp10.000 untuk berkas izin/validasi (4 lembar)", "40.000"],
            ["10", "Barang: amplop cokelat arsip dokumen (10 lembar)", "18.500"],
            ["11", "Jasa: fotokopi/cetak kuesioner pretest-posttest (60 set)", "210.000"],
            ["12", "Jasa: cetak lembar validasi dan angket respons (40 set)", "137.500"],
            ["13", "Jasa: cetak kartu skenario kasus bahasa PRIMA+ (30 set)", "185.000"],
            ["14", "Jasa: jilid laporan proposal/lampiran (3 eksemplar)", "90.000"],
            ["15", "Sewa/akses: paket data internet 25 GB untuk koordinasi dan uji coba (1 nomor)", "135.000"],
            ["16", "Sewa: modem Wi-Fi portabel saat uji coba (1 hari)", "85.000"],
            ["17", "Sewa: tripod HP untuk dokumentasi kegiatan (1 hari)", "45.000"],
            ["18", "Lainnya: transport lokal ke sekolah, validator, dan percetakan (4 kali perjalanan)", "245.000"],
            ["19", "Lainnya: konsumsi ringan validator dan siswa uji coba (35 paket)", "280.000"],
            ["20", "Lainnya: air mineral gelas untuk kegiatan validasi/uji coba (2 dus)", "58.000"],
            ["", "Total", "1.987.100"],
        ],
        caption="Tabel 5. Rancangan anggaran biaya penelitian",
        widths=[0.45, 4.75, 1.3],
    )
    add_heading(doc, "4.2 JADWAL PENELITIAN", 2)
    add_table(
        doc,
        ["No.", "Kegiatan", "Tanggal Pelaksanaan"],
        [
            ["1", "Analisis kebutuhan dan studi literatur", "6-16 April 2026"],
            ["2", "Perancangan media dan instrumen", "17-30 April 2026"],
            ["3", "Penyusunan media PRIMA+", "1-25 Mei 2026"],
            ["4", "Validasi dan revisi awal", "26 Mei-5 Juni 2026"],
            ["5", "Pretest, uji coba kepada 30 siswa, dan posttest", "6-30 Juni 2026"],
            ["6", "Pengolahan data dan revisi akhir", "1-25 Juli 2026"],
            ["7", "Penyusunan laporan akhir", "26 Juli-20 Agustus 2026"],
        ],
        caption="Tabel 6. Jadwal penelitian",
        widths=[0.45, 4.0, 2.05],
    )
    add_heading(doc, "PERNYATAAN PENGGUNAAN KECERDASAN ARTIFISIAL (AI)", 1)
    add_p(doc, "Penelitian ini menggunakan bantuan Artificial Intelligence (AI), yaitu Gemini dan ChatGPT, secara terbatas untuk mencari referensi awal, memahami teori, merapikan kerangka metodologi, dan mencari alternatif ide alur kuis/refleksi bahasa. Keputusan penelitian, penyusunan instrumen, pengembangan media, pengambilan data, analisis data, dan penarikan kesimpulan tetap dilakukan oleh peneliti. AI tidak digunakan untuk membuat data palsu, menggantikan proses uji coba, atau menulis hasil penelitian tanpa verifikasi.")
    add_heading(doc, "DAFTAR PUSTAKA", 1)
    refs = [
        "Balai Bahasa Provinsi Kalimantan Tengah. (2026). Layanan Uji Kemahiran Berbahasa Indonesia (UKBI). https://balaibahasakalteng.kemendikdasmen.go.id/layanan/ukbi/",
        "Branch, R. M. (2009). Instructional Design: The ADDIE Approach. Springer. https://doi.org/10.1007/978-0-387-09506-6",
        "Carter, R. (2003). Language awareness. ELT Journal, 57(1), 64-65. https://doi.org/10.1093/elt/57.1.64",
        "Chaer, A., & Agustina, L. (2014). Sosiolinguistik: Perkenalan awal. Rineka Cipta.",
        "Crystal, D. (2011). Internet Linguistics: A Student Guide. Routledge.",
        "Fairclough, N. (1992). Critical Language Awareness. Longman.",
        "Garvin, P. L., & Mathiot, M. (1968). The urbanization of the Guarani language: A problem in language and culture. In J. A. Fishman (Ed.), Readings in the Sociology of Language (pp. 365-374). De Gruyter Mouton. https://doi.org/10.1515/9783110805376.365",
        "Kementerian Pendidikan Dasar dan Menengah. (2026). Informasi Satuan Pendidikan: MAN Kotawaringin Timur, NPSN 30201526. https://referensi.data.kemendikdasmen.go.id/snpmb/site/sekolah?npsn=30201526",
        "Luo, Z. (2023). The effectiveness of gamified tools for foreign language learning (FLL): A systematic review. Behavioral Sciences, 13(4), 331. https://doi.org/10.3390/bs13040331",
        "Parlindungan Siahaan, A., Aldy Pradana, M., Citra Chairani, D., Heriyani Erizal, A., & Margareta Lase, Y. (2024). Pengaruh era digital terhadap pemakaian bahasa Indonesia di kalangan remaja melalui media sosial. PENG: Jurnal Ekonomi dan Manajemen, 2(1), 879-885. https://teewanjournal.com/index.php/peng/article/view/1026",
        "Prasetyaningrum, R. (2024). Pengaruh media sosial terhadap gaya bahasa dalam penulisan bahasa Indonesia pada remaja. Jurnal Sosial Humaniora dan Pendidikan, 3(1), 127-134. https://doi.org/10.55606/inovasi.v3i1.2734",
        "Shortt, M., Tilak, S., Kuznetcova, I., Martens, B., & Akinkuolie, B. (2021). Gamification in mobile-assisted language learning: A systematic review of Duolingo literature from public release of 2012 to early 2020. Computer Assisted Language Learning, 36(3), 517-554. https://doi.org/10.1080/09588221.2021.1933540",
        "Sugiyono. (2019). Metode Penelitian dan Pengembangan: Research and Development. Alfabeta.",
        "UKBI. (2026). Uji Kemahiran Berbahasa Indonesia. Kementerian Pendidikan Dasar dan Menengah. https://ukbi.kemendikdasmen.go.id/",
        "Widyaningrum, A. Y., Yumarnamto, M., & Prijambodo, V. L. (2020). Analisis resepsi remaja Kota Sampit mengenai keberagaman di media. WACANA: Jurnal Ilmiah Ilmu Komunikasi, 19(1), 51-61.",
    ]
    for ref in refs:
        add_p(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT, after=3)
    add_heading(doc, "LAMPIRAN", 1)
    add_p(doc, "Lampiran 1. Kisi-kisi instrumen loyalitas berbahasa Indonesia", bold=True, size=11)
    add_table(
        doc,
        ["Dimensi", "Indikator", "No. Butir", "Jumlah"],
        [
            ["Sikap positif", "Keyakinan bahwa bahasa Indonesia penting dan bernilai", "1, 2, 3, 4", "4"],
            ["Kesetiaan penggunaan", "Kecenderungan memilih dan menggunakan bahasa Indonesia", "5, 6, 7, 8", "4"],
            ["Kesadaran norma", "Pemahaman perbedaan ragam baku dan tidak baku sesuai konteks", "9, 10, 11, 12", "4"],
            ["Kebanggaan bahasa", "Rasa bangga menggunakan bahasa Indonesia dengan baik", "13, 14, 15, 16", "4"],
            ["Pemilihan ragam", "Kemampuan menyesuaikan bahasa dengan tujuan, lawan bicara, dan situasi", "17, 18, 19, 20", "4"],
        ],
        widths=[1.4, 3.0, 1.2, 0.6],
    )

    add_p(doc, "Lampiran 2. Butir instrumen kuesioner loyalitas berbahasa Indonesia", bold=True, size=11, before=10)
    add_table(
        doc,
        ["No.", "Dimensi", "Pernyataan", "SS", "S", "TS", "STS"],
        [
            ["1", "Sikap positif", "Bahasa Indonesia tetap penting digunakan di media sosial meskipun banyak istilah asing yang populer.", "", "", "", ""],
            ["2", "Sikap positif", "Saya merasa bahasa Indonesia yang baik menunjukkan identitas sebagai pelajar Indonesia.", "", "", "", ""],
            ["3", "Sikap positif", "Penggunaan bahasa Indonesia yang benar lebih membanggakan daripada campur kode Inggris-Indonesia.", "", "", "", ""],
            ["4", "Sikap positif", "Menurut saya, bahasa Indonesia tidak kalah modern dibandingkan bahasa asing.", "", "", "", ""],
            ["5", "Kesetiaan penggunaan", "Saya memilih menggunakan bahasa Indonesia saat menulis tugas sekolah meskipun teman-teman banyak menggunakan istilah asing.", "", "", "", ""],
            ["6", "Kesetiaan penggunaan", "Saya tetap menggunakan bahasa Indonesia dalam diskusi kelompok meskipun ada teman yang menyelipkan bahasa Inggris.", "", "", "", ""],
            ["7", "Kesetiaan penggunaan", "Saya berusaha menghindari penggunaan singkatan tidak baku (misal: \"yg\", \"dg\", \"pdhl\") dalam komunikasi formal.", "", "", "", ""],
            ["8", "Kesetiaan penggunaan", "Saya lebih memilih menulis caption Indonesia yang baik daripada menulis dalam bahasa Inggris agar terlihat keren.", "", "", "", ""],
            ["9", "Kesadaran norma", "Saya dapat membedakan kapan harus menggunakan bahasa Indonesia formal dan kapan boleh menggunakan bahasa santai.", "", "", "", ""],
            ["10", "Kesadaran norma", "Saya menyadari bahwa bahasa gaul tidak selalu tepat digunakan di lingkungan sekolah.", "", "", "", ""],
            ["11", "Kesadaran norma", "Menurut saya, menggunakan campur kode Indonesia-Inggris secara berlebihan dapat mengurangi kualitas komunikasi.", "", "", "", ""],
            ["12", "Kesadaran norma", "Saya memahami bahwa pemilihan kata perlu disesuaikan dengan siapa saya berbicara (guru, teman, atau orang tua).", "", "", "", ""],
            ["13", "Kebanggaan bahasa", "Saya bangga ketika mampu menulis atau berbicara dalam bahasa Indonesia yang baik dan benar.", "", "", "", ""],
            ["14", "Kebanggaan bahasa", "Saya merasa senang ketika ada teman yang memuji cara saya berbahasa Indonesia.", "", "", "", ""],
            ["15", "Kebanggaan bahasa", "Saya percaya diri menggunakan bahasa Indonesia dalam presentasi di kelas.", "", "", "", ""],
            ["16", "Kebanggaan bahasa", "Menurut saya, mampu berbahasa Indonesia dengan baik adalah sesuatu yang patut dibanggakan.", "", "", "", ""],
            ["17", "Pemilihan ragam", "Saya menyesuaikan bahasa yang saya gunakan saat berbicara dengan guru berbeda dengan saat berbicara dengan teman.", "", "", "", ""],
            ["18", "Pemilihan ragam", "Saya memilih kata yang lebih formal saat menulis pengumuman sekolah daripada saat menulis status WhatsApp.", "", "", "", ""],
            ["19", "Pemilihan ragam", "Saya mampu mengubah kalimat tidak baku menjadi kalimat baku tanpa mengubah maksudnya.", "", "", "", ""],
            ["20", "Pemilihan ragam", "Saya mempertimbangkan konteks dan lawan bicara sebelum memilih ragam bahasa yang akan digunakan.", "", "", "", ""],
        ],
        widths=[0.3, 1.0, 3.2, 0.4, 0.4, 0.4, 0.4],
    )

    add_p(doc, "Lampiran 3. Kisi-kisi skenario kasus bahasa PRIMA+ (kesadaran berbahasa)", bold=True, size=11, before=10)
    add_table(
        doc,
        ["No.", "Konstruk", "Jenis Kasus", "Bentuk Respons"],
        [
            ["1", "Mengenali ragam", "Caption Instagram: campur kode Indonesia-Inggris", "Memilih versi caption yang paling sesuai konteks"],
            ["2", "Mengenali ragam", "Chat WhatsApp: singkatan dan bahasa santai", "Mengidentifikasi situasi yang tepat untuk ragam formal"],
            ["3", "Menilai konteks", "Situasi: presentasi kelas vs ngobrol dengan teman", "Memilih ragam bahasa yang sesuai masing-masing situasi"],
            ["4", "Menilai konteks", "Pengumuman sekolah vs story Instagram", "Membedakan pilihan kata untuk konteks formal dan informal"],
            ["5", "Memberi alasan", "Refleksi pilihan bahasa setelah menjawab skenario", "Menjelaskan secara singkat alasan memilih ragam tertentu"],
        ],
        widths=[0.4, 1.2, 2.8, 2.0],
    )

    add_p(doc, "Lampiran 4. Lembar validasi guru/ahli", bold=True, size=11, before=10)
    add_table(
        doc,
        ["No.", "Aspek", "Indikator Penilaian", "1", "2", "3", "4"],
        [
            ["1", "Materi", "Kesesuaian konten dengan konsep kesadaran berbahasa", "", "", "", ""],
            ["2", "Materi", "Relevansi kasus bahasa dengan keseharian remaja", "", "", "", ""],
            ["3", "Materi", "Kedalaman materi sesuai jenjang siswa", "", "", "", ""],
            ["4", "Bahasa", "Penggunaan bahasa Indonesia yang baik dan benar", "", "", "", ""],
            ["5", "Bahasa", "Kejelasan instruksi dan kalimat", "", "", "", ""],
            ["6", "Bahasa", "Kesesuaian bahasa dengan tingkat pemahaman siswa", "", "", "", ""],
            ["7", "Media", "Kemudahan navigasi dan akses", "", "", "", ""],
            ["8", "Media", "Tampilan visual dan daya tarik", "", "", "", ""],
            ["9", "Media", "Kesesuaian media untuk pembelajaran mandiri", "", "", "", ""],
            ["10", "Instrumen", "Kesesuaian butir dengan indikator loyalitas", "", "", "", ""],
            ["11", "Instrumen", "Kejelasan pernyataan dalam kuesioner", "", "", "", ""],
            ["12", "Instrumen", "Kecukupan jumlah butir per indikator", "", "", "", ""],
        ],
        widths=[0.3, 0.7, 3.2, 0.4, 0.4, 0.4, 0.4],
    )

    add_p(doc, "Lampiran 5. Angket respons/refleksi siswa", bold=True, size=11, before=10)
    add_table(
        doc,
        ["No.", "Pertanyaan", "1", "2", "3", "4"],
        [
            ["1", "PRIMA+ mudah saya akses dan gunakan.", "", "", "", ""],
            ["2", "Tampilan PRIMA+ menarik dan tidak membosankan.", "", "", "", ""],
            ["3", "Soal-soal di PRIMA+ membantu saya memahami konteks bahasa.", "", "", "", ""],
            ["4", "Saya mendapat wawasan baru tentang penggunaan bahasa Indonesia yang tepat.", "", "", "", ""],
            ["5", "Saya merasa lebih percaya diri memilih ragam bahasa setelah menggunakan PRIMA+.", "", "", "", ""],
            ["6", "Saya akan merekomendasikan PRIMA+ kepada teman-teman saya.", "", "", "", ""],
        ],
        widths=[0.3, 4.5, 0.4, 0.4, 0.4, 0.4],
    )
    add_p(doc, "Pertanyaan terbuka:\n7. Hal baru apa yang paling berkesan saat menggunakan PRIMA+?\n8. Saran apa yang ingin kamu sampaikan untuk perbaikan PRIMA+?")

    doc.save(PROPOSAL_OUT)


def build_revision_table():
    doc = setup_doc()
    add_heading(doc, "TABEL REVISI REVIEWER PROPOSAL PRIMA+ OPSI 2026", 1)
    add_p(doc, "Tabel ini memuat penyesuaian proposal berdasarkan catatan reviewer dan revisi tambahan, terutama penguatan kategori Ilmu Sosial dan Humaniora (ISH), konteks lokal Sampit/MAN Kotawaringin Timur, serta penjelasan metodologis jumlah responden.", align=WD_ALIGN_PARAGRAPH.LEFT)
    headers = ["No.", "Catatan Reviewer", "Revisi Tambahan", "Perbaikan yang Dilakukan", "Letak Perbaikan", "Penjelasan"]
    rows = [
        ["1", "Judul menarik, tetapi terlalu panjang.", "Judul terlalu ilmu komputer; penelitian kategori sosial humaniora.", "Judul diganti menjadi: Kesadaran Berbahasa Remaja melalui Media PRIMA+ untuk Menguatkan Loyalitas Bahasa Indonesia di Lingkungan Sekolah.", "Cover dan seluruh identitas judul.", "Judul menempatkan kesadaran berbahasa dan loyalitas bahasa sebagai fokus utama. PRIMA+ disebut sebagai media, bukan platform/aplikasi komputer."],
        ["2", "Latar belakang terlalu umum dan belum didukung data empiris.", "Data Jambi tidak sesuai; cari konteks Sampit.", "Latar belakang diganti dengan konteks MAN Kotawaringin Timur, Sampit, menggunakan data Referensi Kemendikdasmen tentang lokasi sekolah dan jumlah peserta didik.", "BAB 1.1.", "Konteks empiris dipindahkan dari Jambi ke lokasi penelitian. Data lokal sekolah digunakan sebagai dasar lingkungan sosial penelitian."],
        ["3", "Rumusan masalah terlalu luas dan belum menunjukkan indikator jelas.", "", "Rumusan masalah diarahkan pada rancangan media, kelayakan media/instrumen, dan efektivitas terhadap loyalitas berbahasa.", "BAB 1.2.", "Rumusan masalah tetap tiga tetapi fokusnya social-humaniora: kesadaran berbahasa, loyalitas bahasa, dan perilaku bahasa remaja."],
        ["4", "Manfaat penelitian masih umum.", "", "Manfaat diperjelas untuk siswa, guru Bahasa Indonesia, sekolah, dan kajian bahasa/sastra.", "BAB 1.4.", "Manfaat tidak lagi terdengar sebagai manfaat aplikasi, tetapi sebagai kontribusi pembelajaran dan pembinaan bahasa."],
        ["5", "Studi pustaka masih sedikit.", "", "Studi pustaka diperkuat dengan sikap bahasa, loyalitas bahasa, language awareness, remaja Sampit, dan bahasa remaja di media sosial.", "BAB 2.1-2.2 dan Daftar Pustaka.", "Rujukan teknologi/gamifikasi tetap ada, tetapi hanya sebagai pendukung media, bukan payung utama penelitian."],
        ["6", "Metode belum merinci jumlah responden.", "", "Jumlah responden ditetapkan 30 siswa.", "BAB 3.3.", "Angka 30 dijelaskan sebagai 30 pasangan data pretest-posttest untuk uji coba terbatas R&D, cukup untuk analisis awal rata-rata, N-gain, dan uji beda berpasangan."],
        ["7", "Metode belum merinci teknik sampling.", "", "Teknik sampling ditetapkan purposive sampling dengan kriteria siswa kelas X/XI, aktif belajar Bahasa Indonesia, menggunakan media digital, bersedia mengikuti seluruh tahapan, dan mendapat izin sekolah.", "BAB 3.3.", "Sampling tidak diklaim mewakili populasi besar, tetapi sesuai untuk uji coba terbatas media kebahasaan."],
        ["8", "Desain uji efektivitas belum jelas.", "", "Desain ditegaskan sebagai R&D sederhana model ADDIE dengan one-group pretest-posttest.", "BAB 3.3 dan BAB 3.4.", "Pengujian diarahkan pada perubahan loyalitas berbahasa, bukan performa teknis aplikasi."],
        ["9", "Indikator pengukuran loyalitas belum diungkapkan.", "", "Indikator loyalitas dijelaskan: sikap positif, kesetiaan penggunaan, kesadaran norma, kebanggaan bahasa, dan pemilihan ragam.", "BAB 3.3 dan Tabel 3.", "Instrumen menjadi terhubung dengan konstruk social-humaniora."],
        ["10", "Prosedur penelitian belum menunjukkan metode yang jelas.", "", "Prosedur disusun berdasarkan ADDIE: Analyze, Design, Develop, Implement, Evaluate.", "BAB 3.3 dan Gambar 1.", "ADDIE dipakai sebagai alur pengembangan media pembelajaran/kebahasaan, bukan rekayasa perangkat lunak."],
        ["11", "Analisis data kurang kuat untuk mengukur peningkatan loyalitas.", "", "Analisis menggunakan rata-rata, persentase peningkatan, N-gain, paired sample t-test jika normal, dan Wilcoxon jika tidak normal.", "BAB 3.4 dan Tabel 4.", "Analisis cocok dengan desain 30 pasangan data pretest-posttest."],
        ["12", "Daftar pustaka belum lengkap dan perlu perbaikan.", "", "Daftar pustaka disesuaikan dengan ISH: language awareness, loyalitas/sikap bahasa, remaja Sampit, bahasa remaja di media sosial, UKBI, dan R&D.", "Daftar Pustaka.", "Rujukan lokal dan sosial-humaniora ditambahkan agar tidak terlalu teknis komputer."],
        ["13", "Anggaran kurang realistis dan komponen belum rinci.", "", "RAB direvisi menjadi Rp1.987.100, dibatasi di bawah Rp2.000.000, dan dipecah menjadi 20 item konkret: barang habis pakai, jasa cetak, jasa jilid, paket data internet, sewa modem, sewa tripod, transport lokal, konsumsi ringan, dan air mineral.", "BAB 4.1 proposal dan file RAB XLSX.", "Komponen abstrak sebelumnya dihapus. RAB kini menyebut nama barang, sewa, jasa, dan kebutuhan lainnya secara langsung beserta volume penggunaannya."],
    ]
    add_table(doc, headers, rows, widths=[0.35, 1.35, 1.25, 1.65, 0.9, 1.8])
    doc.save(REVISION_OUT)


if __name__ == "__main__":
    build_proposal()
    build_revision_table()
    print(PROPOSAL_OUT)
    print(REVISION_OUT)
