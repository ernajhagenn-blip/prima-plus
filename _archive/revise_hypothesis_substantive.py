from docx import Document


PROPOSAL = "Proposal_Penelitian_PRIMA_OPSI_2026_Format_Resmi.docx"
REVISION_TABLE = "Tabel_Revisi_Sistematika_OPSI_PRIMA.docx"


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def revise_proposal():
    doc = Document(PROPOSAL)

    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "1.4 HIPOTESIS":
            doc.paragraphs[idx + 1].text = (
                "Hipotesis penelitian ini adalah: penggunaan platform PRIMA+ berbasis "
                "kesadaran berbahasa diduga dapat meningkatkan loyalitas berbahasa "
                "Indonesia remaja. Hipotesis ini berkaitan dengan rumusan masalah "
                "ketiga tentang efektivitas penggunaan PRIMA+. Rumusan masalah "
                "pertama berfokus pada pengembangan rancangan produk, sedangkan "
                "rumusan masalah kedua berfokus pada kelayakan produk dan instrumen, "
                "sehingga keduanya dijawab melalui proses pengembangan dan validasi, "
                "bukan melalui hipotesis tersendiri."
            )
            while doc.paragraphs[idx + 2].text.startswith(("H0:", "H1:")):
                remove_paragraph(doc.paragraphs[idx + 2])
            break

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Data validasi guru/ahli dihitung dalam bentuk persentase kelayakan."):
            paragraph.text = (
                "Data validasi guru/ahli dihitung dalam bentuk persentase kelayakan. "
                "Skor pretest dan posttest loyalitas berbahasa dibandingkan melalui "
                "rata-rata, persentase peningkatan, dan N-gain. Pengujian perbedaan "
                "skor dilakukan sebagai dasar untuk menilai apakah hipotesis penelitian "
                "didukung oleh data. Jika data memenuhi syarat, digunakan paired sample "
                "t-test; jika tidak memenuhi syarat, digunakan uji Wilcoxon. Refleksi "
                "siswa dianalisis secara tematik untuk melihat pengalaman, kendala, "
                "dan saran setelah menggunakan PRIMA+."
            )
            break

    doc.save(PROPOSAL)


def revise_revision_table():
    doc = Document(REVISION_TABLE)
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if cells and cells[0].text.strip() == "3" and "Hipotesis" in cells[1].text:
                cells[2].text = (
                    "Hipotesis sebelumnya belum dibedakan secara tegas antara "
                    "hipotesis penelitian dan teknis pengujian data."
                )
                cells[3].text = (
                    "Bab 1.4 diubah menjadi hipotesis penelitian/substantif: "
                    "penggunaan PRIMA+ diduga dapat meningkatkan loyalitas berbahasa "
                    "Indonesia remaja. Penjelasan juga ditambahkan bahwa hipotesis "
                    "hanya terkait rumusan masalah ketiga, sedangkan rumusan pertama "
                    "dan kedua dijawab melalui pengembangan produk dan validasi "
                    "kelayakan. Rancangan uji data ditempatkan di Bab 3.4."
                )
                cells[4].text = "BAB 1.4 dan BAB 3.4"
                break
    doc.save(REVISION_TABLE)


if __name__ == "__main__":
    revise_proposal()
    revise_revision_table()
    print("Revised hypothesis wording and revision table.")
