from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches


OUT = "Instrumen_Penelitian_PRIMA.docx"


def set_run(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_p(paragraph, align=None, before=0, after=6, line=1.15):
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_p(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=11, bold=False, after=6, before=0):
    p = doc.add_paragraph()
    set_p(p, align=align, after=after, before=before)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    set_p(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=10 if level == 1 else 6, after=5)
    r = p.add_run(text)
    set_run(r, size=14 if level == 1 else 12, bold=True)
    return p


def shade_cell(cell, fill="EDEDED"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    set_p(p, align=align, after=0, line=1.0)
    r = p.add_run(str(text))
    set_run(r, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, caption=None, widths=None):
    if caption:
        add_p(doc, caption, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, after=3)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)

    # ── COVER ──
    for _ in range(4):
        doc.add_paragraph()
    add_p(doc, "INSTRUMEN PENELITIAN", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, after=16)
    add_p(doc, "PRIMA+", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, after=16)
    add_p(doc, "Kesadaran Berbahasa Remaja melalui Media PRIMA+", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=4)
    add_p(doc, "untuk Menguatkan Loyalitas Bahasa Indonesia", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=4)
    add_p(doc, "di Lingkungan Sekolah", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=20)
    add_p(doc, "MAN Kotawaringin Timur", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=4)
    add_p(doc, "Tahun 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_page_break()

    # ── A. KISI-KISI LOYALITAS ──
    add_heading(doc, "A. KISI-KISI INSTRUMEN LOYALITAS BERBAHASA INDONESIA", 1)
    add_table(
        doc,
        ["Dimensi", "Indikator", "No. Butir", "Jumlah"],
        [
            ["Sikap positif", "Keyakinan bahwa bahasa Indonesia penting dan bernilai", "1, 2, 3, 4", "4"],
            ["Kesetiaan penggunaan", "Kecenderungan memilih dan menggunakan bahasa Indonesia", "5, 6, 7, 8", "4"],
            ["Kesadaran norma", "Pemahaman perbedaan ragam baku dan tidak baku sesuai konteks", "9, 10, 11, 12", "4"],
            ["Kebanggaan bahasa", "Rasa bangga menggunakan bahasa Indonesia dengan baik", "13, 14, 15, 16", "4"],
            ["Pemilihan ragam", "Kemampuan menyesuaikan bahasa dengan tujuan, lawan bicara, dan situasi", "17, 18, 19, 20", "4"],
        ],
        widths=[1.3, 3.0, 1.0, 0.6],
    )

    # ── B. KUESIONER LOYALITAS ──
    add_heading(doc, "B. KUESIONER LOYALITAS BERBAHASA INDONESIA", 1)
    add_p(doc, "Petunjuk: Berilah tanda centang (√) pada kolom yang sesuai dengan pendapatmu.", before=4)
    add_p(doc, "SS = Sangat Setuju, S = Setuju, TS = Tidak Setuju, STS = Sangat Tidak Setuju", after=8)
    add_table(
        doc,
        ["No.", "Dimensi", "Pernyataan", "SS", "S", "TS", "STS"],
        [
            ["1", "Sikap positif", "Bahasa Indonesia tetap penting digunakan di media sosial meskipun banyak istilah asing yang populer.", "", "", "", ""],
            ["2", "Sikap positif", "Saya merasa bahasa Indonesia yang baik menunjukkan identitas sebagai pelajar Indonesia.", "", "", "", ""],
            ["3", "Sikap positif", "Penggunaan bahasa Indonesia yang benar lebih membanggakan daripada campur kode Inggris-Indonesia.", "", "", "", ""],
            ["4", "Sikap positif", "Menurut saya, bahasa Indonesia tidak kalah modern dibandingkan bahasa asing.", "", "", "", ""],
            ["5", "Kesetiaan", "Saya memilih menggunakan bahasa Indonesia saat menulis tugas sekolah meskipun teman-teman banyak menggunakan istilah asing.", "", "", "", ""],
            ["6", "Kesetiaan", "Saya tetap menggunakan bahasa Indonesia dalam diskusi kelompok meskipun ada teman yang menyelipkan bahasa Inggris.", "", "", "", ""],
            ["7", "Kesetiaan", "Saya berusaha menghindari penggunaan singkatan tidak baku (misal: \"yg\", \"dg\", \"pdhl\") dalam komunikasi formal.", "", "", "", ""],
            ["8", "Kesetiaan", "Saya lebih memilih menulis caption Indonesia yang baik daripada menulis dalam bahasa Inggris agar terlihat keren.", "", "", "", ""],
            ["9", "Kesadaran norma", "Saya dapat membedakan kapan harus menggunakan bahasa Indonesia formal dan kapan boleh menggunakan bahasa santai.", "", "", "", ""],
            ["10", "Kesadaran norma", "Saya menyadari bahwa bahasa gaul tidak selalu tepat digunakan di lingkungan sekolah.", "", "", "", ""],
            ["11", "Kesadaran norma", "Menurut saya, menggunakan campur kode Indonesia-Inggris secara berlebihan dapat mengurangi kualitas komunikasi.", "", "", "", ""],
            ["12", "Kesadaran norma", "Saya memahami bahwa pemilihan kata perlu disesuaikan dengan siapa saya berbicara (guru, teman, atau orang tua).", "", "", "", ""],
            ["13", "Kebanggaan", "Saya bangga ketika mampu menulis atau berbicara dalam bahasa Indonesia yang baik dan benar.", "", "", "", ""],
            ["14", "Kebanggaan", "Saya merasa senang ketika ada teman yang memuji cara saya berbahasa Indonesia.", "", "", "", ""],
            ["15", "Kebanggaan", "Saya percaya diri menggunakan bahasa Indonesia dalam presentasi di kelas.", "", "", "", ""],
            ["16", "Kebanggaan", "Menurut saya, mampu berbahasa Indonesia dengan baik adalah sesuatu yang patut dibanggakan.", "", "", "", ""],
            ["17", "Pemilihan ragam", "Saya menyesuaikan bahasa yang saya gunakan saat berbicara dengan guru berbeda dengan saat berbicara dengan teman.", "", "", "", ""],
            ["18", "Pemilihan ragam", "Saya memilih kata yang lebih formal saat menulis pengumuman sekolah daripada saat menulis status WhatsApp.", "", "", "", ""],
            ["19", "Pemilihan ragam", "Saya mampu mengubah kalimat tidak baku menjadi kalimat baku tanpa mengubah maksudnya.", "", "", "", ""],
            ["20", "Pemilihan ragam", "Saya mempertimbangkan konteks dan lawan bicara sebelum memilih ragam bahasa yang akan digunakan.", "", "", "", ""],
        ],
        widths=[0.3, 0.8, 3.2, 0.4, 0.4, 0.4, 0.4],
    )

    # ── C. SKENARIO KUIS PRIMA+ ──
    add_heading(doc, "C. SKENARIO KASUS BAHASA — KUIS PRIMA+", 1)
    add_p(doc, "Petunjuk: Siswa membaca skenario dan memilih respons bahasa yang paling tepat, kemudian menuliskan alasan singkat.", before=4, after=8)
    add_table(
        doc,
        ["No.", "Konstruk", "Jenis Kasus", "Tugas Siswa"],
        [
            ["1", "Mengenali ragam", "Caption Instagram: \"Happy weekend guys! Let's hangout yuk!\"", "Pilih versi caption yang lebih sesuai konteks sekolah"],
            ["2", "Mengenali ragam", "Chat WhatsApp: \"Bsk jm 7 kumpul d ruang kelas yaa\"", "Ubahlah ke bahasa Indonesia yang lebih tertib"],
            ["3", "Menilai konteks", "Situasi A: presentasi di depan kelas\nSituasi B: ngobrol dengan teman saat istirahat", "Pilih ragam bahasa yang sesuai untuk masing-masing situasi"],
            ["4", "Menilai konteks", "Teks pengumuman sekolah vs. story Instagram", "Bandingkan pilihan kata: mana yang formal dan mana yang informal"],
            ["5", "Memberi alasan", "Setelah memilih ragam bahasa pada soal 1-4", "Tuliskan alasan singkat mengapa memilih ragam tersebut"],
        ],
        widths=[0.3, 1.0, 2.5, 2.5],
    )

    # ── D. LEMBAR VALIDASI ──
    add_heading(doc, "D. LEMBAR VALIDASI GURU/AHLI", 1)
    add_p(doc, "Petunjuk: Berilah tanda centang (√) pada kolom skor 1-4 sesuai penilaian Bapak/Ibu.", before=4)
    add_p(doc, "1 = Kurang, 2 = Cukup, 3 = Baik, 4 = Sangat Baik", after=8)
    add_table(
        doc,
        ["No.", "Aspek", "Indikator Penilaian", "1", "2", "3", "4"],
        [
            ["1", "Materi", "Kesesuaian konten dengan konsep kesadaran berbahasa", "", "", "", ""],
            ["2", "Materi", "Relevansi kasus bahasa dengan keseharian remaja", "", "", "", ""],
            ["3", "Materi", "Kedalaman materi sesuai jenjang siswa", "", "", "", ""],
            ["4", "Bahasa", "Penggunaan bahasa Indonesia yang baik dan benar", "", "", "", ""],
            ["5", "Bahasa", "Kejelasan instruksi dan kalimat", "", "", "", ""],
            ["6", "Bahasa", "Kesesuaian bahasa dengan tingkat pemahaman siswa", "", "", "", ""],
            ["7", "Media", "Kemudahan navigasi dan akses", "", "", "", ""],
            ["8", "Media", "Tampilan visual dan daya tarik", "", "", "", ""],
            ["9", "Media", "Kesesuaian media untuk pembelajaran mandiri", "", "", "", ""],
            ["10", "Instrumen", "Kesesuaian butir dengan indikator loyalitas", "", "", "", ""],
            ["11", "Instrumen", "Kejelasan pernyataan dalam kuesioner", "", "", "", ""],
            ["12", "Instrumen", "Kecukupan jumlah butir per indikator", "", "", "", ""],
        ],
        widths=[0.3, 0.7, 3.0, 0.4, 0.4, 0.4, 0.4],
    )
    add_p(doc, "Komentar/Saran: _____________________________________________________________", before=10)
    add_p(doc, "Kesimpulan: [] Layak digunakan [] Layak dengan revisi [] Tidak layak", bold=False)

    # ── E. ANGKET RESPONS SISWA ──
    add_heading(doc, "E. ANGKET RESPONS/REFLEKSI SISWA", 1)
    add_p(doc, "Petunjuk: Berilah tanda centang (√) pada kolom yang sesuai dengan pengalamanmu menggunakan PRIMA+.", before=4)
    add_p(doc, "1 = Sangat Tidak Setuju, 2 = Tidak Setuju, 3 = Setuju, 4 = Sangat Setuju", after=8)
    add_table(
        doc,
        ["No.", "Pernyataan", "1", "2", "3", "4"],
        [
            ["1", "PRIMA+ mudah saya akses dan gunakan.", "", "", "", ""],
            ["2", "Tampilan PRIMA+ menarik dan tidak membosankan.", "", "", "", ""],
            ["3", "Soal-soal di PRIMA+ membantu saya memahami konteks bahasa.", "", "", "", ""],
            ["4", "Saya mendapat wawasan baru tentang penggunaan bahasa Indonesia yang tepat.", "", "", "", ""],
            ["5", "Saya merasa lebih percaya diri memilih ragam bahasa setelah menggunakan PRIMA+.", "", "", "", ""],
            ["6", "Saya akan merekomendasikan PRIMA+ kepada teman-teman saya.", "", "", "", ""],
        ],
        widths=[0.3, 4.5, 0.4, 0.4, 0.4, 0.4],
    )
    add_p(doc, "Pertanyaan terbuka:", before=10)
    add_p(doc, "7. Hal baru apa yang paling berkesan saat menggunakan PRIMA+?")
    add_p(doc, "   Jawab: ___________________________________________________________")
    add_p(doc, "8. Saran apa yang ingin kamu sampaikan untuk perbaikan PRIMA+?")
    add_p(doc, "   Jawab: ___________________________________________________________")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
