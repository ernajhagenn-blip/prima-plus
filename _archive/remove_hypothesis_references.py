from docx import Document


PROPOSAL = "Proposal_Penelitian_PRIMA_OPSI_2026_Format_Resmi_RAB_Rinci.docx"
REVISION_TABLE = "Tabel_Revisi_Sistematika_OPSI_PRIMA.docx"


def revise_proposal():
    doc = Document(PROPOSAL)
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Rancangan pengolahan data disesuaikan dengan tiga hipotesis penelitian."):
            paragraph.text = (
                "Rancangan pengolahan data disesuaikan dengan tiga rumusan masalah dan "
                "tujuan penelitian. Rumusan masalah pertama dijawab melalui deskripsi "
                "proses pengembangan PRIMA+ berdasarkan tahapan ADDIE, kelengkapan "
                "fitur, dan kesesuaian prototipe dengan kebutuhan siswa. Rumusan masalah "
                "kedua dijawab melalui persentase kelayakan dari validasi guru atau ahli "
                "pada aspek materi, bahasa, media, dan instrumen. Rumusan masalah ketiga "
                "dijawab melalui perbandingan skor pretest dan posttest loyalitas "
                "berbahasa, rata-rata, persentase peningkatan, dan N-gain. Jika data "
                "memenuhi syarat, digunakan paired sample t-test; jika tidak memenuhi "
                "syarat, digunakan uji Wilcoxon. Refleksi siswa dianalisis secara tematik "
                "untuk melihat pengalaman, kendala, dan saran setelah menggunakan PRIMA+."
            )
            break
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == "Hipotesis diterima bila peningkatan bermakna.":
                    cell.text = "Peningkatan dinilai bermakna bila hasil uji menunjukkan perbedaan yang signifikan."
    doc.save(PROPOSAL)


def revise_revision_table():
    doc = Document(REVISION_TABLE)
    table = doc.tables[0]
    for row in list(table.rows):
        if all(not cell.text.strip() for cell in row.cells):
            table._tbl.remove(row._tr)
    existing = [
        " || ".join(cell.text.strip() for cell in row.cells)
        for row in table.rows
    ]
    if any("Penghapusan Hipotesis" in row for row in existing):
        doc.save(REVISION_TABLE)
        return
    row = table.add_row()
    row.cells[0].text = "10"
    row.cells[1].text = "Penghapusan Hipotesis"
    row.cells[2].text = (
        "Bab 1.4 Hipotesis dihapus karena desain penelitian lebih tepat diposisikan "
        "sebagai penelitian pengembangan dan uji coba terbatas, bukan penelitian "
        "yang berangkat dari hipotesis formal."
    )
    row.cells[3].text = (
        "Subbab Manfaat disesuaikan menjadi Bab 1.4. Rujukan terhadap hipotesis pada "
        "Bab 3.4 dihapus dan diganti dengan pemetaan pengolahan data berdasarkan "
        "tiga rumusan masalah dan tujuan penelitian."
    )
    row.cells[4].text = "BAB 1.4 dan BAB 3.4"
    doc.save(REVISION_TABLE)


if __name__ == "__main__":
    revise_proposal()
    revise_revision_table()
    print("Removed hypothesis references and aligned analysis section.")
