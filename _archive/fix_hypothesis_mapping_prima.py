from pathlib import Path
from docx import Document


FILES = [
    Path("Proposal_Penelitian_PRIMA_OPSI_2026_Format_Resmi.docx"),
]


EXPLANATION = (
    "Hipotesis hanya dirumuskan untuk rumusan masalah ketiga karena bagian tersebut "
    "menguji efektivitas PRIMA+ melalui perbandingan skor pretest dan posttest. "
    "Rumusan masalah pertama bersifat pengembangan produk, sedangkan rumusan masalah "
    "kedua bersifat penilaian kelayakan, sehingga keduanya tidak memerlukan hipotesis statistik."
)

H0 = (
    "H0: Tidak terdapat peningkatan loyalitas berbahasa Indonesia remaja setelah menggunakan "
    "platform PRIMA+."
)
H1 = (
    "H1: Terdapat peningkatan loyalitas berbahasa Indonesia remaja setelah menggunakan "
    "platform PRIMA+."
)


def set_para_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def update_doc(path):
    doc = Document(str(path))
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "1.4 HIPOTESIS":
            # Replace the two old hypothesis paragraphs and insert the explanation.
            if idx + 1 < len(doc.paragraphs):
                set_para_text(doc.paragraphs[idx + 1], EXPLANATION)
            if idx + 2 < len(doc.paragraphs):
                set_para_text(doc.paragraphs[idx + 2], H0)
            if idx + 3 < len(doc.paragraphs):
                set_para_text(doc.paragraphs[idx + 3], H1)
            else:
                doc.paragraphs[idx + 2].insert_paragraph_before(H1)
            break
    doc.save(str(path))


def update_revision_table():
    path = Path("Tabel_Revisi_Sistematika_OPSI_PRIMA.docx")
    if not path.exists():
        return
    doc = Document(str(path))
    table = doc.tables[0]
    found = False
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) > 1 and cells[1] == "Hipotesis":
            row.cells[2].text = (
                "Hipotesis sebelumnya belum menjelaskan hubungan dengan rumusan masalah, "
                "sehingga tampak tidak seimbang dengan tiga rumusan masalah."
            )
            row.cells[3].text = (
                "Ditambahkan penjelasan bahwa hipotesis hanya berlaku untuk rumusan masalah "
                "ketiga yang menguji efektivitas, sedangkan rumusan pertama dan kedua bersifat "
                "pengembangan dan validasi kelayakan."
            )
            row.cells[4].text = "BAB 1.4"
            found = True
            break
    if not found:
        row = table.add_row()
        values = [
            "10",
            "Hipotesis",
            "Hipotesis belum dipetakan ke rumusan masalah sehingga terlihat tidak konsisten.",
            "Ditambahkan penjelasan bahwa hipotesis hanya untuk rumusan masalah ketiga; rumusan pertama dan kedua tidak memerlukan hipotesis statistik.",
            "BAB 1.4",
        ]
        for i, value in enumerate(values):
            row.cells[i].text = value
    doc.save(str(path))


if __name__ == "__main__":
    for file in FILES:
        if file.exists():
            update_doc(file)
            print(f"updated {file}")
    update_revision_table()
    print("updated revision table")
