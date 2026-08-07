from pathlib import Path
from shutil import copyfile

from docx import Document


PROPOSAL_SOURCE = Path("prima_revisi_HIGH_HUMANIZED_V2.docx")
TABLE_SOURCE = Path("tabel_review_fundamental_HIGH.docx")

PROPOSAL_OUT = Path("Proposal_Penelitian_PRIMA_Revisi_Final.docx")
TABLE_OUT = Path("Tabel_Daftar_Revisi_Proposal_PRIMA.docx")


def scrub_core_properties(doc, title):
    props = doc.core_properties
    props.title = title
    props.subject = ""
    props.author = ""
    props.keywords = ""
    props.comments = ""
    props.category = ""
    props.last_modified_by = ""
    props.revision = 1


def prepare_proposal():
    copyfile(PROPOSAL_SOURCE, PROPOSAL_OUT)
    doc = Document(str(PROPOSAL_OUT))
    scrub_core_properties(doc, "Proposal Penelitian PRIMA+ Revisi Final")
    doc.save(str(PROPOSAL_OUT))


def prepare_revision_table():
    copyfile(TABLE_SOURCE, TABLE_OUT)
    doc = Document(str(TABLE_OUT))

    replacements = {
        "TABEL REVIEW FUNDAMENTAL VERSI HIGH": "TABEL DAFTAR REVISI PROPOSAL PRIMA+",
        "Tabel ini menjelaskan perombakan mendasar yang dilakukan pada naskah revisi high review.": (
            "Tabel ini merangkum bagian proposal yang telah direvisi, alasan perbaikan, "
            "bentuk perubahan, dan letak perbaikannya pada naskah."
        ),
        "Perombakan High Review": "Perbaikan yang Dilakukan",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = replacements[text]
            else:
                paragraph.add_run(replacements[text])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text in replacements:
                    cell.text = replacements[text]

    scrub_core_properties(doc, "Tabel Daftar Revisi Proposal PRIMA+")
    doc.save(str(TABLE_OUT))


if __name__ == "__main__":
    if not PROPOSAL_SOURCE.exists():
        raise FileNotFoundError(PROPOSAL_SOURCE)
    if not TABLE_SOURCE.exists():
        raise FileNotFoundError(TABLE_SOURCE)
    prepare_proposal()
    prepare_revision_table()
    print(PROPOSAL_OUT)
    print(TABLE_OUT)
