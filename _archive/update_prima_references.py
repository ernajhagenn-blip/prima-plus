from docx import Document
from docx.shared import Cm


REFERENCES = [
    "Al-Dosakee, K., & Ozdamli, F. (2021). Gamification in teaching and learning languages: A systematic literature review. Revista Romaneasca pentru Educatie Multidimensionala, 13(2), 559-577. https://doi.org/10.18662/rrem/13.2/436",
    "Alwi, H., & Sugono, D. (2011). Politik Bahasa Nasional. Badan Pengembangan dan Pembinaan Bahasa.",
    "Amelia, S. R., Siroj, M. B., & Afgani, M. W. (2024). Analisis perkembangan bahasa Indonesia di era digital: Tantangan dan peluang. Jurnal Multidisiplin Dehasen, 3(3), 125-132. https://doi.org/10.37676/mude.v3i3.5792",
    "Badan Pengembangan dan Pembinaan Bahasa. (2023). Transformasi digital dalam pemelajaran BIPA: Peluang dan tantangan. Kementerian Pendidikan, Kebudayaan, Riset, dan Teknologi.",
    "Balai Bahasa Provinsi Jambi. (2025). UKBI jadi instrumen strategis wujudkan kemahiran berbahasa Indonesia yang terstandar. Kementerian Pendidikan Dasar dan Menengah.",
    "Carter, R. (2003). Language awareness. ELT Journal, 57(1), 64-65. https://doi.org/10.1093/elt/57.1.64",
    "Chaer, A., & Agustina, L. (2014). Sosiolinguistik: Perkenalan awal. Rineka Cipta.",
    "Crystal, D. (2011). Internet Linguistics: A Student Guide. Routledge.",
    "Fairclough, N. (1992). Critical Language Awareness. Longman.",
    "Garvin, P. L., & Mathiot, M. (1968). The urbanization of the Guarani language: A problem in language and culture. In J. A. Fishman (Ed.), Readings in the Sociology of Language (pp. 365-374). De Gruyter Mouton. https://doi.org/10.1515/9783110805376.365",
    "Luo, Z. (2023). The effectiveness of gamified tools for foreign language learning (FLL): A systematic review. Behavioral Sciences, 13(4), 331. https://doi.org/10.3390/bs13040331",
    "Prensky, M. (2001). Digital natives, digital immigrants. On the Horizon, 9(5), 1-6.",
    "Purba, M. A., Simanjuntak, F., & Sihombing, S. (2024). Menganalisis penggunaan campur kode dalam media sosial. Atmosfer: Jurnal Pendidikan, Bahasa, Sastra, Seni, Budaya, dan Sosial Humaniora, 2(4), 32-38. https://doi.org/10.59024/atmosfer.v2i4.1060",
    "Shortt, M., Tilak, S., Kuznetcova, I., Martens, B., & Akinkuolie, B. (2021). Gamification in mobile-assisted language learning: A systematic review. Computer Assisted Language Learning. https://doi.org/10.1080/09588221.2021.1933540",
    "Sugiyono. (2019). Metode Penelitian dan Pengembangan: Research and Development. Alfabeta.",
    "Sugono, D. (2019). Bahasa Indonesia di era globalisasi. Badan Pengembangan dan Pembinaan Bahasa.",
]


def set_hanging(paragraph):
    paragraph.paragraph_format.first_line_indent = Cm(-0.75)
    paragraph.paragraph_format.left_indent = Cm(0.75)


def update_proposal():
    doc = Document("prima_revisi.docx")
    start = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "DAFTAR PUSTAKA":
            start = idx
            break
    if start is None:
        raise RuntimeError("DAFTAR PUSTAKA heading not found")

    body = doc._body._element
    for paragraph in doc.paragraphs[start + 1 :]:
        body.remove(paragraph._element)

    heading = doc.paragraphs[start]
    heading.style = "Heading 1"
    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        set_hanging(p)

    doc.save("prima_revisi.docx")


def update_review_table():
    doc = Document("tabel_revisi_reviewer.docx")
    table = doc.tables[0]
    for row in table.rows:
        first = row.cells[0].text.strip()
        note = row.cells[1].text.strip().lower()
        if first == "12" or "daftar pustaka" in note:
            row.cells[2].text = (
                "Daftar pustaka diperluas dan disesuaikan dari 8 menjadi 16 sumber. "
                "Referensi ditambahkan untuk teori language awareness, loyalitas/sikap bahasa, "
                "bahasa digital dan campur kode, gamifikasi pembelajaran bahasa, data empiris UKBI, "
                "serta metode R&D."
            )
            row.cells[3].text = "Daftar Pustaka dan BAB 2 Tinjauan Pustaka."
            row.cells[4].text = (
                "Perbaikan tidak hanya menambah jumlah sumber, tetapi juga menyesuaikan fungsi rujukan "
                "dengan bagian naskah: Fairclough/Carter untuk language awareness, Garvin-Mathiot/Chaer "
                "untuk loyalitas bahasa, Crystal dan kajian 2024 untuk bahasa digital/campur kode, "
                "Shortt/Luo/Al-Dosakee untuk gamifikasi, UKBI untuk data empiris, dan Sugiyono untuk metode."
            )
            break
    doc.save("tabel_revisi_reviewer.docx")


if __name__ == "__main__":
    update_proposal()
    update_review_table()
    print("updated prima_revisi.docx references")
    print("updated tabel_revisi_reviewer.docx reviewer row")
