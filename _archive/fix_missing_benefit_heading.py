from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


PATH = Path("Proposal_Penelitian_PRIMA_OPSI_2026_Format_Resmi.docx")


def style_heading_run(run):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = True


doc = Document(str(PATH))
has_heading = any(p.text.strip() == "1.5 MANFAAT PENELITIAN" for p in doc.paragraphs)
if not has_heading:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Bagi siswa, PRIMA+"):
            new_p = paragraph.insert_paragraph_before("1.5 MANFAAT PENELITIAN")
            new_p.style = "Heading 2"
            for run in new_p.runs:
                style_heading_run(run)
            break
doc.save(str(PATH))
print("ensured 1.5 MANFAAT PENELITIAN")
