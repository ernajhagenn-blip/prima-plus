from docx import Document


PROPOSAL = "Proposal_Penelitian_PRIMA_OPSI_2026_Format_Resmi_RAB_Rinci.docx"
PROPOSAL_OUT = "Proposal_Penelitian_PRIMA_OPSI_2026_Format_Resmi_RAB_Rinci_Hipotesis_Sinkron.docx"
REVISION_TABLE = "Tabel_Revisi_Sistematika_OPSI_PRIMA.docx"
REVISION_TABLE_OUT = "Tabel_Revisi_Sistematika_OPSI_PRIMA_Hipotesis_Sinkron.docx"


RQ = [
    "1. Bagaimana rancangan platform PRIMA+ berbasis kesadaran berbahasa untuk menguatkan loyalitas bahasa Indonesia remaja?",
    "2. Bagaimana kelayakan materi, bahasa, media, dan instrumen PRIMA+ berdasarkan validasi guru atau ahli?",
    "3. Bagaimana efektivitas penggunaan PRIMA+ terhadap peningkatan loyalitas berbahasa Indonesia remaja?",
]

OBJECTIVES = [
    "1. Mengembangkan platform PRIMA+ berbasis kesadaran berbahasa sesuai kebutuhan remaja pengguna media digital.",
    "2. Menilai kelayakan PRIMA+ dari aspek materi, bahasa, media, dan instrumen penelitian.",
    "3. Menguji efektivitas PRIMA+ dalam menguatkan loyalitas berbahasa Indonesia remaja melalui desain pretest-posttest.",
]

HYPOTHESES = [
    "1. PRIMA+ diduga dapat dirancang sebagai platform digital berbasis kesadaran berbahasa yang memuat kasus bahasa digital, kuis kontekstual, refleksi pilihan bahasa, skor, dan umpan balik untuk menguatkan loyalitas bahasa Indonesia remaja.",
    "2. PRIMA+ diduga layak digunakan berdasarkan validasi guru atau ahli pada aspek materi, bahasa, media, dan instrumen.",
    "3. Penggunaan PRIMA+ diduga dapat meningkatkan loyalitas berbahasa Indonesia remaja setelah siswa mengikuti kegiatan uji coba terbatas.",
]


def remove_following_until_heading(doc, start_idx):
    while start_idx + 1 < len(doc.paragraphs):
        next_paragraph = doc.paragraphs[start_idx + 1]
        if next_paragraph.style.name.startswith("Heading"):
            break
        element = next_paragraph._element
        element.getparent().remove(element)


def insert_after(paragraph, text, style="Normal"):
    new_p = paragraph._p.addnext(paragraph._p.__class__())
    inserted = paragraph._parent.paragraphs[paragraph._parent.paragraphs.index(paragraph) + 1]
    inserted.text = text
    inserted.style = style
    return inserted


def add_paragraph_after(paragraph, text, style="Normal"):
    new_paragraph = paragraph.insert_paragraph_before(text)
    paragraph._p.addprevious(new_paragraph._p)


def append_after(paragraph, text, style="Normal"):
    body = paragraph._parent
    new_p = paragraph._p.addnext(paragraph._p.__class__())
    # python-docx does not expose the newly inserted XML as a Paragraph directly;
    # recreate the paragraph object by locating the insertion position.
    idx = list(body._element).index(paragraph._p) + 1
    new_paragraph = body.paragraphs[idx]
    new_paragraph.text = text
    new_paragraph.style = style
    return new_paragraph


def replace_section_items(doc, heading_text, items):
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == heading_text:
            remove_following_until_heading(doc, idx)
            anchor = paragraph
            for item in items:
                anchor = append_after(anchor, item, "Normal")
            return
    raise RuntimeError(f"Heading not found: {heading_text}")


def revise_proposal():
    doc = Document(PROPOSAL)
    replace_section_items(doc, "1.2 RUMUSAN MASALAH", RQ)
    replace_section_items(doc, "1.3 TUJUAN PENELITIAN", OBJECTIVES)
    replace_section_items(doc, "1.4 HIPOTESIS", HYPOTHESES)

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Data validasi guru/ahli dihitung dalam bentuk persentase kelayakan."):
            paragraph.text = (
                "Rancangan pengolahan data disesuaikan dengan tiga hipotesis penelitian. "
                "Hipotesis pertama dinilai melalui kelengkapan rancangan dan prototipe PRIMA+ "
                "berdasarkan tahapan ADDIE. Hipotesis kedua dinilai melalui persentase "
                "kelayakan dari validasi guru atau ahli pada aspek materi, bahasa, media, "
                "dan instrumen. Hipotesis ketiga dinilai melalui perbandingan skor pretest "
                "dan posttest loyalitas berbahasa, rata-rata, persentase peningkatan, dan "
                "N-gain. Jika data memenuhi syarat, digunakan paired sample t-test; jika "
                "tidak memenuhi syarat, digunakan uji Wilcoxon. Refleksi siswa dianalisis "
                "secara tematik untuk melihat pengalaman, kendala, dan saran setelah "
                "menggunakan PRIMA+."
            )
            break

    try:
        doc.save(PROPOSAL)
        return PROPOSAL
    except PermissionError:
        doc.save(PROPOSAL_OUT)
        return PROPOSAL_OUT


def revise_revision_table():
    doc = Document(REVISION_TABLE)
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if cells and cells[0].text.strip() == "3" and "Hipotesis" in cells[1].text:
                cells[2].text = (
                    "Rumusan masalah berjumlah tiga, tetapi hipotesis sebelumnya belum "
                    "ditulis sejajar dalam tiga pointer. Bentuk paragraf membuat relasi "
                    "RM-Tujuan-Hipotesis kurang tegas."
                )
                cells[3].text = (
                    "Bab 1.2, 1.3, dan 1.4 disinkronkan menjadi tiga pointer sejajar. "
                    "RM1 terhubung dengan tujuan 1 dan hipotesis 1 tentang rancangan "
                    "PRIMA+; RM2 terhubung dengan tujuan 2 dan hipotesis 2 tentang "
                    "kelayakan; RM3 terhubung dengan tujuan 3 dan hipotesis 3 tentang "
                    "efektivitas peningkatan loyalitas berbahasa."
                )
                cells[4].text = "BAB 1.2, BAB 1.3, BAB 1.4, dan BAB 3.4"
                break
    try:
        doc.save(REVISION_TABLE)
        return REVISION_TABLE
    except PermissionError:
        doc.save(REVISION_TABLE_OUT)
        return REVISION_TABLE_OUT


if __name__ == "__main__":
    proposal_path = revise_proposal()
    revision_table_path = revise_revision_table()
    print(f"proposal={proposal_path}")
    print(f"revision_table={revision_table_path}")
