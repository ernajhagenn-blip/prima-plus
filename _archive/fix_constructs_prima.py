from pathlib import Path
from docx import Document


FILES = [
    Path("Proposal_Penelitian_PRIMA_Revisi_Final.docx"),
    Path("prima_revisi_HIGH_HUMANIZED_V2.docx"),
]


CONSTRUCT_ROWS = [
    [
        "Kesadaran berbahasa",
        "Kepekaan siswa dalam mengenali bentuk bahasa, fungsi bahasa, konteks pemakaian, dan dampak sosial dari pilihan bahasa.",
        "Mengenali ragam bahasa; menilai kesesuaian bahasa dengan konteks; menjelaskan alasan memilih ragam bahasa; merefleksikan dampak pilihan bahasa.",
        "Tes skenario/kuis PRIMA+ dan refleksi singkat.",
        "Skor kuis dan catatan refleksi.",
    ],
    [
        "Loyalitas berbahasa Indonesia",
        "Kemauan siswa untuk tetap menempatkan bahasa Indonesia secara tepat sebagai bagian dari identitas, terutama pada situasi yang membutuhkan bahasa Indonesia yang jelas dan tertib.",
        "Sikap positif; kesetiaan penggunaan; kesadaran norma; kebanggaan bahasa; kemampuan memilih ragam bahasa sesuai konteks.",
        "Kuesioner Likert 20 butir dan pretest-posttest.",
        "Skor pretest dan posttest.",
    ],
    [
        "Kelayakan platform PRIMA+",
        "Tingkat kelayakan produk sebagai media pembelajaran bahasa berdasarkan aspek materi, bahasa, tampilan, navigasi, dan kegunaan.",
        "Kesesuaian materi; kejelasan bahasa; kemudahan penggunaan; tampilan; kualitas umpan balik; kesesuaian fitur dengan tujuan pembelajaran.",
        "Lembar validasi guru/ahli.",
        "Skor validasi skala 1-4.",
    ],
    [
        "Respons pengguna",
        "Tanggapan siswa setelah menggunakan PRIMA+, terutama terkait kemudahan, kemenarikan, manfaat, dan kendala penggunaan.",
        "Kemudahan akses; kemenarikan aktivitas; manfaat terhadap pemahaman bahasa; kendala teknis; saran perbaikan.",
        "Angket respons dan refleksi/wawancara singkat.",
        "Persentase dan ringkasan tema.",
    ],
]

GRID_ROWS = [
    [
        "Loyalitas berbahasa Indonesia",
        "Sikap positif",
        "4",
        "Bahasa Indonesia tetap penting digunakan di media digital.",
        "Kuesioner Likert",
    ],
    [
        "Loyalitas berbahasa Indonesia",
        "Kesetiaan penggunaan",
        "4",
        "Saya memilih bahasa Indonesia yang jelas saat menulis informasi sekolah.",
        "Kuesioner Likert",
    ],
    [
        "Loyalitas berbahasa Indonesia",
        "Kesadaran norma",
        "4",
        "Saya dapat membedakan bahasa santai dan bahasa formal.",
        "Kuesioner Likert dan tes skenario",
    ],
    [
        "Loyalitas berbahasa Indonesia",
        "Kebanggaan bahasa",
        "4",
        "Saya bangga menggunakan bahasa Indonesia dengan baik.",
        "Kuesioner Likert",
    ],
    [
        "Loyalitas berbahasa Indonesia",
        "Pemilihan ragam bahasa",
        "4",
        "Saya menyesuaikan bahasa dengan lawan bicara dan tujuan komunikasi.",
        "Kuesioner Likert, kuis PRIMA+, dan refleksi",
    ],
]


def clear_table(table):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)


def fill_table(table, headers, rows):
    # Reuse the existing column count. This script expects the selected tables
    # to already have the intended number of columns after creation/rebuild.
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    clear_table(table)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            row.cells[i].text = value


def insert_paragraph_after(paragraph, text):
    new_p = paragraph._p.addnext(paragraph._p.__class__())
    # The above keeps python-docx from creating a detached paragraph object.
    # Re-open via document after save is not necessary for simple text insert,
    # so use a lower-level fallback by adding a sibling through OXML.


def add_explanation_after_heading(doc, heading_text):
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == heading_text:
            # If an explanation already exists, replace it; otherwise insert before
            # the next table by placing text in the following blank/normal paragraph
            # when available.
            explanation = (
                "Pada penelitian ini, konstruk dipahami sebagai konsep utama yang diukur "
                "atau dinilai. Konstruk utama yang diuji adalah loyalitas berbahasa Indonesia, "
                "sedangkan kesadaran berbahasa menjadi konstruk proses yang dilatih melalui "
                "PRIMA+. Selain itu, kelayakan platform dan respons pengguna dinilai untuk "
                "memastikan produk yang dikembangkan dapat digunakan dalam uji coba terbatas."
            )
            if idx + 1 < len(doc.paragraphs) and not doc.paragraphs[idx + 1].text.strip():
                doc.paragraphs[idx + 1].text = explanation
            else:
                paragraph.insert_paragraph_before("")
                # Safer fallback: put explanation immediately in the paragraph after
                # the heading if it exists and is not a heading.
                if idx + 1 < len(doc.paragraphs):
                    doc.paragraphs[idx + 1].text = explanation
            return


def replace_table_with_new(doc, old_table_index, headers, rows):
    old = doc.tables[old_table_index]
    parent = old._element.getparent()
    pos = parent.index(old._element)
    parent.remove(old._element)
    new = doc.add_table(rows=1, cols=len(headers))
    new.style = "Table Grid"
    for i, h in enumerate(headers):
        new.rows[0].cells[i].text = h
    for row_values in rows:
        row = new.add_row()
        for i, value in enumerate(row_values):
            row.cells[i].text = value
    parent.insert(pos, new._element)


def update_doc(path):
    doc = Document(str(path))

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "3.5 Variabel, Indikator, dan Instrumen":
            paragraph.text = "3.5 Konstruk, Indikator, dan Instrumen Penelitian"
        elif text == "LAMPIRAN A. KISI-KISI INSTRUMEN LOYALITAS BAHASA":
            paragraph.text = "LAMPIRAN A. KISI-KISI INSTRUMEN KONSTRUK LOYALITAS BERBAHASA INDONESIA"

    add_explanation_after_heading(doc, "3.5 Konstruk, Indikator, dan Instrumen Penelitian")

    # Table 4 in the final proposal is the old variable/instrument table.
    replace_table_with_new(
        doc,
        3,
        ["Konstruk", "Definisi Konseptual", "Dimensi/Indikator", "Instrumen", "Data"],
        CONSTRUCT_ROWS,
    )

    # Table 10 is the appendix grid.
    replace_table_with_new(
        doc,
        9,
        ["Konstruk", "Indikator", "Jumlah Butir", "Contoh Pernyataan", "Instrumen/Sumber Data"],
        GRID_ROWS,
    )

    doc.save(str(path))


def update_revision_table():
    path = Path("Tabel_Daftar_Revisi_Proposal_PRIMA.docx")
    if not path.exists():
        return
    doc = Document(str(path))
    table = doc.tables[0]
    found = False
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if cells and cells[1] == "Instrumen":
            row.cells[2].text = "Bagian instrumen sebelumnya belum menampilkan konstruk secara eksplisit."
            row.cells[3].text = "Ditambahkan konstruk kesadaran berbahasa, loyalitas berbahasa Indonesia, kelayakan platform PRIMA+, dan respons pengguna."
            row.cells[4].text = "BAB 3.5 dan Lampiran A."
            found = True
            break
    if not found:
        row = table.add_row()
        values = [
            "14",
            "Konstruk dan instrumen",
            "Bagian 3.5 sebelumnya langsung menampilkan variabel/aspek, indikator, dan instrumen, tetapi belum menjelaskan konstruk yang diukur.",
            "Bagian 3.5 diubah menjadi 'Konstruk, Indikator, dan Instrumen Penelitian'. Tabelnya kini memuat konstruk kesadaran berbahasa, loyalitas berbahasa Indonesia, kelayakan platform PRIMA+, dan respons pengguna. Lampiran A juga diselaraskan menjadi kisi-kisi instrumen konstruk loyalitas berbahasa Indonesia.",
            "BAB 3.5 dan Lampiran A.",
        ]
        for i, value in enumerate(values):
            row.cells[i].text = value
    doc.save(str(path))


if __name__ == "__main__":
    for file in FILES:
        if file.exists():
            update_doc(file)
            print(f"updated {file}")
    update_revision_table()
    print("updated revision table")
