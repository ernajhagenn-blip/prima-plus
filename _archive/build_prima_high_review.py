from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DOCX = "prima_revisi_HIGH_REVIEW.docx"
OUT_MATRIX = "tabel_review_fundamental_HIGH.docx"
FLOWCHART = "prima_prosedur_ADDIE_HIGH.png"


REFERENCES = [
    "Al-Dosakee, K., & Ozdamli, F. (2021). Gamification in teaching and learning languages: A systematic literature review. Revista Romaneasca pentru Educatie Multidimensionala, 13(2), 559-577. https://doi.org/10.18662/rrem/13.2/436",
    "Alwi, H., & Sugono, D. (2011). Politik Bahasa Nasional. Badan Pengembangan dan Pembinaan Bahasa.",
    "Amelia, S. R., Siroj, M. B., & Afgani, M. W. (2024). Analisis perkembangan bahasa Indonesia di era digital: Tantangan dan peluang. Jurnal Multidisiplin Dehasen, 3(3), 125-132. https://doi.org/10.37676/mude.v3i3.5792",
    "Balai Bahasa Provinsi Jambi. (2025). UKBI jadi instrumen strategis wujudkan kemahiran berbahasa Indonesia yang terstandar. Kementerian Pendidikan Dasar dan Menengah.",
    "Branch, R. M. (2009). Instructional Design: The ADDIE Approach. Springer. https://doi.org/10.1007/978-0-387-09506-6",
    "Carter, R. (2003). Language awareness. ELT Journal, 57(1), 64-65. https://doi.org/10.1093/elt/57.1.64",
    "Chaer, A., & Agustina, L. (2014). Sosiolinguistik: Perkenalan awal. Rineka Cipta.",
    "Crystal, D. (2011). Internet Linguistics: A Student Guide. Routledge.",
    "Fairclough, N. (1992). Critical Language Awareness. Longman.",
    "Garvin, P. L., & Mathiot, M. (1968). The urbanization of the Guarani language: A problem in language and culture. In J. A. Fishman (Ed.), Readings in the Sociology of Language (pp. 365-374). De Gruyter Mouton. https://doi.org/10.1515/9783110805376.365",
    "Luo, Z. (2023). The effectiveness of gamified tools for foreign language learning (FLL): A systematic review. Behavioral Sciences, 13(4), 331. https://doi.org/10.3390/bs13040331",
    "Prensky, M. (2001). Digital natives, digital immigrants. On the Horizon, 9(5), 1-6.",
    "Purba, E. N., Togatorop, D. P., Simbolon, A., & Sari, Y. (2024). Analisis pengaruh media sosial terhadap keberagaman bahasa: Campur kode sebagai tren komunikasi anak muda. Atmosfer: Jurnal Pendidikan, Bahasa, Sastra, Seni, Budaya, dan Sosial Humaniora, 2(4), 184-194. https://doi.org/10.59024/atmosfer.v2i4.1060",
    "Shortt, M., Tilak, S., Kuznetcova, I., Martens, B., & Akinkuolie, B. (2021). Gamification in mobile-assisted language learning: A systematic review of Duolingo literature from public release of 2012 to early 2020. Computer Assisted Language Learning. https://doi.org/10.1080/09588221.2021.1933540",
    "Sugiyono. (2019). Metode Penelitian dan Pengembangan: Research and Development. Alfabeta.",
    "Sugono, D. (2019). Bahasa Indonesia di era globalisasi. Badan Pengembangan dan Pembinaan Bahasa.",
]


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def clear_cell(cell):
    cell.text = ""


def write_cell(cell, text, bold=False, size=9.2, align=None):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROPOSAL PENELITIAN")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "PENGEMBANGAN PLATFORM PRIMA+ BERBASIS KESADARAN BERBAHASA "
        "UNTUK MENGUATKAN LOYALITAS BAHASA INDONESIA REMAJA"
    )
    r.bold = True
    r.font.size = Pt(13)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for line in [
        "Ketua: Larisa Dianti",
        "Anggota: Gusti Nazwa Azizah",
        "Bidang Ilmu Sosial Humaniora (ISH), Subbidang Bahasa dan Sastra",
        "MAN Kotawaringin Timur, Kalimantan Tengah",
        "Tahun 2026",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def para(doc, text):
    return doc.add_paragraph(text)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def table(doc, headers, rows, widths=None, font_size=9.2):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        write_cell(t.rows[0].cells[i], h, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(t.rows[0].cells[i], "F4F6F9")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            write_cell(cells[i], val, size=font_size)
    if widths:
        for row in t.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return t


def make_flowchart():
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return None

    width, height = 1500, 1800
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arialbd.ttf", 42)
        font_box = ImageFont.truetype("arial.ttf", 30)
        font_small = ImageFont.truetype("arial.ttf", 24)
    except Exception:
        font_title = font_box = font_small = ImageFont.load_default()

    draw.text((width // 2, 55), "Alur Penelitian PRIMA+ Berbasis ADDIE", anchor="mm", font=font_title, fill=(20, 55, 85))
    boxes = [
        ("ANALYZE", "Analisis masalah, kebutuhan siswa,\nindikator loyalitas, dan contoh bahasa digital"),
        ("DESIGN", "Merancang fitur PRIMA+, kisi-kisi instrumen,\nskenario kuis, dan rubrik validasi"),
        ("DEVELOP", "Membangun prototipe, aset visual/audio,\nkonten kuis, dan umpan balik"),
        ("VALIDATE", "Validasi ahli/guru: materi, bahasa,\nmedia, dan kelayakan instrumen"),
        ("IMPLEMENT", "Pretest -> penggunaan PRIMA+ -> posttest\npada 30-40 responden uji coba terbatas"),
        ("EVALUATE", "Analisis skor, N-gain, uji beda,\nrefleksi pengguna, dan revisi akhir"),
    ]
    x1, x2 = 170, 1330
    y = 150
    box_h = 185
    gap = 70
    for idx, (title, body) in enumerate(boxes):
        y1 = y + idx * (box_h + gap)
        y2 = y1 + box_h
        fill = (235, 242, 249) if idx % 2 == 0 else (246, 248, 250)
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill=fill, outline=(60, 110, 155), width=4)
        draw.text((x1 + 40, y1 + 35), title, font=font_title, fill=(31, 77, 120))
        draw.multiline_text((x1 + 40, y1 + 92), body, font=font_box, fill=(35, 35, 35), spacing=7)
        if idx < len(boxes) - 1:
            ax = width // 2
            draw.line((ax, y2 + 10, ax, y2 + gap - 10), fill=(60, 110, 155), width=5)
            draw.polygon([(ax, y2 + gap - 10), (ax - 15, y2 + gap - 38), (ax + 15, y2 + gap - 38)], fill=(60, 110, 155))
    draw.text((width // 2, height - 80), "Output: platform PRIMA+ tervalidasi, data efektivitas, dan rekomendasi revisi produk", anchor="mm", font=font_small, fill=(60, 60, 60))
    img.save(FLOWCHART)
    return FLOWCHART


def build_proposal():
    make_flowchart()
    doc = Document()
    style_doc(doc)
    add_title(doc)
    doc.add_page_break()

    doc.add_heading("BAB 1. PENDAHULUAN", level=1)
    doc.add_heading("1.1 Latar Belakang", level=2)
    for text in [
        "Bahasa Indonesia memiliki fungsi strategis sebagai bahasa nasional, bahasa persatuan, bahasa pendidikan, dan penanda identitas kebangsaan. Pada kalangan remaja, bahasa Indonesia tidak hanya dipakai untuk menyampaikan pesan, tetapi juga untuk membentuk identitas sosial, menunjukkan kedekatan kelompok, dan berpartisipasi dalam ruang digital. Karena itu, penguatan loyalitas bahasa Indonesia perlu diarahkan bukan hanya pada hafalan kaidah, tetapi juga pada kesadaran memilih ragam bahasa sesuai konteks.",
        "Perkembangan media sosial mengubah pola komunikasi remaja. Ungkapan singkat, slang, emotikon, campur kode Indonesia-Inggris, dan istilah populer seperti happy, healing, deadline, vibes, atau relate sering muncul dalam percakapan digital. Fenomena ini tidak boleh disederhanakan sebagai kesalahan semata karena variasi bahasa merupakan gejala sosiolinguistik yang wajar. Namun, masalah penelitian muncul ketika remaja tidak memiliki kesadaran untuk membedakan kapan bahasa santai, campur kode, bahasa baku, dan bahasa Indonesia yang sesuai norma perlu digunakan.",
        "Kebutuhan penguatan bahasa Indonesia pada pelajar juga memiliki dasar empiris. Balai Bahasa Provinsi Jambi (2025) melaporkan pemanfaatan UKBI sebagai instrumen strategis untuk mewujudkan kemahiran berbahasa Indonesia yang terstandar, serta menyoroti bahwa data pelajar tahun 2024 masih menunjukkan kebutuhan penguatan kemahiran berbahasa. Temuan ini memperlihatkan bahwa pembinaan bahasa Indonesia pada kelompok pelajar perlu terus dikembangkan melalui strategi yang dekat dengan kehidupan digital mereka.",
        "Pendekatan kesadaran berbahasa atau language awareness relevan karena menekankan kepekaan terhadap bentuk, fungsi, konteks, dan dampak sosial penggunaan bahasa (Carter, 2003; Fairclough, 1992). Dengan pendekatan ini, remaja tidak sekadar diberi larangan memakai campur kode atau bahasa gaul, tetapi diajak memahami fungsi bahasa, menilai konteks, dan mengambil keputusan kebahasaan secara sadar. Pendekatan tersebut lebih adil secara sosiolinguistik dan lebih dekat dengan realitas komunikasi remaja.",
        "Berdasarkan persoalan tersebut, penelitian ini mengembangkan PRIMA+ sebagai platform digital berbasis kesadaran berbahasa. PRIMA+ dirancang sebagai media interaktif yang memadukan contoh kasus bahasa digital, kuis kontekstual, refleksi pilihan bahasa, sistem skor, dan umpan balik. Penelitian ini tidak hanya bertujuan membuat produk, tetapi juga menguji apakah penggunaan PRIMA+ dapat menguatkan loyalitas bahasa Indonesia remaja secara terukur.",
    ]:
        para(doc, text)

    doc.add_heading("1.2 Fokus Masalah dan Gap Penelitian", level=2)
    para(
        doc,
        "Gap utama penelitian ini adalah belum kuatnya model penguatan loyalitas bahasa Indonesia remaja yang menghubungkan pendekatan kesadaran berbahasa dengan media digital interaktif. Kajian tentang campur kode dan bahasa digital umumnya menjelaskan fenomena, sedangkan kajian gamifikasi bahasa banyak berfokus pada pembelajaran bahasa asing. PRIMA+ ditempatkan sebagai upaya mengisi celah tersebut melalui platform yang menumbuhkan kesadaran konteks, kebanggaan, dan kesetiaan penggunaan bahasa Indonesia.",
    )

    doc.add_heading("1.3 Rumusan Masalah", level=2)
    numbers(
        doc,
        [
            "Bagaimana kebutuhan dan rancangan platform PRIMA+ berbasis kesadaran berbahasa untuk menguatkan loyalitas bahasa Indonesia remaja?",
            "Bagaimana kelayakan materi, bahasa, media, dan instrumen PRIMA+ berdasarkan validasi guru atau ahli?",
            "Bagaimana efektivitas penggunaan PRIMA+ terhadap peningkatan loyalitas bahasa Indonesia remaja berdasarkan indikator sikap positif, kesetiaan penggunaan, kesadaran norma, kebanggaan bahasa, dan kemampuan memilih ragam bahasa sesuai konteks?",
        ],
    )

    doc.add_heading("1.4 Tujuan Penelitian", level=2)
    numbers(
        doc,
        [
            "Mengembangkan platform PRIMA+ berbasis kesadaran berbahasa sesuai kebutuhan remaja pengguna media digital.",
            "Menilai kelayakan PRIMA+ dari aspek materi, bahasa, media, dan instrumen penelitian.",
            "Menguji efektivitas PRIMA+ dalam menguatkan loyalitas bahasa Indonesia remaja melalui desain pretest-posttest.",
        ],
    )

    doc.add_heading("1.5 Manfaat Penelitian", level=2)
    doc.add_heading("1.5.1 Manfaat Teoretis", level=3)
    para(doc, "Penelitian ini memperkaya kajian sosiolinguistik terapan dan pembelajaran bahasa Indonesia, terutama pada pengembangan media digital berbasis kesadaran berbahasa untuk menguatkan sikap dan loyalitas bahasa.")
    doc.add_heading("1.5.2 Manfaat Praktis", level=3)
    bullets(
        doc,
        [
            "Bagi siswa, PRIMA+ membantu memahami penggunaan bahasa Indonesia secara sadar, kontekstual, dan tetap relevan dengan ruang digital.",
            "Bagi guru Bahasa Indonesia, PRIMA+ dapat menjadi contoh media pembelajaran sikap bahasa yang tidak hanya berisi materi kaidah, tetapi juga latihan pengambilan keputusan berbahasa.",
            "Bagi sekolah, penelitian ini dapat mendukung program literasi, karakter kebahasaan, dan penguatan identitas nasional.",
            "Bagi pengembang media pendidikan, penelitian ini memberi rancangan awal produk digital yang dapat dikembangkan lebih lanjut.",
        ],
    )

    doc.add_heading("1.6 Batasan Penelitian", level=2)
    bullets(
        doc,
        [
            "Penelitian dilakukan pada uji coba terbatas di MAN Kotawaringin Timur.",
            "Produk PRIMA+ difokuskan pada prototipe pembelajaran interaktif, bukan aplikasi komersial penuh.",
            "Loyalitas bahasa Indonesia diukur melalui indikator sikap dan perilaku berbahasa yang dapat diamati melalui kuesioner, tes skenario, dan aktivitas platform.",
            "Data empiris lokal tidak diasumsikan sebelum pengambilan data; data lokal akan diperoleh melalui survei kebutuhan, pretest, posttest, dan refleksi pengguna.",
        ],
    )

    doc.add_heading("BAB 2. TINJAUAN PUSTAKA", level=1)
    doc.add_heading("2.1 Bahasa Indonesia, Remaja, dan Komunikasi Digital", level=2)
    para(doc, "Ruang digital membuat komunikasi remaja berlangsung cepat, multimodal, dan penuh variasi. Kajian internet linguistics menjelaskan bahwa teknologi digital dapat mengubah praktik bahasa, bentuk interaksi, dan norma komunikasi (Crystal, 2011). Dalam konteks Indonesia, penggunaan bahasa gaul dan campur kode di media sosial perlu dipahami sebagai gejala sosial sekaligus tantangan pembinaan bahasa.")
    doc.add_heading("2.2 Kesadaran Berbahasa", level=2)
    para(doc, "Kesadaran berbahasa menekankan kepekaan terhadap struktur, fungsi, konteks, dan konsekuensi sosial bahasa. Carter (2003) menjelaskan language awareness sebagai peningkatan kesadaran terhadap bentuk dan fungsi bahasa, sedangkan Fairclough (1992) menekankan dimensi kritis bahasa dalam relasi sosial. Pada penelitian ini, kesadaran berbahasa diarahkan pada kemampuan siswa menilai pilihan bahasa dalam konteks formal, informal, sekolah, dan media sosial.")
    doc.add_heading("2.3 Loyalitas Bahasa Indonesia", level=2)
    para(doc, "Loyalitas bahasa merupakan sikap mempertahankan dan menggunakan bahasa sebagai bagian dari identitas sosial. Garvin dan Mathiot (1968) menempatkan loyalitas bahasa sebagai unsur sikap bahasa, bersama kebanggaan bahasa dan kesadaran norma. Dalam konteks penelitian ini, loyalitas bahasa Indonesia tidak dimaknai sebagai penolakan total terhadap bahasa asing atau bahasa gaul, melainkan sebagai kemampuan dan kemauan menjaga fungsi bahasa Indonesia secara tepat sesuai konteks.")

    table(
        doc,
        ["Indikator", "Definisi Operasional", "Sumber Data", "Contoh Butir/Perilaku"],
        [
            ["Sikap positif", "Penilaian bahwa bahasa Indonesia penting dan bernilai", "Kuesioner Likert", "Saya merasa bahasa Indonesia tetap penting di ruang digital."],
            ["Kesetiaan penggunaan", "Kecenderungan memilih bahasa Indonesia pada konteks sekolah dan komunikasi resmi", "Skenario pilihan bahasa", "Memilih ragam Indonesia yang tepat untuk pengumuman kelas."],
            ["Kesadaran norma", "Pemahaman perbedaan ragam baku, santai, slang, dan campur kode", "Tes skenario/kuis", "Mengenali konteks saat campur kode kurang tepat."],
            ["Kebanggaan bahasa", "Rasa bangga dan percaya diri menggunakan bahasa Indonesia", "Kuesioner Likert", "Tidak menganggap bahasa Indonesia kalah modern."],
            ["Pemilihan ragam", "Kemampuan menyesuaikan bahasa dengan tujuan, audiens, dan media", "Kuis PRIMA+ dan refleksi", "Menulis ulang caption agar sesuai konteks formal."],
        ],
        widths=[3.0, 4.7, 3.1, 5.2],
        font_size=8.7,
    )

    doc.add_heading("2.4 Platform Digital, Gamifikasi, dan Pembelajaran Bahasa", level=2)
    para(doc, "Gamifikasi menggunakan elemen permainan seperti poin, tantangan, misi, umpan balik, dan capaian untuk meningkatkan keterlibatan belajar. Tinjauan sistematis menunjukkan bahwa gamifikasi dapat mendukung motivasi belajar bahasa, tetapi efektivitasnya bergantung pada desain pedagogis, kualitas umpan balik, dan cara pengukuran hasil belajar (Al-Dosakee & Ozdamli, 2021; Luo, 2023; Shortt et al., 2021). Karena itu, PRIMA+ tidak cukup hanya dibuat menarik, tetapi harus memiliki indikator, instrumen, dan mekanisme evaluasi yang jelas.")

    doc.add_heading("2.5 Posisi Penelitian", level=2)
    table(
        doc,
        ["Klaster Kajian", "Temuan Utama", "Keterbatasan yang Disasar", "Posisi PRIMA+"],
        [
            ["Bahasa digital dan campur kode", "Media sosial memunculkan variasi bahasa, campur kode, dan gaya komunikasi baru.", "Sering berhenti pada deskripsi fenomena.", "Mengubah fenomena menjadi materi refleksi dan kuis kontekstual."],
            ["Kesadaran berbahasa", "Kesadaran bahasa membantu siswa memahami fungsi dan konteks bahasa.", "Belum selalu diwujudkan dalam media digital interaktif.", "Menjadikan language awareness sebagai prinsip desain platform."],
            ["Gamifikasi bahasa", "Gamifikasi dapat meningkatkan motivasi, tetapi hasilnya bergantung pada desain.", "Banyak kajian berfokus pada bahasa asing.", "Menerapkan gamifikasi untuk penguatan loyalitas bahasa Indonesia."],
            ["R&D media pembelajaran", "ADDIE memberi tahapan pengembangan sistematis.", "Sering kurang kuat dalam uji efektivitas.", "Menggabungkan ADDIE dengan pretest-posttest dan validasi ahli."],
        ],
        widths=[3.5, 4.2, 4.2, 4.1],
        font_size=8.7,
    )

    doc.add_heading("BAB 3. METODE PENELITIAN", level=1)
    doc.add_heading("3.1 Desain Penelitian", level=2)
    para(doc, "Penelitian ini menggunakan metode Research and Development (R&D) dengan model ADDIE yang terdiri atas Analyze, Design, Develop, Implement, dan Evaluate (Branch, 2009). Untuk menguji efektivitas produk, penelitian menggunakan desain one-group pretest-posttest. Desain ini dipilih karena penelitian berfokus pada pengembangan prototipe dan uji coba terbatas, bukan perbandingan antarkelas secara luas.")

    doc.add_heading("3.2 Waktu dan Tempat Penelitian", level=2)
    para(doc, "Penelitian dilaksanakan pada April-Agustus 2026 di MAN Kotawaringin Timur, Kalimantan Tengah. Tahapan penelitian meliputi analisis kebutuhan, perancangan, pengembangan prototipe, validasi, uji coba terbatas, analisis data, revisi produk, dan penyusunan laporan.")

    doc.add_heading("3.3 Subjek Penelitian dan Sampling", level=2)
    para(doc, "Subjek penelitian adalah siswa MAN Kotawaringin Timur kelas X atau XI yang aktif menggunakan media digital. Uji coba terbatas direncanakan melibatkan 30-40 siswa agar data pretest-posttest lebih stabil daripada uji coba sangat kecil. Teknik sampling menggunakan purposive sampling dengan kriteria: siswa aktif, bersedia mengikuti rangkaian penelitian, memiliki akses perangkat untuk mencoba platform, dan memperoleh izin sesuai ketentuan sekolah.")

    doc.add_heading("3.4 Spesifikasi Produk PRIMA+", level=2)
    table(
        doc,
        ["Komponen Produk", "Isi/Fungsi", "Keluaran yang Diukur"],
        [
            ["Modul kasus bahasa digital", "Contoh caption, chat, story, dan situasi sekolah yang memuat pilihan ragam bahasa.", "Kemampuan mengenali konteks bahasa."],
            ["Kuis kontekstual", "Soal pilihan dan skenario tentang bahasa baku, santai, slang, dan campur kode.", "Skor pemahaman dan kesadaran norma."],
            ["Refleksi bahasa", "Pertanyaan singkat mengapa pilihan bahasa tertentu tepat/kurang tepat.", "Alasan siswa dalam memilih ragam bahasa."],
            ["Umpan balik otomatis", "Penjelasan setelah siswa menjawab kuis.", "Perbaikan pemahaman setelah latihan."],
            ["Skor dan capaian", "Poin, level, atau badge sederhana.", "Keterlibatan dan penyelesaian aktivitas."],
        ],
        widths=[3.5, 7.0, 5.5],
        font_size=8.8,
    )

    doc.add_heading("3.5 Variabel, Indikator, dan Instrumen", level=2)
    table(
        doc,
        ["Variabel/Aspek", "Indikator", "Instrumen", "Skala/Data"],
        [
            ["Kebutuhan pengguna", "Kebiasaan berbahasa digital, kebutuhan media, preferensi fitur", "Angket analisis kebutuhan", "Deskriptif persentase"],
            ["Kelayakan produk", "Materi, bahasa, media, tampilan, navigasi", "Lembar validasi guru/ahli", "Skala 1-4"],
            ["Loyalitas bahasa Indonesia", "Sikap positif, kesetiaan, kesadaran norma, kebanggaan, pemilihan ragam", "Kuesioner Likert 20 butir", "Skor pretest-posttest"],
            ["Pemahaman bahasa kontekstual", "Ketepatan memilih ragam bahasa sesuai konteks", "Tes skenario/kuis PRIMA+", "Skor benar-salah"],
            ["Respons pengguna", "Kemenarikan, kemudahan, manfaat, kendala", "Refleksi/wawancara singkat", "Data kualitatif ringkas"],
        ],
        widths=[3.4, 5.3, 4.0, 3.3],
        font_size=8.7,
    )

    doc.add_heading("3.6 Alat dan Bahan Penelitian", level=2)
    table(
        doc,
        ["Kategori", "Nama", "Jumlah", "Fungsi"],
        [
            ["Perangkat", "Laptop pengembang", "1 unit", "Mendesain platform, mengolah data, dan menyusun laporan."],
            ["Perangkat", "Telepon pintar/laptop siswa", "30-40 perangkat milik responden", "Menguji akses dan penggunaan PRIMA+."],
            ["Perangkat lunak", "Godot Engine atau platform web sederhana", "1 sistem", "Membangun prototipe interaktif PRIMA+."],
            ["Perangkat lunak", "Spreadsheet/statistik sederhana", "1 paket", "Mengolah skor pretest-posttest, N-gain, dan grafik."],
            ["Bahan", "Skenario kasus bahasa digital", "1 paket", "Materi utama kuis dan refleksi bahasa."],
            ["Bahan", "Kuesioner loyalitas bahasa", "20 butir", "Mengukur indikator loyalitas bahasa Indonesia."],
            ["Bahan", "Lembar validasi ahli/guru", "2-3 validator", "Menilai kelayakan materi, bahasa, dan media."],
            ["Bahan", "Aset visual/audio", "1 paket", "Mendukung tampilan dan pengalaman belajar."],
        ],
        widths=[2.8, 4.6, 3.1, 5.5],
        font_size=8.7,
    )

    doc.add_heading("3.7 Prosedur Penelitian", level=2)
    para(doc, "Prosedur penelitian direvisi menggunakan model ADDIE agar hubungan antara pengembangan produk, validasi, uji coba, dan evaluasi menjadi lebih jelas.")
    if Path(FLOWCHART).exists():
        doc.add_picture(FLOWCHART, width=Cm(15.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph("Gambar 1. Prosedur penelitian PRIMA+ berbasis ADDIE")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table(
        doc,
        ["Tahap ADDIE", "Aktivitas", "Luaran"],
        [
            ["Analyze", "Analisis kebutuhan, observasi contoh bahasa digital, kajian literatur, dan pemetaan indikator loyalitas.", "Profil kebutuhan, indikator penelitian, dan bahan awal kasus bahasa."],
            ["Design", "Merancang fitur PRIMA+, alur pengguna, kisi-kisi instrumen, skenario kuis, dan rubrik validasi.", "Blueprint platform, kisi-kisi, dan instrumen awal."],
            ["Develop", "Membangun prototipe, menyusun konten kuis/refleksi, menyiapkan aset visual/audio, dan uji teknis internal.", "Prototipe PRIMA+ versi uji."],
            ["Validate", "Validasi oleh guru/ahli terhadap materi, bahasa, media, dan instrumen.", "Skor kelayakan dan daftar revisi."],
            ["Implement", "Pretest, penggunaan PRIMA+, posttest, dan pengumpulan respons pengguna.", "Data skor, log aktivitas, dan respons siswa."],
            ["Evaluate", "Analisis efektivitas, interpretasi data, revisi akhir produk, dan penyusunan laporan.", "Produk akhir, hasil penelitian, dan rekomendasi pengembangan."],
        ],
        widths=[2.6, 8.0, 5.4],
        font_size=8.8,
    )

    doc.add_heading("3.8 Teknik Pengumpulan Data", level=2)
    numbers(
        doc,
        [
            "Angket analisis kebutuhan untuk mengetahui kebiasaan siswa menggunakan bahasa di ruang digital.",
            "Validasi ahli/guru terhadap materi, bahasa, media, dan instrumen.",
            "Pretest loyalitas bahasa Indonesia dan tes pemahaman penggunaan bahasa kontekstual.",
            "Uji coba PRIMA+ dan pencatatan skor aktivitas platform.",
            "Posttest setelah penggunaan PRIMA+.",
            "Refleksi singkat atau wawancara terbatas untuk mengetahui pengalaman siswa.",
        ],
    )

    doc.add_heading("3.9 Teknik Analisis Data", level=2)
    para(doc, "Data dianalisis secara kuantitatif dan kualitatif sederhana. Data validasi ahli dihitung dengan persentase kelayakan. Data pretest-posttest dianalisis melalui rata-rata, persentase peningkatan, dan N-gain. Jika data memenuhi syarat, uji paired sample t-test digunakan; jika tidak, digunakan uji Wilcoxon. Data refleksi siswa dianalisis secara tematik ringkas untuk menjelaskan kendala dan pengalaman pengguna.")
    table(
        doc,
        ["Data", "Analisis", "Kriteria Interpretasi"],
        [
            ["Validasi ahli/guru", "Persentase kelayakan", "Sangat layak, layak, cukup, atau perlu revisi."],
            ["Pretest-posttest loyalitas", "Rata-rata, persentase peningkatan, N-gain", "Ada peningkatan skor setelah penggunaan PRIMA+."],
            ["Tes skenario bahasa", "Skor benar-salah dan persentase ketepatan", "Pemahaman konteks bahasa meningkat."],
            ["Uji beda", "Paired t-test atau Wilcoxon", "Perbedaan dinilai bermakna jika p < 0,05."],
            ["Refleksi pengguna", "Koding tema ringkas", "Menjelaskan kelebihan, kendala, dan saran revisi."],
        ],
        widths=[3.6, 5.5, 6.9],
        font_size=8.8,
    )

    doc.add_heading("3.10 Kriteria Keberhasilan", level=2)
    bullets(
        doc,
        [
            "Produk memperoleh kategori minimal layak dari validator.",
            "Rata-rata skor posttest lebih tinggi daripada pretest.",
            "N-gain minimal berada pada kategori sedang atau menunjukkan peningkatan yang dapat dijelaskan secara akademik.",
            "Sebagian besar responden menyatakan PRIMA+ mudah digunakan dan membantu memahami pilihan bahasa sesuai konteks.",
        ],
    )

    doc.add_heading("BAB 4. RANCANGAN ANGGARAN BIAYA DAN JADWAL", level=1)
    doc.add_heading("4.1 Rancangan Anggaran Biaya", level=2)
    table(
        doc,
        ["No.", "Komponen", "Jumlah", "Subtotal", "Keterangan"],
        [
            ["1", "Pengembangan prototipe PRIMA+", "1 paket", "Rp2.500.000", "Desain alur, fitur kuis, skor, umpan balik, dan integrasi konten."],
            ["2", "Aset visual antarmuka/game", "1 paket", "Rp1.200.000", "Karakter, ikon, map sederhana, tombol, dan elemen UI."],
            ["3", "Aset audio permainan", "1 paket", "Rp600.000", "Efek suara, musik latar, dan audio pendukung."],
            ["4", "Perangkat pendukung uji coba", "1 paket", "Rp1.200.000", "Headset, penyimpanan, kabel, dan kebutuhan teknis lapangan."],
            ["5", "Internet, hosting/domain uji coba", "1 paket", "Rp1.000.000", "Akses referensi, pengembangan, unggah prototipe, dan uji akses."],
            ["6", "Cetak instrumen penelitian", "1 paket", "Rp700.000", "Kuesioner, lembar validasi, pretest-posttest, dan formulir observasi."],
            ["7", "Dokumentasi dan bahan presentasi", "1 paket", "Rp800.000", "Dokumentasi kegiatan, poster, slide, dan laporan visual."],
            ["8", "Operasional uji coba", "1 paket", "Rp1.200.000", "Koordinasi responden, konsumsi terbatas, dan kebutuhan kegiatan."],
            ["9", "Validasi ahli/guru dan revisi materi", "1 paket", "Rp800.000", "Honorarium/transport terbatas validator dan perbaikan konten."],
            ["", "Total", "", "Rp10.000.000", ""],
        ],
        widths=[1.0, 5.0, 2.3, 2.8, 4.9],
        font_size=8.5,
    )
    para(doc, "Total anggaran Rp10.000.000 dipilih agar realistis untuk pengembangan prototipe dan uji coba, tetapi tetap berada jauh di bawah batas maksimal Rp15.000.000. Godot Engine atau perangkat lunak open-source tidak dihitung sebagai biaya lisensi; biaya diarahkan pada pembuatan prototipe, aset, instrumen, uji coba, validasi, dan dokumentasi.")

    doc.add_heading("4.2 Jadwal Penelitian", level=2)
    table(
        doc,
        ["No.", "Tahap", "Waktu", "Luaran"],
        [
            ["1", "Analyze: studi literatur dan analisis kebutuhan", "6-16 April 2026", "Data kebutuhan dan indikator penelitian."],
            ["2", "Design: rancangan platform dan instrumen", "17-30 April 2026", "Blueprint PRIMA+, kisi-kisi, dan instrumen."],
            ["3", "Develop: pembuatan prototipe", "1-25 Mei 2026", "Prototipe PRIMA+ versi uji."],
            ["4", "Validate: validasi dan revisi awal", "26 Mei-5 Juni 2026", "Skor kelayakan dan daftar revisi."],
            ["5", "Implement: pretest, uji coba, posttest", "6-30 Juni 2026", "Data efektivitas dan respons siswa."],
            ["6", "Evaluate: analisis dan revisi akhir", "1-25 Juli 2026", "Produk revisi dan hasil analisis."],
            ["7", "Penyusunan laporan akhir", "26 Juli-20 Agustus 2026", "Laporan penelitian final."],
        ],
        widths=[1.0, 5.5, 4.0, 5.5],
        font_size=8.7,
    )

    doc.add_heading("PERNYATAAN PENGGUNAAN ARTIFICIAL INTELLIGENCE (AI)", level=1)
    para(doc, "Penelitian ini menggunakan bantuan Artificial Intelligence (AI), yaitu Gemini dan ChatGPT, untuk mencari referensi awal, memperdalam teori, merapikan metodologi, dan menyusun alternatif alur permainan atau kuis bahasa. Seluruh keputusan penelitian, penyusunan instrumen, pengembangan produk, pengambilan data, analisis data, dan penarikan kesimpulan tetap dilakukan secara mandiri oleh peneliti. AI tidak digunakan untuk membuat data palsu, menggantikan uji coba, atau menulis hasil penelitian tanpa verifikasi.")

    doc.add_heading("DAFTAR PUSTAKA", level=1)
    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading("LAMPIRAN A. KISI-KISI INSTRUMEN LOYALITAS BAHASA", level=1)
    table(
        doc,
        ["Indikator", "Jumlah Butir", "Contoh Pernyataan"],
        [
            ["Sikap positif", "4", "Bahasa Indonesia tetap penting digunakan di media digital."],
            ["Kesetiaan penggunaan", "4", "Saya memilih bahasa Indonesia yang jelas saat menulis informasi sekolah."],
            ["Kesadaran norma", "4", "Saya dapat membedakan bahasa santai dan bahasa formal."],
            ["Kebanggaan bahasa", "4", "Saya bangga menggunakan bahasa Indonesia dengan baik."],
            ["Pemilihan ragam", "4", "Saya menyesuaikan bahasa dengan lawan bicara dan tujuan komunikasi."],
        ],
        widths=[4.0, 3.0, 9.0],
        font_size=8.8,
    )

    doc.save(OUT_DOCX)


def build_matrix():
    doc = Document()
    style_doc(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TABEL REVIEW FUNDAMENTAL VERSI HIGH")
    r.bold = True
    r.font.size = Pt(14)
    para(doc, "Tabel ini menjelaskan perombakan mendasar yang dilakukan pada naskah revisi high review.")
    table(
        doc,
        ["No.", "Bagian", "Masalah Sebelumnya", "Perombakan High Review", "Letak di Naskah"],
        [
            ["1", "Judul", "Masih campuran istilah dan kurang tegas sebagai pengembangan produk.", "Judul diganti menjadi pengembangan platform berbasis kesadaran berbahasa untuk menguatkan loyalitas bahasa Indonesia remaja.", "Halaman judul"],
            ["2", "Latar belakang", "Fenomena campur kode berisiko dibaca sebagai kesalahan semata.", "Ditulis ulang lebih hati-hati: campur kode dipahami sebagai gejala sosiolinguistik, masalahnya adalah kurangnya kesadaran konteks.", "BAB 1.1"],
            ["3", "Gap penelitian", "Gap belum dinyatakan eksplisit.", "Ditambahkan bagian fokus masalah dan gap penelitian.", "BAB 1.2"],
            ["4", "Rumusan masalah", "Belum memuat validasi produk.", "Ditambah pertanyaan tentang kelayakan materi, bahasa, media, dan instrumen.", "BAB 1.3"],
            ["5", "Metode", "R&D masih terlalu umum.", "Diganti menjadi R&D model ADDIE + one-group pretest-posttest.", "BAB 3.1"],
            ["6", "Sampel", "Jumlah responden masih 15-30 dan kurang kuat.", "Dinaikkan menjadi 30-40 responden uji coba terbatas dengan purposive sampling.", "BAB 3.3"],
            ["7", "Produk", "Spesifikasi platform belum rinci.", "Ditambahkan komponen produk: kasus digital, kuis, refleksi, umpan balik, skor/capaian.", "BAB 3.4"],
            ["8", "Instrumen", "Alat-bahan bercampur dengan instrumen.", "Instrumen dipisah: angket kebutuhan, validasi, kuesioner, tes skenario, refleksi.", "BAB 3.5-3.6"],
            ["9", "Gambar prosedur", "Gambar lama terlalu linear dan belum menunjukkan metode penelitian.", "Diganti dengan flowchart ADDIE lengkap: Analyze, Design, Develop, Validate, Implement, Evaluate.", "BAB 3.7"],
            ["10", "Analisis data", "Analisis belum cukup kuat.", "Ditambah kelayakan produk, N-gain, paired t-test/Wilcoxon, dan analisis tematik refleksi.", "BAB 3.9"],
            ["11", "Kriteria keberhasilan", "Belum eksplisit.", "Ditambah kriteria minimal kelayakan, peningkatan posttest, N-gain, dan respons pengguna.", "BAB 3.10"],
            ["12", "RAB", "Sudah Rp10 juta tetapi perlu narasi akademik.", "RAB dipertahankan Rp10 juta dengan pembagian lebih rapi dan alasan di bawah plafon Rp15 juta.", "BAB 4.1"],
            ["13", "Daftar pustaka", "Referensi sudah bertambah tetapi perlu selaras dengan ADDIE.", "Ditambahkan Branch (2009) untuk ADDIE dan daftar pustaka diselaraskan dengan teori/metode baru.", "Daftar Pustaka"],
        ],
        widths=[1.0, 3.2, 4.5, 5.0, 2.3],
        font_size=8.3,
    )
    doc.save(OUT_MATRIX)


if __name__ == "__main__":
    build_proposal()
    build_matrix()
    print(OUT_DOCX)
    print(OUT_MATRIX)
    print(FLOWCHART)
