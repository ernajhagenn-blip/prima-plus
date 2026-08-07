from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from pathlib import Path


OUT_PROPOSAL = "prima_revisi.docx"
OUT_TABLE = "tabel_revisi_reviewer.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(9.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    r.font.size = Pt(14)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(11)
        r2.font.name = "Calibri"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "F4F6F9")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def proposal_doc():
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        'PROPOSAL PENELITIAN',
        'PRIMA+: Platform Language Awareness untuk Meningkatkan Loyalitas Berbahasa Indonesia Remaja',
    )
    for text in [
        "Ketua: Larisa Dianti",
        "Anggota: Gusti Nazwa Azizah",
        "Bidang Ilmu Sosial Humaniora (ISH), Subbidang Bahasa dan Sastra",
        "MAN Kotawaringin Timur, Kalimantan Tengah",
        "Tahun 2026",
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("BAB 1. PENDAHULUAN", level=1)
    doc.add_heading("1.1 Latar Belakang", level=2)
    paragraphs = [
        "Bahasa Indonesia memiliki kedudukan penting sebagai bahasa nasional, bahasa persatuan, dan identitas bangsa. Dalam kehidupan remaja, bahasa Indonesia tidak hanya berfungsi sebagai alat komunikasi, tetapi juga menjadi sarana membangun identitas sosial, mengekspresikan pikiran, serta berpartisipasi dalam lingkungan pendidikan dan ruang digital. Oleh karena itu, kemampuan dan kemauan remaja untuk menggunakan bahasa Indonesia secara tepat perlu terus diperkuat.",
        "Perkembangan media sosial dan budaya digital membawa perubahan besar pada cara remaja berkomunikasi. Pada berbagai platform seperti TikTok, Instagram, WhatsApp, dan gim daring, remaja sering menggunakan bahasa singkat, slang, serta campur kode antara bahasa Indonesia dan bahasa asing. Fenomena tersebut tidak selalu dapat dinilai negatif karena variasi bahasa merupakan bagian dari dinamika sosial. Namun, masalah muncul ketika remaja tidak mampu membedakan konteks penggunaan bahasa santai, bahasa baku, campur kode, dan bahasa Indonesia yang sesuai norma.",
        "Data kebahasaan terbaru juga menunjukkan perlunya penguatan kemampuan berbahasa Indonesia pada jenjang pelajar. Badan Bahasa melalui diseminasi UKBI menyatakan bahwa data tahun 2024 menunjukkan pelajar SMA/SMK masih berada di bawah standar kemahiran yang ditetapkan. Informasi ini memperlihatkan bahwa penguatan bahasa Indonesia pada kelompok pelajar masih menjadi kebutuhan nyata, bukan hanya asumsi umum.",
        "Di lingkungan remaja, penggunaan ungkapan seperti 'aku happy banget', 'lagi healing', 'deadline tugas', atau 'vibes kelas hari ini' menunjukkan bahwa bahasa Indonesia sering bercampur dengan bahasa asing dalam komunikasi sehari-hari. Pola ini perlu dilihat secara kritis: campur kode dapat menjadi strategi komunikasi, tetapi loyalitas berbahasa Indonesia tetap perlu dibangun agar remaja memiliki kebanggaan, kesadaran norma, dan kemampuan memilih ragam bahasa sesuai situasi.",
        "Salah satu pendekatan yang relevan untuk menjawab persoalan tersebut adalah language awareness atau kesadaran berbahasa. Pendekatan ini menekankan pemahaman terhadap fungsi, struktur, norma, dan peran sosial bahasa. Melalui language awareness, remaja tidak sekadar diminta menggunakan bahasa Indonesia yang baik dan benar, tetapi diajak memahami mengapa pilihan bahasa tertentu tepat atau tidak tepat dalam konteks komunikasi tertentu.",
        "Berdasarkan kebutuhan tersebut, penelitian ini mengembangkan PRIMA+ sebagai platform digital berbasis language awareness. Platform ini dirancang dalam bentuk media interaktif dan gamifikasi sederhana yang memuat kuis, refleksi bahasa, contoh kasus komunikasi digital, dan umpan balik pembelajaran. PRIMA+ diharapkan dapat meningkatkan loyalitas berbahasa Indonesia remaja melalui pengalaman belajar yang dekat dengan kebiasaan digital mereka.",
    ]
    for para in paragraphs:
        doc.add_paragraph(para)

    doc.add_heading("1.2 Rumusan Masalah", level=2)
    add_numbered(
        doc,
        [
            "Bagaimana rancangan platform PRIMA+ berbasis language awareness untuk meningkatkan loyalitas berbahasa Indonesia remaja?",
            "Bagaimana efektivitas penggunaan platform PRIMA+ terhadap peningkatan loyalitas berbahasa Indonesia remaja berdasarkan indikator sikap positif, kesetiaan penggunaan, kesadaran norma, kebanggaan bahasa, dan kemampuan memilih ragam bahasa sesuai konteks?",
        ],
    )

    doc.add_heading("1.3 Tujuan Penelitian", level=2)
    add_numbered(
        doc,
        [
            "Mengembangkan rancangan platform PRIMA+ berbasis language awareness sebagai media penguatan loyalitas berbahasa Indonesia remaja.",
            "Menguji efektivitas penggunaan platform PRIMA+ dalam meningkatkan loyalitas berbahasa Indonesia remaja berdasarkan hasil pretest, posttest, kuesioner sikap, dan aktivitas kuis pada platform.",
        ],
    )

    doc.add_heading("1.4 Manfaat Penelitian", level=2)
    doc.add_heading("1.4.1 Manfaat Teoretis", level=3)
    doc.add_paragraph(
        "Penelitian ini diharapkan memberikan kontribusi bagi kajian sosiolinguistik dan pembelajaran bahasa Indonesia, khususnya pada pengembangan model language awareness berbasis platform digital untuk membangun loyalitas berbahasa pada remaja."
    )
    doc.add_heading("1.4.2 Manfaat Praktis", level=3)
    add_bullets(
        doc,
        [
            "Bagi remaja, PRIMA+ dapat membantu meningkatkan kesadaran norma bahasa, kebanggaan terhadap bahasa Indonesia, dan kemampuan memilih ragam bahasa sesuai konteks.",
            "Bagi guru Bahasa Indonesia, penelitian ini dapat menjadi contoh pemanfaatan media digital interaktif untuk pembelajaran sikap bahasa.",
            "Bagi sekolah, penelitian ini dapat menjadi program penguatan literasi dan karakter kebahasaan di lingkungan madrasah.",
            "Bagi pengembang media pendidikan, penelitian ini memberikan rancangan awal platform gamifikasi untuk pembelajaran bahasa Indonesia.",
        ],
    )

    doc.add_heading("BAB 2. TINJAUAN PUSTAKA", level=1)
    doc.add_heading("2.1 Platform Digital dan Gamifikasi Pembelajaran Bahasa", level=2)
    doc.add_paragraph(
        "Platform digital memberi peluang bagi pembelajaran bahasa yang lebih fleksibel, interaktif, dan dekat dengan kebiasaan remaja sebagai pengguna teknologi. Dalam pembelajaran bahasa, gamifikasi dapat meningkatkan motivasi dan keterlibatan, tetapi efektivitasnya tetap bergantung pada kualitas desain aktivitas, umpan balik, dan kesesuaian tujuan pembelajaran."
    )
    doc.add_heading("2.2 Language Awareness", level=2)
    doc.add_paragraph(
        "Language awareness merupakan kesadaran seseorang terhadap fungsi, struktur, norma, dan peran sosial bahasa. Dalam konteks penelitian ini, language awareness digunakan untuk membantu remaja memahami kapan bahasa baku, bahasa santai, slang, atau campur kode dapat digunakan secara tepat."
    )
    doc.add_heading("2.3 Loyalitas Berbahasa", level=2)
    doc.add_paragraph(
        "Loyalitas berbahasa adalah sikap positif untuk mempertahankan dan menggunakan bahasa tertentu sebagai bagian dari identitas sosial dan budaya. Dalam penelitian ini, loyalitas berbahasa Indonesia diukur melalui lima indikator: sikap positif terhadap bahasa Indonesia, kesetiaan penggunaan, kesadaran norma bahasa, kebanggaan bahasa, dan kemampuan memilih ragam bahasa sesuai konteks."
    )
    add_table(
        doc,
        ["Indikator", "Makna Operasional", "Contoh Pengukuran"],
        [
            ["Sikap positif", "Pandangan bahwa bahasa Indonesia penting dan bernilai", "Skor persetujuan terhadap pernyataan sikap"],
            ["Kesetiaan penggunaan", "Kecenderungan memilih bahasa Indonesia pada konteks sekolah dan digital", "Pilihan bahasa dalam skenario komunikasi"],
            ["Kesadaran norma", "Pemahaman terhadap ragam baku, santai, slang, dan campur kode", "Jawaban kuis tentang kesesuaian bahasa"],
            ["Kebanggaan bahasa", "Rasa bangga menggunakan bahasa Indonesia", "Respons skala Likert tentang kebanggaan bahasa"],
            ["Pemilihan ragam", "Kemampuan menyesuaikan bahasa dengan konteks", "Tes kasus komunikasi formal dan informal"],
        ],
        widths=[3.3, 6.2, 6.5],
    )
    doc.add_heading("2.4 Studi Pustaka Terdahulu", level=2)
    add_table(
        doc,
        ["No.", "Sumber", "Fokus", "Relevansi terhadap PRIMA+"],
        [
            ["1", "Fairclough (1992)", "Critical language awareness", "Menjadi dasar bahwa kesadaran bahasa terkait konteks sosial dan sikap kritis."],
            ["2", "Prensky (2001)", "Digital natives", "Menjelaskan karakter remaja yang dekat dengan teknologi digital."],
            ["3", "Crystal (2011)", "Internet linguistics", "Menjelaskan perubahan bahasa dalam komunikasi internet."],
            ["4", "Shortt dkk. (2021)", "Gamification in mobile-assisted language learning", "Mendukung penggunaan gamifikasi untuk meningkatkan keterlibatan pembelajaran bahasa."],
            ["5", "Badan Bahasa (2025)", "Data UKBI pelajar", "Menjadi data empiris bahwa penguatan kemahiran bahasa Indonesia pada pelajar masih diperlukan."],
            ["6", "Kajian campur kode media sosial 2024-2025", "Campur kode remaja di media sosial", "Memperkuat fenomena bahasa campuran sebagai konteks masalah penelitian."],
        ],
        widths=[1.2, 4.1, 4.3, 6.4],
    )
    doc.add_paragraph(
        "Berdasarkan kajian tersebut, penelitian ini memiliki posisi pada pengembangan media pembelajaran bahasa Indonesia yang tidak hanya menguji pemahaman kebahasaan, tetapi juga membangun loyalitas berbahasa melalui pendekatan language awareness dan gamifikasi."
    )

    doc.add_heading("BAB 3. METODE PENELITIAN", level=1)
    doc.add_heading("3.1 Desain Penelitian", level=2)
    doc.add_paragraph(
        "Penelitian ini menggunakan pendekatan Research and Development (R&D) sederhana dengan uji efektivitas one-group pretest-posttest. R&D digunakan untuk mengembangkan platform PRIMA+, sedangkan pretest-posttest digunakan untuk melihat perubahan loyalitas berbahasa Indonesia sebelum dan sesudah penggunaan platform."
    )
    doc.add_heading("3.2 Waktu dan Tempat Penelitian", level=2)
    doc.add_paragraph(
        "Penelitian dilaksanakan pada April-Agustus 2026 di MAN Kotawaringin Timur, Kalimantan Tengah. Tahapan penelitian mencakup studi literatur, perancangan platform, validasi sederhana, uji coba kepada responden, analisis data, dan penyusunan laporan akhir."
    )
    doc.add_heading("3.3 Subjek dan Teknik Pengambilan Sampel", level=2)
    doc.add_paragraph(
        "Subjek penelitian adalah siswa MAN Kotawaringin Timur yang aktif menggunakan media digital. Uji coba awal direncanakan melibatkan 15-30 responden. Teknik pengambilan sampel menggunakan purposive sampling dengan kriteria: siswa aktif, bersedia mengikuti seluruh rangkaian uji coba, memiliki akses ke telepon pintar, dan memperoleh izin sesuai ketentuan sekolah."
    )
    doc.add_heading("3.4 Variabel dan Indikator Penelitian", level=2)
    add_table(
        doc,
        ["Variabel", "Indikator", "Sumber Data"],
        [
            ["Platform PRIMA+", "Ketersediaan fitur kuis, contoh kasus bahasa digital, skor, dan umpan balik", "Validasi media dan uji coba sistem"],
            ["Loyalitas berbahasa Indonesia", "Sikap positif, kesetiaan penggunaan, kesadaran norma, kebanggaan, pemilihan ragam bahasa", "Kuesioner pretest-posttest dan hasil kuis"],
            ["Efektivitas", "Peningkatan skor posttest dibanding pretest", "Analisis skor, persentase peningkatan, dan N-gain"],
        ],
        widths=[3.7, 7.1, 5.2],
    )
    doc.add_heading("3.5 Alat dan Bahan", level=2)
    add_table(
        doc,
        ["No.", "Alat/Bahan", "Jumlah", "Keterangan"],
        [
            ["1", "Laptop pengembang", "1 unit", "Mendesain platform dan mengolah data"],
            ["2", "Telepon pintar responden", "15-30 unit", "Perangkat uji coba platform"],
            ["3", "Godot Engine", "1 aplikasi", "Game engine gratis/open-source untuk pengembangan prototipe"],
            ["4", "Kuesioner pretest dan posttest", "1 paket", "Mengukur loyalitas dan pemahaman bahasa"],
            ["5", "Skenario soal PRIMA+", "1 paket", "Materi kuis dan contoh kasus bahasa digital"],
            ["6", "Aset visual/audio", "1 paket", "Mendukung tampilan dan pengalaman pengguna"],
        ],
        widths=[1.2, 5.0, 3.0, 6.8],
    )
    doc.add_heading("3.6 Prosedur Penelitian", level=2)
    add_numbered(
        doc,
        [
            "Studi literatur dan analisis kebutuhan melalui kajian teori, penelitian terdahulu, serta observasi awal penggunaan bahasa remaja.",
            "Perancangan platform PRIMA+ yang memuat kuis bahasa, contoh kasus komunikasi digital, skor, dan umpan balik.",
            "Validasi sederhana oleh guru Bahasa Indonesia atau ahli media untuk menilai kesesuaian materi dan tampilan.",
            "Pelaksanaan pretest untuk mengukur kondisi awal loyalitas berbahasa dan pemahaman penggunaan bahasa Indonesia.",
            "Uji coba penggunaan platform PRIMA+ oleh responden.",
            "Pelaksanaan posttest setelah responden menggunakan platform.",
            "Analisis data untuk melihat peningkatan skor dan efektivitas platform.",
            "Revisi akhir platform dan penyusunan laporan penelitian.",
        ],
    )
    flow = Path("prima_media/image2.jpeg")
    if flow.exists():
        doc.add_picture(str(flow), width=Cm(8.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph("Gambar 1. Alur prosedur penelitian PRIMA+")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3.7 Teknik Pengumpulan Data", level=2)
    add_bullets(
        doc,
        [
            "Kuesioner pretest-posttest skala Likert untuk mengukur loyalitas berbahasa.",
            "Tes pemahaman penggunaan bahasa Indonesia melalui kuis PRIMA+.",
            "Log aktivitas platform berupa skor, jumlah jawaban benar/salah, dan penyelesaian misi.",
            "Validasi ahli/guru terhadap materi, bahasa, dan tampilan platform.",
            "Catatan observasi atau refleksi singkat responden setelah menggunakan platform.",
        ],
    )
    doc.add_heading("3.8 Teknik Analisis Data", level=2)
    doc.add_paragraph(
        "Data dianalisis secara kuantitatif deskriptif dengan menghitung rata-rata skor pretest dan posttest, persentase ketepatan jawaban, serta persentase peningkatan. Untuk mengukur efektivitas, penelitian dapat menggunakan N-gain. Jika jumlah data dan distribusi memungkinkan, peningkatan skor dapat diuji menggunakan paired sample t-test; jika data tidak berdistribusi normal, digunakan uji Wilcoxon."
    )
    add_table(
        doc,
        ["Aspek", "Teknik Analisis", "Kriteria"],
        [
            ["Pemahaman bahasa", "Persentase jawaban benar dan rata-rata skor kuis", "Meningkat setelah penggunaan PRIMA+"],
            ["Loyalitas berbahasa", "Perbandingan skor pretest-posttest kuesioner", "Posttest lebih tinggi daripada pretest"],
            ["Efektivitas platform", "N-gain atau persentase peningkatan", "Kategori rendah, sedang, atau tinggi sesuai hasil"],
            ["Kelayakan media", "Rata-rata skor validasi ahli/guru", "Layak digunakan setelah revisi"],
        ],
        widths=[4.0, 6.6, 5.4],
    )

    doc.add_heading("BAB 4. RANCANGAN ANGGARAN BIAYA DAN JADWAL", level=1)
    doc.add_heading("4.1 Rancangan Anggaran Biaya", level=2)
    add_table(
        doc,
        ["No.", "Komponen", "Jumlah", "Harga Satuan", "Subtotal", "Keterangan"],
        [
            ["1", "Headset/audio uji coba", "1 unit", "Rp224.000", "Rp224.000", "Menguji audio permainan"],
            ["2", "Flashdisk/penyimpanan", "1 unit", "Rp96.000", "Rp96.000", "Backup data dan platform"],
            ["3", "Aset visual game", "1 paket", "Rp50.000", "Rp50.000", "Sprite karakter dan objek"],
            ["4", "Aset audio permainan", "1 paket", "Rp150.000", "Rp150.000", "Efek suara dan musik latar"],
            ["5", "Canva/desain grafis", "1 bulan", "Rp20.000", "Rp20.000", "Desain pendukung instrumen dan poster"],
            ["6", "Akses internet", "1 paket", "Rp350.000", "Rp350.000", "Pengembangan, referensi, dan uji coba"],
            ["7", "Cetak instrumen dan dokumentasi", "1 paket", "Rp150.000", "Rp150.000", "Pretest, posttest, dan dokumentasi"],
            ["8", "Biaya operasional uji coba", "1 paket", "Rp152.000", "Rp152.000", "Koordinasi dan kebutuhan teknis lapangan"],
            ["", "Total", "", "", "Rp1.192.000", ""],
        ],
        widths=[1.0, 4.2, 2.0, 2.7, 2.7, 3.4],
    )
    doc.add_paragraph(
        "Catatan: Godot Engine tidak dimasukkan sebagai biaya lisensi karena bersifat gratis/open-source. Biaya yang dicantumkan difokuskan pada aset, perangkat pendukung, internet, dokumentasi, dan operasional uji coba."
    )

    doc.add_heading("4.2 Jadwal Penelitian", level=2)
    add_table(
        doc,
        ["No.", "Kegiatan", "Tanggal Pelaksanaan", "Luaran"],
        [
            ["1", "Studi literatur dan analisis kebutuhan", "6-16 April 2026", "Dasar teori dan indikator penelitian"],
            ["2", "Perancangan konsep dan desain PRIMA+", "17-30 April 2026", "Rancangan fitur dan skenario kuis"],
            ["3", "Pengembangan prototipe dan validasi sederhana", "1-25 Mei 2026", "Prototipe awal dan masukan guru/ahli"],
            ["4", "Uji coba platform kepada responden", "26 Mei-20 Juni 2026", "Data pretest, posttest, dan log kuis"],
            ["5", "Revisi platform dan pengolahan data", "21 Juni-25 Juli 2026", "Data siap analisis"],
            ["6", "Analisis data dan penyusunan laporan akhir", "26 Juli-20 Agustus 2026", "Laporan penelitian akhir"],
        ],
        widths=[1.0, 5.4, 4.4, 5.2],
    )

    doc.add_heading("PERNYATAAN PENGGUNAAN ARTIFICIAL INTELLIGENCE (AI)", level=1)
    doc.add_paragraph(
        "Penelitian ini menggunakan bantuan Artificial Intelligence (AI), yaitu Gemini dan ChatGPT, untuk membantu mencari referensi awal, memperdalam landasan teori, merapikan kerangka metodologi, dan menyusun ide alur permainan atau kuis bahasa. Seluruh keputusan penelitian, penyusunan instrumen, pengembangan produk, analisis data, dan penarikan kesimpulan tetap dilakukan secara mandiri oleh peneliti. AI tidak digunakan untuk membuat data penelitian palsu, menggantikan proses uji coba, atau menulis hasil penelitian tanpa verifikasi."
    )

    doc.add_heading("DAFTAR PUSTAKA", level=1)
    refs = [
        "Alwi, H., & Sugono, D. (2011). Politik Bahasa Nasional. Badan Pengembangan dan Pembinaan Bahasa.",
        "Badan Bahasa. (2025). UKBI jadi instrumen strategis wujudkan kemahiran berbahasa Indonesia yang terstandar. Balai Bahasa Provinsi Jambi/Kemendikdasmen.",
        "Carter, R. (2003). Language awareness. ELT Journal, 57(1), 64-65.",
        "Crystal, D. (2011). Internet Linguistics: A Student Guide. Routledge.",
        "Fairclough, N. (1992). Critical Language Awareness. Longman.",
        "Prensky, M. (2001). Digital natives, digital immigrants. On the Horizon, 9(5).",
        "Shortt, M., Tilak, S., Kuznetcova, I., Martens, B., & Akinkuolie, B. (2021). Gamification in mobile-assisted language learning: A systematic review. Computer Assisted Language Learning.",
        "Sugono, D. (2019). Bahasa Indonesia di Era Globalisasi. Badan Pengembangan dan Pembinaan Bahasa.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(OUT_PROPOSAL)


def revision_table_doc():
    doc = Document()
    style_doc(doc)
    add_title(doc, "TABEL DAFTAR REVISI BERDASARKAN CATATAN REVIEWER")
    doc.add_paragraph(
        "Dokumen ini memetakan catatan reviewer dari revisi.jpeg dengan bentuk perbaikan yang sudah diterapkan pada naskah prima_revisi.docx."
    )
    rows = [
        ["1", "Judul menarik, tetapi terlalu panjang.", "Judul diringkas menjadi 'PRIMA+: Platform Language Awareness untuk Meningkatkan Loyalitas Berbahasa Indonesia Remaja'.", "Halaman judul/cover.", "Judul tetap memuat PRIMA+, language awareness, loyalitas berbahasa Indonesia, dan sasaran remaja, tetapi lebih singkat dan langsung."],
        ["2", "Latar belakang terlalu umum dan belum didukung data empiris.", "Ditambahkan data empiris UKBI 2024 tentang pelajar SMA/SMK yang masih di bawah standar kemahiran, serta contoh nyata bahasa campuran remaja di media digital.", "BAB 1, bagian 1.1 Latar Belakang.", "Latar belakang kini tidak hanya menjelaskan fenomena umum, tetapi memberi dasar kebutuhan penelitian dari data kebahasaan dan contoh komunikasi digital."],
        ["3", "Rumusan masalah terlalu luas dan belum menunjukkan indikator jelas.", "Rumusan masalah kedua ditajamkan dengan indikator loyalitas: sikap positif, kesetiaan penggunaan, kesadaran norma, kebanggaan bahasa, dan kemampuan memilih ragam bahasa.", "BAB 1, bagian 1.2 Rumusan Masalah.", "Pertanyaan penelitian menjadi lebih terukur dan langsung terhubung dengan instrumen serta analisis data."],
        ["4", "Manfaat penelitian masih umum dan belum menunjukkan kontribusi spesifik.", "Manfaat dipisahkan menjadi manfaat teoretis dan praktis untuk remaja, guru Bahasa Indonesia, sekolah, dan pengembang media pendidikan.", "BAB 1, bagian 1.4 Manfaat Penelitian.", "Kontribusi setiap pihak dijelaskan secara spesifik sehingga manfaat penelitian lebih konkret."],
        ["5", "Studi pustaka masih sedikit, tambahkan.", "Studi pustaka diperluas dengan platform digital, gamifikasi pembelajaran bahasa, language awareness, loyalitas berbahasa, internet linguistics, data UKBI, dan kajian campur kode media sosial.", "BAB 2, bagian 2.1-2.4 dan Daftar Pustaka.", "Landasan teori dan posisi riset menjadi lebih kuat serta lebih sesuai dengan isu remaja digital."],
        ["6", "Metode belum merinci jumlah responden.", "Ditambahkan rencana uji coba awal 15-30 responden.", "BAB 3, bagian 3.3 Subjek dan Teknik Pengambilan Sampel.", "Jumlah responden dibuat eksplisit sehingga desain uji coba lebih jelas."],
        ["7", "Metode belum merinci teknik pengambilan sampel.", "Ditambahkan teknik purposive sampling dengan kriteria siswa aktif, bersedia mengikuti uji coba, memiliki akses telepon pintar, dan mengikuti ketentuan izin sekolah.", "BAB 3, bagian 3.3 Subjek dan Teknik Pengambilan Sampel.", "Pemilihan responden menjadi memiliki dasar metodologis."],
        ["8", "Desain uji efektivitas belum jelas.", "Desain ditegaskan sebagai Research and Development sederhana dengan one-group pretest-posttest.", "BAB 3, bagian 3.1 Desain Penelitian.", "Reviewer dapat melihat hubungan antara pengembangan platform dan pengukuran efektivitas."],
        ["9", "Indikator pengukuran loyalitas belum diungkapkan.", "Ditambahkan tabel indikator loyalitas berbahasa: sikap positif, kesetiaan penggunaan, kesadaran norma, kebanggaan, dan pemilihan ragam.", "BAB 2, bagian 2.3 dan BAB 3, bagian 3.4.", "Variabel loyalitas berbahasa kini dapat diukur melalui kuesioner, kuis, dan skenario komunikasi."],
        ["10", "Prosedur penelitian belum menunjukkan metode yang jelas.", "Prosedur disusun menjadi delapan tahap: studi literatur, perancangan, validasi, pretest, uji coba, posttest, analisis, revisi akhir.", "BAB 3, bagian 3.6 Prosedur Penelitian.", "Alur penelitian lebih runtut dan sesuai desain R&D + pretest-posttest."],
        ["11", "Analisis data kurang kuat untuk mengukur peningkatan loyalitas.", "Ditambahkan teknik analisis rata-rata pretest-posttest, persentase peningkatan, N-gain, dan opsi uji paired sample t-test atau Wilcoxon.", "BAB 3, bagian 3.8 Teknik Analisis Data.", "Analisis tidak hanya deskriptif sederhana, tetapi memiliki ukuran peningkatan yang lebih kuat."],
        ["12", "Daftar pustaka belum lengkap dan perlu perbaikan.", "Daftar pustaka dirapikan, kesalahan 'Logman' diperbaiki menjadi Longman, Carter disesuaikan, dan referensi gamifikasi serta data UKBI ditambahkan.", "Daftar Pustaka.", "Referensi menjadi lebih konsisten dan mendukung revisi teori/metode."],
        ["13", "Anggaran Rp1.192.000 kurang realistis karena komponen belum rinci.", "RAB dipecah menjadi komponen, jumlah, harga satuan, subtotal, dan keterangan. Godot dijelaskan sebagai gratis/open-source.", "BAB 4, bagian 4.1 Rancangan Anggaran Biaya.", "Anggaran tetap pada total Rp1.192.000, tetapi rincian dan alasan kebutuhannya lebih jelas."],
    ]
    add_table(
        doc,
        ["No.", "Catatan Reviewer", "Perbaikan yang Dilakukan", "Letak Perbaikan", "Penjelasan Bentuk Perbaikan"],
        rows,
        widths=[1.0, 3.6, 4.4, 3.2, 4.0],
    )
    doc.save(OUT_TABLE)


if __name__ == "__main__":
    proposal_doc()
    revision_table_doc()
    print(OUT_PROPOSAL)
    print(OUT_TABLE)
