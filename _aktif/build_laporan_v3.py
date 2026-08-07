# -*- coding: utf-8 -*-
"""Bangun Laporan (Full Paper) PRIMA+ OPSI 2026 — versi v3 sesuai Panduan.

Perbedaan dengan v2 (build_full_paper_opsi_2026.py):
1. Urutan BAB mengikuti Panduan OPSI 2026 (BAB 1 Pendahuluan -> BAB 5 Kesimpulan & Saran).
2. RAB dipindah ke LAMPIRAN (sesuai ketentuan Lampiran pada panduan), bukan BAB tersendiri.
3. Ditambah bagian UCAPAN TERIMA KASIH dan PERNYATAAN PENGGUNAAN AI (mandatory panduan).
4. Abstrak dinetralkan (tidak lagi mengklaim hasil sebelum data terkumpul).
5. BAB 4 & 5 memakai placeholder jujur [isi data] — data penelitian belum dikumpulkan.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

import os

OUT = "Laporan_Penelitian_PRIMA_OPSI_2026_v3_Struktur_Panduan.docx"
FIGURE = "prima_prosedur_ADDIE_HIGH.png"

TITLE = "KESADARAN BERBAHASA REMAJA MELALUI MEDIA PRIMA+ UNTUK MENGUATKAN LOYALITAS BAHASA INDONESIA DI LINGKUNGAN SEKOLAH"


def set_run(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_p(paragraph, align=None, before=0, after=6, line=1.15):
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_p(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, bold=False, italic=False, after=6, before=0):
    p = doc.add_paragraph()
    set_p(p, align=align, after=after, before=before)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, italic=italic)
    return p


def add_center(doc, text, size=12, bold=True, after=8):
    return add_p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=size, bold=bold, after=after)


def add_heading(doc, text, level=1, size=None):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    set_p(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=10 if level == 1 else 6, after=5)
    r = p.add_run(text)
    set_run(r, size=size if size is not None else (14 if level == 1 else 12), bold=True)
    return p


def add_list(doc, text):
    return add_p(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3)


def shade_cell(cell, fill="EDEDED"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    set_p(p, align=align, after=0, line=1.0)
    r = p.add_run(str(text))
    set_run(r, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, caption=None, widths=None):
    if caption:
        p = add_p(doc, caption, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, after=3)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    return doc


def build():
    doc = setup_doc()

    # ── HALAMAN SAMPUL ──
    for _ in range(2):
        doc.add_paragraph()
    add_center(doc, "LAPORAN PENELITIAN OPSI", size=14, bold=True, after=18)
    add_center(doc, TITLE, size=14, bold=True, after=22)
    add_center(doc, "Nama Tim Peneliti", size=12, bold=True, after=4)
    add_center(doc, "Ketua : Larisa Dianti", size=12, bold=True, after=2)
    add_center(doc, "Anggota : Gusti Nazwa Azizah", size=12, bold=True, after=14)
    add_center(doc, "Bidang Kompetisi Penelitian: Ilmu Sosial dan Humaniora (ISH)", size=12, bold=True, after=2)
    add_center(doc, "Subbidang: Bahasa dan Sastra", size=12, bold=True, after=14)
    add_center(doc, "MAN KOTAWARINGIN TIMUR", size=12, bold=True, after=2)
    add_center(doc, "Kotawaringin Timur, Kalimantan Tengah", size=12, bold=True, after=14)
    add_center(doc, "Tahun 2026", size=12, bold=True)
    doc.add_page_break()

    # ── ABSTRAK (dinetralkan: belum ada klaim hasil) ──
    add_heading(doc, "ABSTRAK", 1)
    add_p(doc, "Penelitian ini bertujuan mengembangkan media PRIMA+ berbasis kesadaran berbahasa untuk menguatkan loyalitas bahasa Indonesia remaja di lingkungan sekolah. PRIMA+ dirancang sebagai media pembelajaran/kebahasaan yang menyajikan latihan pilihan ragam bahasa, refleksi konteks komunikasi, dan umpan balik. Penelitian menggunakan metode Research and Development (R&D) dengan model ADDIE. Uji coba terbatas direncanakan kepada 30 siswa MAN Kotawaringin Timur dengan desain one-group pretest-posttest. Kelayakan media dinilai melalui validasi ahli pada aspek materi, bahasa, media, dan instrumen, sedangkan efektivitas diukur melalui perbandingan skor pretest dan posttest serta uji beda berpasangan. Hasil pengukuran akan disajikan pada bagian hasil dan pembahasan setelah pengumpulan dan pengolahan data penelitian.")
    add_p(doc, "Kata kunci: kesadaran berbahasa, loyalitas bahasa Indonesia, remaja, media pembelajaran, PRIMA+", italic=True, before=8)
    doc.add_page_break()

    # ── DAFTAR ISI ──
    add_heading(doc, "DAFTAR ISI", 1)
    for line in [
        "BAB 1. PENDAHULUAN",
        "1.1 Latar Belakang",
        "1.2 Rumusan Masalah",
        "1.3 Tujuan Penelitian",
        "1.4 Hipotesis",
        "1.5 Manfaat Penelitian",
        "BAB 2. LANDASAN TEORI DAN STUDI PUSTAKA",
        "2.1 Landasan Teori",
        "2.2 Studi Pustaka",
        "BAB 3. METODE PENELITIAN",
        "3.1 Waktu dan Tempat Penelitian",
        "3.2 Alat dan Bahan",
        "3.3 Prosedur Penelitian",
        "3.4 Pengolahan Data",
        "BAB 4. HASIL DAN PEMBAHASAN",
        "4.1 Hasil Pengembangan Media PRIMA+",
        "4.2 Hasil Validasi Media PRIMA+",
        "4.3 Hasil Uji Efektivitas PRIMA+",
        "4.4 Pembahasan",
        "BAB 5. ANGGARAN BIAYA DAN KEGIATAN",
        "5.1 Anggaran Biaya Penelitian",
        "5.2 Kegiatan",
        "KESIMPULAN DAN SARAN",
        "KESIMPULAN",
        "SARAN",
        "UCAPAN TERIMA KASIH",
        "PERNYATAAN PENGGUNAAN AI",
        "DAFTAR PUSTAKA",
        "LAMPIRAN",
    ]:
        add_p(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)

    add_heading(doc, "DAFTAR TABEL", 1)
    for line in [
        "Tabel 1. Studi pustaka terdahulu",
        "Tabel 2. Alat dan bahan penelitian",
        "Tabel 3. Konstruk/aspek, indikator, dan instrumen",
        "Tabel 4. Rancangan pengolahan data",
        "Tabel 5. Hasil validasi media PRIMA+",
        "Tabel 6. Perbandingan skor pretest dan posttest",
        "Tabel 7. Rencana dan realisasi anggaran biaya penelitian",
        "Tabel 8. Jadwal dan realisasi kegiatan penelitian",
    ]:
        add_p(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, after=2)

    add_heading(doc, "DAFTAR GAMBAR", 1)
    add_p(doc, "Gambar 1. Alur prosedur penelitian PRIMA+ berbasis ADDIE", align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

    # ── BAB 1. PENDAHULUAN ──
    add_heading(doc, "BAB 1. PENDAHULUAN", 1, size=16)
    add_heading(doc, "1.1 LATAR BELAKANG", 2)
    for text in [
        "Bahasa Indonesia digunakan siswa dalam kegiatan belajar, diskusi kelas, organisasi, penulisan tugas, dan komunikasi digital. Di lingkungan sekolah, kemampuan memilih ragam bahasa menjadi bagian penting dari kedisiplinan akademik dan sikap berbahasa. Tantangannya, remaja kini bergerak di dua ruang sekaligus: ruang sekolah yang menuntut bahasa Indonesia yang jelas dan tertib, serta ruang digital yang cenderung cepat, ringkas, santai, dan campur kode.",
        "Konteks penelitian ini ditempatkan di MAN Kotawaringin Timur, Sampit. Data Referensi Kemendikdasmen mencatat MAN Kotawaringin Timur sebagai madrasah negeri di Jalan H.M. Arsyad No. 68 Sampit, Kecamatan Mentawa Baru Ketapang, Kabupaten Kotawaringin Timur, dengan total peserta didik 402 orang pada pembaruan data 26 Juni 2026. Jumlah peserta didik tersebut menunjukkan bahwa sekolah memiliki lingkungan sosial yang cukup besar untuk mengamati praktik bahasa remaja dalam komunikasi akademik dan digital.",
        "Masalah yang tampak bukan sekadar penggunaan slang atau istilah asing, melainkan melemahnya kepekaan siswa dalam membedakan kapan bahasa santai dapat digunakan dan kapan bahasa Indonesia yang lebih tertib diperlukan. Dalam konteks pembelajaran Bahasa Indonesia, kondisi ini berkaitan dengan kesadaran berbahasa, sikap terhadap bahasa Indonesia, dan loyalitas berbahasa Indonesia sebagai bagian dari identitas dan komunikasi sekolah.",
        "Kajian tentang remaja Sampit juga menunjukkan bahwa media dapat berperan dalam membentuk cara remaja memaknai identitas dan kehidupan sosialnya. Widyaningrum, Yumarnamto, dan Prijambodo (2020) meneliti remaja Kota Sampit sebagai khalayak media dan menunjukkan pentingnya memahami remaja melalui konteks lokal, identitas, dan pengalaman sosialnya. Oleh karena itu, penelitian bahasa remaja di Sampit perlu ditempatkan sebagai kajian sosial-humaniora, bukan semata-mata sebagai pengembangan teknologi.",
        "PRIMA+ dalam penelitian ini diposisikan sebagai media pembelajaran/kebahasaan, bukan sebagai produk rekayasa perangkat lunak. Unsur digital dipakai untuk menghadirkan kasus bahasa, pilihan ragam, kuis kontekstual, refleksi, dan umpan balik. Fokus penelitian tetap berada pada perubahan kesadaran berbahasa dan loyalitas bahasa Indonesia remaja di lingkungan sekolah.",
        "Kebaruan penelitian ini terletak pada penggunaan media PRIMA+ sebagai sarana latihan pengambilan keputusan berbahasa. Penelitian tidak hanya mendeskripsikan fenomena bahasa digital remaja, tetapi mengembangkan media yang dapat membantu siswa menilai konteks, memilih ragam bahasa, dan merefleksikan alasan penggunaan bahasa Indonesia secara tepat.",
    ]:
        add_p(doc, text)

    add_heading(doc, "1.2 RUMUSAN MASALAH", 2)
    for text in [
        "1. Bagaimana rancangan media PRIMA+ berbasis kesadaran berbahasa untuk menguatkan loyalitas bahasa Indonesia remaja di lingkungan sekolah?",
        "2. Bagaimana kelayakan materi, bahasa, media, dan instrumen PRIMA+ berdasarkan validasi guru atau ahli?",
        "3. Bagaimana efektivitas penggunaan PRIMA+ terhadap peningkatan loyalitas berbahasa Indonesia remaja?",
    ]:
        add_list(doc, text)

    add_heading(doc, "1.3 TUJUAN PENELITIAN", 2)
    for text in [
        "1. Mengembangkan media PRIMA+ berbasis kesadaran berbahasa sesuai kebutuhan remaja di lingkungan sekolah.",
        "2. Menilai kelayakan PRIMA+ dari aspek materi, bahasa, media, dan instrumen penelitian.",
        "3. Menguji efektivitas PRIMA+ dalam menguatkan loyalitas berbahasa Indonesia remaja melalui desain pretest-posttest.",
    ]:
        add_list(doc, text)

    add_heading(doc, "1.4 HIPOTESIS", 2)
    add_p(doc, "Berdasarkan landasan teori dan penelitian terdahulu yang dituangkan pada Bab 2, dirumuskan hipotesis penelitian sebagai berikut: penggunaan media PRIMA+ efektif meningkatkan loyalitas berbahasa Indonesia remaja di lingkungan sekolah. Hipotesis ini akan diuji melalui perbandingan skor pretest dan posttest pada desain one-group pretest-posttest. Jika data tidak memenuhi asumsi uji parametrik, pengujian menggunakan alternatif nonparametrik yang sesuai.")

    add_heading(doc, "1.5 MANFAAT PENELITIAN", 2)
    for text in [
        "Bagi siswa, PRIMA+ membantu mereka lebih sadar dalam memilih ragam bahasa sesuai konteks sekolah dan media digital.",
        "Bagi guru Bahasa Indonesia, PRIMA+ dapat menjadi contoh media pembelajaran yang menghubungkan materi bahasa dengan situasi komunikasi nyata siswa.",
        "Bagi sekolah, penelitian ini mendukung pembiasaan berbahasa Indonesia yang tertib, santun, dan sesuai konteks akademik.",
        "Bagi kajian bahasa dan sastra, penelitian ini memberi contoh pengembangan media berbasis kesadaran berbahasa untuk menguatkan sikap dan loyalitas bahasa Indonesia remaja.",
    ]:
        add_list(doc, text)

    # ── BAB 2. LANDASAN TEORI DAN STUDI PUSTAKA ──
    add_heading(doc, "BAB 2. LANDASAN TEORI DAN STUDI PUSTAKA", 1, size=16)
    add_heading(doc, "2.1 LANDASAN TEORI", 2)
    for text in [
        "Kesadaran berbahasa menekankan kemampuan memperhatikan bentuk, fungsi, konteks, dan dampak sosial dari bahasa yang digunakan. Carter (2003) memandang language awareness sebagai kesadaran terhadap bentuk dan fungsi bahasa, sedangkan Fairclough (1992) menekankan pentingnya sikap kritis dalam melihat hubungan bahasa dan kehidupan sosial.",
        "Loyalitas bahasa merupakan bagian dari sikap bahasa. Garvin dan Mathiot (1968) menempatkan loyalitas bahasa bersama kebanggaan bahasa dan kesadaran norma. Dalam penelitian ini, loyalitas berbahasa Indonesia dimaknai sebagai kemauan siswa untuk menempatkan bahasa Indonesia secara tepat, terutama pada situasi sekolah, komunikasi resmi, dan ruang digital yang berkaitan dengan kegiatan akademik.",
        "Bahasa remaja di media sosial memperlihatkan perubahan pilihan kata, campur kode, dan ragam santai. Penelitian tentang pengaruh media sosial terhadap bahasa remaja menunjukkan bahwa media digital memengaruhi gaya penulisan bahasa Indonesia (Prasetyaningrum, 2024) dan penggunaan ragam informal di kalangan remaja (Parlindungan Siahaan dkk., 2024). Temuan tersebut menjadi dasar bahwa pembinaan bahasa perlu dekat dengan kebiasaan komunikasi remaja.",
        "Media digital dalam penelitian ini dipahami sebagai sarana pedagogis untuk menghadirkan kasus bahasa dan refleksi, bukan sebagai objek utama rekayasa komputer. Dengan demikian, pengembangan PRIMA+ tetap ditempatkan dalam kajian sosial-humaniora, terutama pembelajaran bahasa, sikap bahasa, dan perilaku komunikasi remaja.",
    ]:
        add_p(doc, text)

    add_heading(doc, "2.2 STUDI PUSTAKA", 2)
    add_table(
        doc,
        ["No.", "Sumber", "Temuan Utama", "Posisi terhadap PRIMA+"],
        [
            ["1", "Carter (2003); Fairclough (1992)", "Kesadaran berbahasa berkaitan dengan bentuk, fungsi, konteks, dan sikap kritis.", "Menjadi dasar latihan refleksi pilihan bahasa."],
            ["2", "Garvin dan Mathiot (1968); Chaer dan Agustina (2014)", "Loyalitas bahasa merupakan bagian dari sikap bahasa dan berkaitan dengan norma serta kebanggaan bahasa.", "Menjadi dasar konstruk loyalitas berbahasa Indonesia."],
            ["3", "Widyaningrum dkk. (2020)", "Remaja Kota Sampit dapat dikaji melalui konteks media, identitas, dan pengalaman sosial lokal.", "Menegaskan konteks lokal Sampit dan kategori sosial-humaniora."],
            ["4", "Prasetyaningrum (2024); Parlindungan Siahaan dkk. (2024)", "Media sosial memengaruhi gaya bahasa dan ragam informal remaja.", "Memperkuat masalah bahasa digital remaja."],
            ["5", "Shortt dkk. (2021); Luo (2023)", "Gamifikasi dapat mendukung pembelajaran bahasa jika desain dan evaluasinya jelas.", "Menjadi pendukung bentuk media interaktif PRIMA+."],
        ],
        caption="Tabel 1. Studi pustaka terdahulu",
        widths=[0.45, 1.55, 2.35, 2.15],
    )

    # ── BAB 3. METODE PENELITIAN ──
    add_heading(doc, "BAB 3. METODE PENELITIAN", 1, size=16)
    add_heading(doc, "3.1 WAKTU DAN TEMPAT PENELITIAN", 2)
    add_p(doc, "Penelitian dilaksanakan pada April hingga Agustus 2026 di MAN Kotawaringin Timur, Sampit, Kabupaten Kotawaringin Timur, Kalimantan Tengah. Kegiatan penelitian meliputi analisis kebutuhan, perancangan media PRIMA+, penyusunan instrumen, validasi, uji coba terbatas, pengolahan data, revisi media, dan penyusunan laporan.")

    add_heading(doc, "3.2 ALAT DAN BAHAN", 2)
    add_table(
        doc,
        ["No.", "Alat/Bahan", "Jumlah", "Fungsi"],
        [
            ["1", "Laptop peneliti", "1 unit", "Menyusun materi, mengolah data, dan mengembangkan PRIMA+."],
            ["2", "Perangkat siswa", "30 perangkat/sesuai ketersediaan", "Mengakses PRIMA+ (web/PWA) saat uji coba terbatas."],
            ["3", "ReactJS (frontend)", "1 framework", "Membangun antarmuka web interaktif dan PWA untuk PRIMA+."],
            ["4", "Express.js + Node.js (backend)", "1 sistem", "Mengelola API, autentikasi, dan logika aplikasi PRIMA+."],
            ["5", "Firebase (database & autentikasi)", "1 layanan", "Menyimpan data pengguna, skor pretest-posttest, dan progress."],
            ["6", "Gemini API (LLM)", "1 layanan", "Menghasilkan soal kuis dinamis, feedback otomatis, dan refleksi bahasa."],
            ["7", "Google AI Studio", "1 platform", "Mengembangkan dan menguji prompt Gemini API untuk konten PRIMA+."],
            ["8", "Skenario kasus bahasa remaja", "1 paket", "Materi latihan pilihan ragam bahasa dalam PRIMA+."],
            ["9", "Kuesioner loyalitas berbahasa", "20 butir", "Mengukur sikap positif, kesetiaan penggunaan, kesadaran norma, kebanggaan, dan pemilihan ragam."],
            ["10", "Lembar validasi guru/ahli", "2-3 validator", "Menilai kelayakan materi, bahasa, media, dan instrumen."],
        ],
        caption="Tabel 2. Alat dan bahan penelitian",
        widths=[0.45, 1.7, 1.25, 3.1],
    )

    add_heading(doc, "3.3 PROSEDUR PENELITIAN", 2)
    add_p(doc, "Penelitian menggunakan pendekatan Research and Development (R&D) sederhana dengan model ADDIE. Model ini dipilih karena penelitian bertujuan mengembangkan media pembelajaran/kebahasaan dan menguji penggunaannya secara terbatas. Desain uji efektivitas menggunakan one-group pretest-posttest, yaitu siswa mengisi kuesioner loyalitas berbahasa sebelum dan sesudah menggunakan PRIMA+.")
    add_p(doc, "Subjek uji coba terbatas berjumlah 30 siswa kelas X atau XI MAN Kotawaringin Timur. Jumlah 30 siswa dipilih karena penelitian ini bukan survei populasi besar, melainkan uji coba terbatas media dalam penelitian pengembangan. Dengan 30 siswa, peneliti memperoleh 30 pasangan data pretest-posttest sehingga analisis peningkatan dapat dilakukan melalui rata-rata, persentase peningkatan, N-gain, serta uji beda berpasangan.")
    add_p(doc, "Teknik pengambilan sampel menggunakan purposive sampling. Kriteria responden adalah siswa kelas X atau XI, aktif mengikuti pembelajaran Bahasa Indonesia, menggunakan media digital dalam komunikasi sehari-hari, bersedia mengikuti pretest, uji coba PRIMA+, dan posttest, serta memperoleh izin dari sekolah sesuai ketentuan penelitian yang melibatkan manusia.")
    for text in [
        "Analyze: menganalisis kebutuhan siswa, contoh bahasa remaja di lingkungan sekolah dan media digital, serta indikator loyalitas berbahasa.",
        "Design: merancang skenario kasus bahasa, kisi-kisi instrumen, alur penggunaan PRIMA+, dan rubrik validasi.",
        "Develop: menyusun media PRIMA+ sebagai aplikasi web berbasis ReactJS dengan backend Express.js, database Firebase, serta integrasi Gemini API untuk menghasilkan soal kuis dinamis, feedback otomatis, dan refleksi pilihan bahasa.",
        "Implement: melaksanakan pretest, uji coba PRIMA+ kepada 30 siswa, posttest, dan pengumpulan respons siswa.",
        "Evaluate: menganalisis data, menafsirkan hasil, dan merevisi media berdasarkan validasi serta respons siswa.",
    ]:
        add_list(doc, text)
    if os.path.exists(FIGURE):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(FIGURE, width=Inches(5.6))
            add_center(doc, "Gambar 1. Alur prosedur penelitian PRIMA+ berbasis ADDIE", size=10, bold=False)
        except Exception:
            add_center(doc, "Gambar 1. Alur prosedur penelitian PRIMA+ berbasis ADDIE", size=10, bold=False)

    add_table(
        doc,
        ["Konstruk/Aspek", "Definisi", "Indikator", "Instrumen"],
        [
            ["Kesadaran berbahasa", "Kepekaan terhadap bentuk, fungsi, konteks, dan dampak sosial bahasa.", "Mengenali ragam; menilai konteks; memberi alasan pilihan bahasa.", "Kuis PRIMA+ dan refleksi."],
            ["Loyalitas berbahasa Indonesia", "Kemauan menempatkan bahasa Indonesia secara tepat sebagai identitas dan alat komunikasi.", "Sikap positif; kesetiaan penggunaan; kesadaran norma; kebanggaan bahasa; pemilihan ragam.", "Kuesioner Likert 20 butir."],
            ["Kelayakan media", "Kelayakan PRIMA+ sebagai media pembelajaran/kebahasaan.", "Materi; bahasa; tampilan; navigasi; umpan balik.", "Lembar validasi guru/ahli."],
            ["Respons pengguna", "Tanggapan siswa setelah menggunakan PRIMA+.", "Kemudahan; kemenarikan; manfaat; kendala.", "Angket respons/refleksi."],
        ],
        caption="Tabel 3. Konstruk/aspek, indikator, dan instrumen",
        widths=[1.35, 1.75, 2.0, 1.4],
    )

    add_heading(doc, "3.4 PENGOLAHAN DATA", 2)
    add_p(doc, "Pengolahan data disesuaikan dengan tiga rumusan masalah dan tujuan penelitian. Rumusan masalah pertama dijawab melalui deskripsi proses pengembangan media PRIMA+ berdasarkan tahapan ADDIE. Rumusan masalah kedua dijawab melalui persentase kelayakan dari validasi guru atau ahli. Rumusan masalah ketiga dijawab melalui perbandingan skor pretest dan posttest loyalitas berbahasa Indonesia.")
    add_table(
        doc,
        ["Data", "Teknik Pengolahan", "Interpretasi"],
        [
            ["Validasi media", "Persentase kelayakan", "Kategori sangat layak, layak, cukup, atau perlu revisi."],
            ["Pretest-posttest", "Rata-rata, persentase peningkatan, dan N-gain", "Posttest lebih tinggi menunjukkan peningkatan awal loyalitas berbahasa."],
            ["Uji beda berpasangan", "Paired sample t-test jika selisih data berdistribusi normal; Wilcoxon signed-rank test jika tidak normal", "Peningkatan dinilai bermakna bila hasil uji menunjukkan perbedaan yang signifikan."],
            ["Refleksi siswa", "Analisis tema ringkas", "Menjelaskan pengalaman, kendala, dan saran pengguna."],
        ],
        caption="Tabel 4. Rancangan pengolahan data",
        widths=[1.45, 2.6, 2.45],
    )

    # ── BAB 4. HASIL DAN PEMBAHASAN (kerangka; data belum terkumpul) ──
    add_heading(doc, "BAB 4. HASIL DAN PEMBAHASAN", 1)
    add_p(doc, "Bab ini menyajikan hasil penelitian dan pembahasan berdasarkan tiga rumusan masalah yang telah ditetapkan. Data diperoleh melalui validasi ahli, uji coba terbatas kepada 30 siswa MAN Kotawaringin Timur, serta analisis pretest-posttest loyalitas berbahasa Indonesia. Bagian ini masih berupa kerangka; angka hasil akan diisi setelah pengumpulan dan pengolahan data.")

    add_heading(doc, "4.1 HASIL PENGEMBANGAN MEDIA PRIMA+", 2)
    add_p(doc, "Bagian ini mendeskripsikan proses pengembangan media PRIMA+ berdasarkan tahapan ADDIE, mulai dari analisis kebutuhan siswa, perancangan skenario kasus bahasa, penyusunan konten kuis dan refleksi, hingga validasi dan uji coba terbatas.")
    add_p(doc, "[Diisi nanti: narasi hasil tiap tahap ADDIE — analisis kebutuhan bahasa remaja di lingkungan sekolah, desain skenario kasus bahasa dan kisi-kisi instrumen, penyusunan konten kuis dan refleksi, hasil validasi, serta hasil uji coba terbatas.]")

    add_heading(doc, "4.2 HASIL VALIDASI MEDIA PRIMA+", 2)
    add_p(doc, "Validasi dilakukan oleh guru atau ahli terhadap aspek materi, bahasa, media, dan instrumen penelitian. Hasil validasi disajikan pada Tabel 5.")
    add_table(
        doc,
        ["Aspek Validasi", "Skor", "Kategori"],
        [
            ["Materi", "[Diisi nanti: skor]", "[Diisi nanti: kategori kelayakan]"],
            ["Bahasa", "[Diisi nanti: skor]", "[Diisi nanti: kategori kelayakan]"],
            ["Media", "[Diisi nanti: skor]", "[Diisi nanti: kategori kelayakan]"],
            ["Instrumen", "[Diisi nanti: skor]", "[Diisi nanti: kategori kelayakan]"],
        ],
        caption="Tabel 5. Hasil validasi media PRIMA+",
        widths=[1.8, 1.2, 2.5],
    )
    add_p(doc, "[Diisi nanti: narasi hasil validasi — skor rata-rata, persentase kelayakan, dan catatan/saran validator.]")

    add_heading(doc, "4.3 HASIL UJI EFEKTIVITAS PRIMA+", 2)
    add_p(doc, "Efektivitas PRIMA+ diukur melalui perbandingan skor pretest dan posttest loyalitas berbahasa Indonesia pada 30 siswa. Hasil perbandingan disajikan pada Tabel 6.")
    add_table(
        doc,
        ["Indikator", "Rata-rata Pretest", "Rata-rata Posttest", "Peningkatan", "N-gain"],
        [
            ["Sikap positif", "[Diisi nanti: skor]", "[Diisi nanti: skor]", "[Diisi nanti: %]", "[Diisi nanti: kategori]"],
            ["Kesetiaan penggunaan", "[Diisi nanti: skor]", "[Diisi nanti: skor]", "[Diisi nanti: %]", "[Diisi nanti: kategori]"],
            ["Kesadaran norma", "[Diisi nanti: skor]", "[Diisi nanti: skor]", "[Diisi nanti: %]", "[Diisi nanti: kategori]"],
            ["Kebanggaan bahasa", "[Diisi nanti: skor]", "[Diisi nanti: skor]", "[Diisi nanti: %]", "[Diisi nanti: kategori]"],
            ["Pemilihan ragam", "[Diisi nanti: skor]", "[Diisi nanti: skor]", "[Diisi nanti: %]", "[Diisi nanti: kategori]"],
            ["Rata-rata total", "[Diisi nanti: skor]", "[Diisi nanti: skor]", "[Diisi nanti: %]", "[Diisi nanti: kategori]"],
        ],
        caption="Tabel 6. Perbandingan skor pretest dan posttest",
        widths=[1.5, 1.2, 1.2, 1.0, 1.0],
    )
    add_p(doc, "[Diisi nanti: narasi hasil uji efektivitas — perbandingan rata-rata pretest-posttest, persentase peningkatan, hasil uji beda berpasangan (paired t-test/Wilcoxon), dan interpretasi N-gain.]")

    add_heading(doc, "4.4 PEMBAHASAN", 2)
    add_p(doc, "[Diisi nanti: pembahasan hasil penelitian dikaitkan dengan teori dan studi pustaka Bab 2; perbandingan dengan penelitian terdahulu; implikasi temuan terhadap penguatan loyalitas berbahasa Indonesia remaja.]")

    # ── BAB 5. ANGGARAN BIAYA DAN KEGIATAN ──
    add_heading(doc, "BAB 5. ANGGARAN BIAYA DAN KEGIATAN", 1, size=14)
    add_heading(doc, "5.1 ANGGARAN BIAYA PENELITIAN", 2)
    add_p(doc, "Bagian ini menyajikan rencana dan realisasi anggaran biaya penelitian. Anggaran rencana disusun sesuai batas yang ditetapkan, sedangkan kolom realisasi diisi setelah seluruh kegiatan penelitian terlaksana.")
    add_table(
        doc,
        ["No.", "Komponen Pengeluaran", "Anggaran Rencana (Rp)", "Realisasi (Rp)"],
        [
            ["1", "Barang: kertas HVS A4 80 gsm (2 rim)", "112.000", "[Diisi nanti: realisasi]"],
            ["2", "Barang: tinta printer hitam refill (1 botol)", "62.500", "[Diisi nanti: realisasi]"],
            ["3", "Barang: tinta printer warna refill (1 botol)", "78.500", "[Diisi nanti: realisasi]"],
            ["4", "Barang: map plastik/snelhecter untuk berkas responden (35 lembar)", "52.500", "[Diisi nanti: realisasi]"],
            ["5", "Barang: label kode responden (2 pak)", "24.000", "[Diisi nanti: realisasi]"],
            ["6", "Barang: pulpen untuk pengisian instrumen (2 lusin)", "48.000", "[Diisi nanti: realisasi]"],
            ["7", "Barang: binder clip, staples, dan isolasi arsip (1 paket)", "35.750", "[Diisi nanti: realisasi]"],
            ["8", "Barang: flashdisk 32 GB untuk cadangan data penelitian (1 unit)", "44.850", "[Diisi nanti: realisasi]"],
            ["9", "Barang: materai Rp10.000 untuk berkas izin/validasi (4 lembar)", "40.000", "[Diisi nanti: realisasi]"],
            ["10", "Barang: amplop cokelat arsip dokumen (10 lembar)", "18.500", "[Diisi nanti: realisasi]"],
            ["11", "Jasa: fotokopi/cetak kuesioner pretest-posttest (60 set)", "210.000", "[Diisi nanti: realisasi]"],
            ["12", "Jasa: cetak lembar validasi dan angket respons (40 set)", "137.500", "[Diisi nanti: realisasi]"],
            ["13", "Jasa: cetak kartu skenario kasus bahasa PRIMA+ (30 set)", "185.000", "[Diisi nanti: realisasi]"],
            ["14", "Jasa: jilid laporan proposal/lampiran (3 eksemplar)", "90.000", "[Diisi nanti: realisasi]"],
            ["15", "Sewa/akses: paket data internet 25 GB untuk koordinasi dan uji coba (1 nomor)", "135.000", "[Diisi nanti: realisasi]"],
            ["16", "Sewa: modem Wi-Fi portabel saat uji coba (1 hari)", "85.000", "[Diisi nanti: realisasi]"],
            ["17", "Sewa: tripod HP untuk dokumentasi kegiatan (1 hari)", "45.000", "[Diisi nanti: realisasi]"],
            ["18", "Lainnya: transport lokal ke sekolah, validator, dan percetakan (4 kali perjalanan)", "245.000", "[Diisi nanti: realisasi]"],
            ["19", "Lainnya: konsumsi ringan validator dan siswa uji coba (35 paket)", "280.000", "[Diisi nanti: realisasi]"],
            ["20", "Lainnya: air mineral gelas untuk kegiatan validasi/uji coba (2 dus)", "58.000", "[Diisi nanti: realisasi]"],
            ["", "Total", "1.987.100", "[Diisi nanti: total realisasi]"],
        ],
        caption="Tabel 7. Rencana dan realisasi anggaran biaya penelitian",
        widths=[0.45, 4.35, 1.3, 1.3],
    )
    add_p(doc, "[Diisi nanti: ringkasan realisasi per kelompok pengeluaran dan selisih terhadap rencana.]")

    add_heading(doc, "5.2 KEGIATAN", 2)
    add_p(doc, "Bagian ini menyajikan jadwal pelaksanaan penelitian. Kolom realisasi dan bukti isi diisi setelah kegiatan dilaksanakan.")
    add_table(
        doc,
        ["No.", "Kegiatan", "Waktu", "Status", "Bukti"],
        [
            ["1", "Analisis kebutuhan dan penyusunan rancangan media PRIMA+", "April 2026", "[Diisi nanti]", "[Diisi nanti]"],
            ["2", "Pengembangan media dan penyusunan instrumen", "Mei 2026", "[Diisi nanti]", "[Diisi nanti]"],
            ["3", "Validasi ahli dan revisi media", "Juni 2026", "[Diisi nanti]", "[Diisi nanti]"],
            ["4", "Uji coba terbatas (pretest, penggunaan PRIMA+, posttest)", "Juli 2026", "[Diisi nanti]", "[Diisi nanti]"],
            ["5", "Pengolahan data dan penyusunan laporan", "Juli-Agustus 2026", "[Diisi nanti]", "[Diisi nanti]"],
        ],
        caption="Tabel 8. Jadwal dan realisasi kegiatan penelitian",
        widths=[0.45, 3.6, 1.4, 1.2, 1.2],
    )
    add_p(doc, "[Diisi nanti: status penyelesaian dan dokumentasi setiap kegiatan.]")

    # ── KESIMPULAN DAN SARAN ──
    add_heading(doc, "KESIMPULAN DAN SARAN", 1, size=14)
    add_heading(doc, "KESIMPULAN", 2)
    add_p(doc, "[Diisi nanti: kesimpulan berdasarkan tiga rumusan masalah — (1) rancangan media PRIMA+ berbasis kesadaran berbahasa, (2) kelayakan PRIMA+ berdasarkan validasi, (3) efektivitas PRIMA+ dalam meningkatkan loyalitas berbahasa Indonesia.]")

    add_heading(doc, "SARAN", 2)
    add_p(doc, "[Diisi nanti: saran untuk pengembangan media selanjutnya, saran bagi guru dan sekolah dalam memanfaatkan PRIMA+, serta saran bagi peneliti berikutnya.]")

    # ── UCAPAN TERIMA KASIH ──
    add_heading(doc, "UCAPAN TERIMA KASIH", 1)
    add_p(doc, "Peneliti menyampaikan terima kasih kepada pihak-pihak yang membantu terlaksananya penelitian ini, antara lain kepala madrasah, guru Bahasa Indonesia, siswa yang menjadi subjek uji coba, validator, serta pihak yang mendukung pendanaan dan fasilitas pelaksanaan penelitian. [Lengkapi nama/lembaga yang sesuai.]")

    # ── PERNYATAAN PENGGUNAAN AI ──
    add_heading(doc, "PERNYATAAN PENGGUNAAN AI", 1)
    add_p(doc, "Sesuai ketentuan Panduan OPSI 2026, penelitian ini menggunakan perangkat AI jenis [Google Gemini / LLM chatbot, tuliskan nama dan versi] sebagai alat bantu untuk: menemukan dan memahami referensi, menyusun ide, serta membantu pengembangan konten kuis dan refleksi PRIMA+. Penggunaan AI tersebut telah dicantumkan pada Surat Pernyataan Penggunaan AI (lampiran). Isi, analisis, interpretasi data, dan kesimpulan naskah ditulis dan dipertanggungjawabkan sendiri oleh peserta.")

    # ── DAFTAR PUSTAKA ──
    add_heading(doc, "DAFTAR PUSTAKA", 1)
    refs = [
        "Balai Bahasa Provinsi Kalimantan Tengah. (2026). Layanan Uji Kemahiran Berbahasa Indonesia (UKBI). https://balaibahasakalteng.kemendikdasmen.go.id/layanan/ukbi/",
        "Branch, R. M. (2009). Instructional Design: The ADDIE Approach. Springer. https://doi.org/10.1007/978-0-387-09506-6",
        "Carter, R. (2003). Language awareness. ELT Journal, 57(1), 64-65. https://doi.org/10.1093/elt/57.1.64",
        "Chaer, A., & Agustina, L. (2014). Sosiolinguistik: Perkenalan awal. Rineka Cipta.",
        "Crystal, D. (2011). Internet Linguistics: A Student Guide. Routledge.",
        "Fairclough, N. (1992). Critical Language Awareness. Longman.",
        "Garvin, P. L., & Mathiot, M. (1968). The urbanization of the Guarani language: A problem in language and culture. In J. A. Fishman (Ed.), Readings in the Sociology of Language (pp. 365-374). De Gruyter Mouton. https://doi.org/10.1515/9783110805376.365",
        "Kementerian Pendidikan Dasar dan Menengah. (2026). Informasi Satuan Pendidikan: MAN Kotawaringin Timur, NPSN 30201526. https://referensi.data.kemendikdasmen.go.id/snpmb/site/sekolah?npsn=30201526",
        "Luo, Z. (2023). The effectiveness of gamified tools for foreign language learning (FLL): A systematic review. Behavioral Sciences, 13(4), 331. https://doi.org/10.3390/bs13040331",
        "Parlindungan Siahaan, A., Aldy Pradana, M., Citra Chairani, D., Heriyani Erizal, A., & Margareta Lase, Y. (2024). Pengaruh era digital terhadap pemakaian bahasa Indonesia di kalangan remaja melalui media sosial. PENG: Jurnal Ekonomi dan Manajemen, 2(1), 879-885. https://teewanjournal.com/index.php/peng/article/view/1026",
        "Prasetyaningrum, R. (2024). Pengaruh media sosial terhadap gaya bahasa dalam penulisan bahasa Indonesia pada remaja. Jurnal Sosial Humaniora dan Pendidikan, 3(1), 127-134. https://doi.org/10.55606/inovasi.v3i1.2734",
        "Shortt, M., Tilak, S., Kuznetcova, I., Martens, B., & Akinkuolie, B. (2021). Gamification in mobile-assisted language learning: A systematic review of Duolingo literature from public release of 2012 to early 2020. Computer Assisted Language Learning, 36(3), 517-554. https://doi.org/10.1080/09588221.2021.1933540",
        "Sugiyono. (2019). Metode Penelitian dan Pengembangan: Research and Development. Alfabeta.",
        "UKBI. (2026). Uji Kemahiran Berbahasa Indonesia. Kementerian Pendidikan Dasar dan Menengah. https://ukbi.kemendikdasmen.go.id/",
        "Widyaningrum, A. Y., Yumarnamto, M., & Prijambodo, V. L. (2020). Analisis resepsi remaja Kota Sampit mengenai keberagaman di media. WACANA: Jurnal Ilmiah Ilmu Komunikasi, 19(1), 51-61.",
    ]
    for ref in refs:
        add_p(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT, after=3)

# ── LAMPIRAN ──
    add_heading(doc, "LAMPIRAN", 1)
    add_p(doc, "Lampiran 1. Kisi-kisi instrumen loyalitas berbahasa Indonesia (lihat Instrumen_Penelitian_PRIMA.docx)", bold=True)
    add_p(doc, "Lampiran 2. Lembar validasi guru/ahli (lihat Instrumen_Penelitian_PRIMA.docx)", bold=True)
    add_p(doc, "Lampiran 3. Angket respons/refleksi siswa (lihat Instrumen_Penelitian_PRIMA.docx)", bold=True)
    add_p(doc, "Lampiran 4. Surat pengantar kepada responden dan formulir persetujuan mengikuti penelitian (informed consent, sesuai Lampiran 7-8 Panduan OPSI)", bold=True)
    add_p(doc, "Lampiran 5. Logbook penelitian (sesuai Lampiran 9 Panduan OPSI)", bold=True)
    add_p(doc, "Lampiran 6. Tampilan antarmuka PRIMA+", bold=True)
    add_p(doc, "Lampiran 7. Dokumentasi kegiatan uji coba", bold=True)
    add_p(doc, "Lampiran 8. Surat pernyataan orisinalitas karya", bold=True)
    add_p(doc, "Lampiran 9. Surat pernyataan penggunaan AI (sesuai Lampiran 2 Panduan OPSI)", bold=True)

    doc.save(OUT)
    print("OK ->", OUT)


if __name__ == "__main__":
    build()